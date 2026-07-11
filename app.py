"""
Criterion Assessment Metrics (CAM) - unified Streamlit dashboard.

The master suite of the CAM grading ecosystem: this dashboard (formerly
"Art Criterion Metrics" / ACM) plus the Flask grading sub-app that now lives
in ./cam_grading_workspace (formerly "GCG"). A wide three-window cockpit over
the ``engine`` package (models / ingestion / aggregation / criterion_d /
rubrics). Run with:

    streamlit run app.py

Layout (st.columns([4, 3, 5])):
    Window 1  Timeline & assignment system   (staging, sync, dates, toggles)
    Window 2  Roster & submission tracker     (order, missing-work highlights)
    Window 3  Evaluation cockpit              (edit, trend, comments, AI deck)
    Tray      Deliverables (Excel, report-card pack, class comments)

All mutable interaction state lives in st.session_state and is initialised
defensively in init_state() so reruns never KeyError.
"""

from __future__ import annotations

import atexit
import csv
import glob
import hashlib
import io
import json
import os
import shutil
import socket
import subprocess
import sys
import tempfile
import time
import webbrowser
from datetime import date, datetime
from collections import Counter
from statistics import pstdev, mean
from typing import Optional
from urllib.parse import urlencode
from urllib.request import Request, urlopen

import streamlit as st

from engine import (
    Criterion,
    CriterionScore,
    Gradebook,
    Assignment,
    IngestionPipeline,
    CriterionDInitializer,
    CALCULATION_METHODS,
    METHOD_LABELS,
    myp_grade,
    school_grade,
    aggregate_student_criterion,
    map_criterion_columns,
    detect_criterion_in_header,
    clean_assignment_name,
    is_exam_csv,
    parse_iso_date,
    parse_date_from_filename,
    exam_question_columns,
    parse_unit_plan,
    parse_classroom_roster,
    gojuon_sort_key,
    student_id_from_email,
    DEFAULT_DB_FILENAME,
    save_database,
    load_database,
    db_file_state,
    unit_plan_to_dict,
    unit_plan_from_dict,
    Evidence,
    select_evidence,
    trend_for_series,
    format_trend_sentence,
    load_term_summaries,
    load_class_mirror,
    save_class_mirror,
    score_to_dict,
    score_from_dict,
    assignment_to_dict,
    assignment_from_dict,
    exam_result_to_dict,
    exam_result_from_dict,
)
# METHOD_60_40 / METHOD_WEIGHTED_MEDIAN are the two methods the auto default
# picks between; they live in the aggregation module (not re-exported by the
# package) and this keeps all changes Streamlit-side per the plan.
from engine.aggregation import METHOD_60_40, METHOD_WEIGHTED_MEDIAN
from engine import rubrics

# --------------------------------------------------------------------------
# Constants
# --------------------------------------------------------------------------

CRIT_COLORS = {"A": "#1f77b4", "B": "#ff7f0e", "C": "#2ca02c", "D": "#9467bd"}
CRIT_ORDER = ["A", "B", "C", "D"]

# Effort / English Use: per-term teacher filler grade feeding the School
# grade lookup. 4 = normal, 3 = behaviour issues, 5 = exceptional standout.
EFFORT_DEFAULT = 4
# Academic year runs April -> March. Map a month to its position in the year.
ACADEMIC_MONTHS = [4, 5, 6, 7, 8, 9, 10, 11, 12, 1, 2, 3]
# School terms (Japanese 3-term year). Assignments are tagged with one of these;
# the prompt engine treats earlier terms as "past" (summarised, not raw).
TERMS = ["Term 1", "Term 2", "Term 3"]

# Explicit term backup files (docs/TERM_BACKUP_RESTORE_PLAN.md): a teacher-
# initiated, self-describing snapshot of one term written OUTSIDE the database,
# and the loader that can put that term back after a database disaster.
TERM_BACKUP_VERSION = 1
TERM_BACKUP_KIND = "cam_term_backup"

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# The Flask grading sub-app (formerly GCG) now lives inside CAM. The bridge
# spawns it on its own port so it never blocks the Streamlit server.
GRADING_WORKSPACE_DIR = os.path.join(BASE_DIR, "cam_grading_workspace")
GRADING_PORT = 5001

# Window-2 unmatched-work thumbnails (Phase 4). Rendered first-page/image PNGs
# are disk-cached here, keyed by source path + mtime + width, so re-opening the
# matching dialog is instant. Gitignored (thumb_cache/) and deliberately beside
# app.py — NEVER inside the (possibly OneDrive-synced) db folder. Safe to
# delete; it repopulates on next open.
THUMB_CACHE_DIR = os.path.join(BASE_DIR, "thumb_cache")
THUMB_GRID_WIDTH = 400        # grid-tile render width (px)
THUMB_ENLARGE_WIDTH = 1600    # click-to-enlarge render width (px)
# File extensions rendered directly by Pillow (everything else falls to fitz for
# PDFs, or a filename-only tile); mirrors exam_engine's IMAGE_EXTS.
THUMB_IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff",
                    ".gif"}

# Gender selection (Window 2) -> report-comment pronouns. Blank/unknown is
# treated inclusively as they/them.
GENDER_OPTIONS = ["Male", "Female", "Non-Binary"]
GENDER_PRONOUNS = {"Male": "he/him", "Female": "she/her", "Non-Binary": "they/them"}

# ---- Device-local UI preferences -----------------------------------------
# These live in a SEPARATE file from the shared cloud database so a desktop and
# a laptop can keep their own column widths / scroll heights while pointing the
# main acm_database.json at the same OneDrive / Google Drive folder.
PREFS_FILENAME = "local_device_prefs.json"
PREFS_PATH = os.path.join(BASE_DIR, PREFS_FILENAME)
DEFAULT_PREFS = {
    "db_custom_path": "",   # blank -> acm_database.json beside app.py
    # Folder the explicit end-of-term backups are written to. Per-device (like
    # every path pref) — the teacher may point it at a USB stick or a non-cloud
    # folder for an off-site copy. Blank until they choose one in ⚙ Settings.
    "term_backup_folder": "",
    "col_w1": 4, "col_w2": 3, "col_w3": 5,   # 3-column width ratios
    "h1": 520, "h2": 520, "h3": 640,         # per-window scroll heights (px)
    "h_remarks": 80, "h_comment": 90,        # Window 3 editable text-area heights (px)
    # Phase-4 first-boot bootstrap marker: False (or a missing prefs file) means
    # this machine has not yet chosen a data home, so CAM shows the one-time
    # setup panel instead of silently booting the sample DB. Set True by any of
    # the panel's choices (adopt a discovered DB / use another folder / start
    # fresh) so the panel never nags again — even for "start fresh", which keeps
    # db_custom_path blank. See docs/CROSS_DEVICE_AND_DB_SAFETY_PLAN.md Phase 4.
    "setup_done": False,
}


def load_prefs() -> dict:
    """Read device-local UI prefs, merged over defaults (never raises)."""
    prefs = dict(DEFAULT_PREFS)
    try:
        with open(PREFS_PATH, "r", encoding="utf-8") as fh:
            saved = json.load(fh)
        if isinstance(saved, dict):
            prefs.update({k: saved[k] for k in DEFAULT_PREFS if k in saved})
    except (OSError, ValueError):
        pass
    return prefs


def save_prefs(prefs: dict) -> None:
    """Persist device-local UI prefs to PREFS_PATH (never raises)."""
    try:
        with open(PREFS_PATH, "w", encoding="utf-8") as fh:
            json.dump(prefs, fh, ensure_ascii=False, indent=2)
    except OSError:
        pass


def resolve_db_path(custom: str) -> str:
    """Resolve a custom-path preference value to a concrete database file path.

    Smart interpretation: a value ending in ``.json`` is used verbatim as the
    file path; any other non-empty value is treated as a directory and
    ``acm_database.json`` is placed inside it (the OneDrive/Drive folder case).
    Blank falls back to a file beside this script. Pure — takes the pref value
    directly so callers (e.g. the Phase-2 settings panel) can resolve a
    *candidate* path without mutating the active pref first.
    """
    custom = (custom or "").strip()
    if not custom:
        return os.path.join(BASE_DIR, DEFAULT_DB_FILENAME)
    if custom.lower().endswith(".json"):
        return custom
    return os.path.join(custom, DEFAULT_DB_FILENAME)


def db_path() -> str:
    """The active database path, resolved from the current custom-path pref.

    The ``CAM_DB_PATH`` environment variable overrides the pref entirely (Phase
    4): a one-liner new-machine setup, and — crucially — a way for tests and
    harnesses to force a sandbox path that **cannot** fall through to the real
    device prefs (which point at live OneDrive data). Set it and the app can
    never resolve the teacher's real data folder, closing the ``.wiped-by-test``
    hazard class. It is resolved through the same folder-or-.json logic as the
    pref. See docs/CROSS_DEVICE_AND_DB_SAFETY_PLAN.md Phase 4."""
    env = os.environ.get("CAM_DB_PATH", "").strip()
    if env:
        return resolve_db_path(env)
    return resolve_db_path(st.session_state.get("prefs", DEFAULT_PREFS).get("db_custom_path"))


def db_folder() -> str:
    """The folder holding the active database — the CAM data home.

    Every class gets its own subfolder here (grading exports, exam scans,
    grading caches, finalized term summaries), so all of a class's data
    travels together inside the one cloud-synced directory."""
    return os.path.dirname(os.path.abspath(db_path())) or BASE_DIR


# --- Phase-4 cross-device bootstrap: discover an existing cloud database -----
# A second computer should reach the shared database without hand-copying
# local_device_prefs.json. The first-boot setup panel probes well-known cloud-
# mirror roots for a folder holding acm_database.json and offers each as a
# candidate; the teacher recognises the real one by its assignment/roster counts.
# Convenience only — the safety guards (Phases 1-3) are what actually protect the
# file; this just saves the manual path-paste. USB/removable drives are
# deliberately NOT auto-scanned (Phase 1's storage-missing quarantine is the
# safety net for a forgotten drive); the teacher points at those via "Use another
# folder". See docs/CROSS_DEVICE_AND_DB_SAFETY_PLAN.md Phase 4.

# Directory names never worth descending into during the shallow scan — noisy,
# huge, or system trees that will not hold a teacher's cloud data folder.
_SCAN_SKIP_DIRS = {
    "node_modules", "__pycache__", ".git", ".svn", "$recycle.bin",
    "system volume information", "windows", "program files",
    "program files (x86)", "programdata", "appdata", "thumb_cache",
    ".thumbnails", "cache", ".cache", "venv", ".venv", "site-packages",
}
# Hard cap on directories visited per root so a pathologically large cloud tree
# can never hang the first-boot panel; depth ≤ 3 already bounds it in practice.
_SCAN_MAX_DIRS = 6000


def _cloud_search_roots() -> list:
    """Well-known local cloud-mirror roots to probe for a CAM data folder.

    OneDrive, Google Drive for Desktop and Dropbox, discovered from their
    standard environment variables / mount points / config, de-duplicated and
    filtered to those that actually exist on this machine. Best-effort — a
    missing service simply contributes no roots."""
    roots: list = []

    def add(path: str) -> None:
        if not path:
            return
        try:
            if os.path.isdir(path):
                key = os.path.normcase(os.path.abspath(path))
                if key not in {os.path.normcase(os.path.abspath(r)) for r in roots}:
                    roots.append(path)
        except OSError:
            pass

    # OneDrive (personal + the two business-tenant variables).
    for var in ("OneDrive", "OneDriveConsumer", "OneDriveCommercial"):
        add(os.environ.get(var, ""))
    home = os.path.expanduser("~")
    # Google Drive for Desktop: the virtualised "My Drive" under the profile,
    # a legacy "Google Drive" folder, and the DriveFS letter mounts (G:\My Drive
    # etc. — Drive defaults to G: but the letter is user-configurable).
    add(os.path.join(home, "My Drive"))
    add(os.path.join(home, "Google Drive"))
    for letter in "GHIJKLMNOPQRSTUVWXYZ":
        add(f"{letter}:\\My Drive")
    # Dropbox: the default profile folder, plus any account paths recorded in
    # info.json (Dropbox's own location file) when the folder was moved.
    add(os.path.join(home, "Dropbox"))
    try:
        info = os.path.join(os.environ.get("LOCALAPPDATA", ""), "Dropbox", "info.json")
        with open(info, "r", encoding="utf-8") as fh:
            data = json.load(fh)
        if isinstance(data, dict):
            for acct in data.values():
                if isinstance(acct, dict):
                    add(acct.get("path", ""))
    except (OSError, ValueError):
        pass
    return roots


def _scan_for_db_files(root: str, max_depth: int) -> list:
    """Shallow walk under ``root`` (directory depth ≤ ``max_depth``) collecting
    every ``acm_database.json``. Prunes system/noise directories and caps total
    directories visited so a huge cloud tree cannot hang the scan. os.walk
    swallows permission errors, so this never raises."""
    hits: list = []
    try:
        root = os.path.abspath(root)
    except OSError:
        return hits
    base = root.rstrip("\\/").count(os.sep)
    visited = 0
    for dirpath, dirnames, filenames in os.walk(root):
        visited += 1
        if visited > _SCAN_MAX_DIRS:
            break
        if DEFAULT_DB_FILENAME in filenames:
            hits.append(os.path.join(dirpath, DEFAULT_DB_FILENAME))
        depth = dirpath.rstrip("\\/").count(os.sep) - base
        if depth >= max_depth:
            dirnames[:] = []  # at the depth limit: check files, do not descend
            continue
        dirnames[:] = [d for d in dirnames
                       if not d.startswith(".")
                       and d.lower() not in _SCAN_SKIP_DIRS]
    return hits


def discover_db_candidates(max_depth: int = 3) -> list:
    """Existing CAM databases found under the known cloud roots (Phase 4).

    Returns a list of ``{"path", "folder", "counts"}`` dicts — one per distinct
    ``acm_database.json`` discovered — each carrying the assignment / roster /
    class counts (``_db_file_counts``) so the teacher can recognise their real
    data. De-duplicated by resolved path; unreadable hits are dropped (a
    placeholder / not-yet-synced cloud file has nothing to recognise)."""
    seen: set = set()
    out: list = []
    for root in _cloud_search_roots():
        for db in _scan_for_db_files(root, max_depth):
            key = os.path.normcase(os.path.abspath(db))
            if key in seen:
                continue
            seen.add(key)
            if db_file_state(db) != "ok":
                continue
            out.append({"path": db,
                        "folder": os.path.dirname(db),
                        "counts": _db_file_counts(db)})
    return out


# A freshly-created, genuinely-empty database file is tiny (schema envelope +
# empty gradebook, a few hundred bytes). A real year of grades is hundreds of
# KB. So a file that carries appreciably more than an empty envelope yet
# deserializes to zero students AND zero assignments is damaged, not a legit
# first-run file — the boot guard quarantines it rather than letting the demo
# session overwrite it. Generous on purpose (see Phase 1 of
# docs/CROSS_DEVICE_AND_DB_SAFETY_PLAN.md).
EMPTY_DB_MAX_BYTES = 4096


def diagnose_db_load(path: str) -> Optional[dict]:
    """Decide whether the boot hydrate may safely run against ``path``.

    Returns ``None`` when it is safe to proceed with the normal load-or-start-
    empty behaviour, or a ``{"reason": ..., "path": ...}`` dict when the app
    must **not** silently fall back to demo state on top of a real (but
    currently unreadable / unavailable) database. Wipe mechanism 2 in
    docs/CROSS_DEVICE_AND_DB_SAFETY_PLAN.md: an unreadable file, an empty-but-
    heavy file, or a configured path whose storage is gone must never be
    replaced by the demo gradebook via the next autosave.

    Reasons: ``"unreadable"`` (file present, cannot be parsed), ``"empty-load"``
    (parses but yields no students/assignments despite carrying real bytes),
    ``"storage-missing"`` (the file is absent *and* its parent folder / volume
    is not currently accessible — an unplugged drive, an unmounted cloud folder,
    a disconnected share).
    """
    state = db_file_state(path)
    if state == "unreadable":
        return {"reason": "unreadable", "path": path}
    if state == "ok":
        loaded = load_database(path)
        if loaded and not (loaded["gradebook"].students
                           or loaded["gradebook"].assignments):
            try:
                size = os.path.getsize(path)
            except OSError:
                size = 0
            if size > EMPTY_DB_MAX_BYTES:
                return {"reason": "empty-load", "path": path}
        return None
    # state == "absent": an absent file in an existing folder is a legitimate
    # first run (start empty, create on first save). An absent file whose parent
    # folder / volume is itself missing is NOT a first run — never seed a fresh
    # DB onto a reassigned drive letter or a not-yet-synced cloud folder.
    parent = os.path.dirname(os.path.abspath(path))
    if os.path.isdir(parent):
        return None
    return {"reason": "storage-missing", "path": path}


def _safe_dirname(name: str) -> str:
    """Filesystem-safe folder name for a class label (mirrors the grading
    workspace's sanitizer so both apps resolve the same class subfolder)."""
    import re
    return re.sub(r'[\\/*?:"<>|]', "_", (name or "").strip()).strip() or "Unsorted"


def class_data_dir(class_name: str, create: bool = False) -> str:
    """``[db folder]/[class name]/`` — one consolidated folder per class."""
    d = os.path.join(db_folder(), _safe_dirname(class_name))
    if create:
        try:
            os.makedirs(d, exist_ok=True)
        except OSError:
            pass
    return d

# Dense layout CSS: shrink Streamlit's default padding/gaps/margins so the
# three-window cockpit packs into a 1920x1080 frame, while max-width:100% lets
# it scale cleanly up to 4K. Injected once per run from main().
DENSE_CSS = """
<style>
  /* Outer page frame */
  .block-container {
      padding-top: 3.25rem; padding-bottom: 0.8rem;
      padding-left: 1.4rem; padding-right: 1.4rem;
      max-width: 100%;
  }
  /* Tighten the vertical/horizontal flex gaps between every element */
  [data-testid="stVerticalBlock"] { gap: 0.3rem; }
  [data-testid="stHorizontalBlock"] { gap: 0.4rem; }
  [data-testid="column"] { padding: 0 0.15rem; }
  /* Compact widgets — two-tier control sizing (docs/UI_STYLE_GUIDE.md).
     SHORT tier (the default for every command button): st.button,
     st.download_button (deliverable downloads) and st.form_submit_button
     (Save settings / Save changes / Create / Apply) all share one compact
     height. Descendant selector (space, not >) on purpose: a button given
     help=... is wrapped in a tooltip container, which a direct-child
     selector misses — that's how the deliverable Build buttons, 🔄 Sync,
     👁 Watch and "Generate for whole class" ended up a different size. */
  .stButton button, .stDownloadButton button, .stFormSubmitButton button {
      padding: 0.12rem 0.55rem; margin: 0; min-height: 0;
  }
  /* TALL tier (~2.5rem): dropdowns, text/number/date inputs, popover
     triggers (Remarks, ⚠ missing-work) and file-uploader buttons keep
     Streamlit's taller field height — anything that opens a picker, takes
     typed input, or is read for copy/paste stays big. Selects, inputs,
     popovers and uploader buttons already default to this height; only the
     grade chips (theme_css) need an explicit match. */
  div[data-testid="stCheckbox"] { margin: 0; }
  .stTextInput input, .stNumberInput input, .stDateInput input { padding: 0.2rem 0.45rem; }
  div[data-baseweb="select"] > div { min-height: 2rem; }
  /* Safety net so text areas never render shorter than ~3 lines. */
  div[data-testid="stTextArea"] textarea { min-height: 4rem; }
  /* File-upload dropzones EVERYWHERE (Upload & stage files modal, Window 2
     roster intake): one shared grey container. Pin Streamlit's default size
     explicitly so every dropzone matches, and center the row's content
     vertically (the default is align-items:flex-start, which leaves the
     ⬆ Upload button hugging the top) while keeping it flush left. */
  [data-testid="stFileUploaderDropzone"] {
      min-height: 4.25rem; padding: 0.75rem;
      align-items: center; justify-content: flex-start;
  }
  /* Window 2 roster intake: the format hint lives in the caption below the
     box, so the dropzone's own drag-and-drop instructions are noise — hide
     the whole node (a zero-height flex item would skew the centering). */
  .st-key-roster_intake [data-testid="stFileUploaderDropzoneInstructions"]
      { display: none; }
  .st-key-roster_intake [data-testid="stFileUploader"] { margin: 0; }
  /* Single roster file only: once one is chosen, its name should fill the grey
     box (chip stretched full-width) and the now-redundant "+" add-another
     button goes — the file is replaced via its ✕ then a fresh upload. The
     chip container is [data-testid="stFileChips"]; its own ✕ button is a
     BaseButton-minimal nested in the chip row, so hiding only the
     borderlessIcon "+" leaves the ✕ untouched. Dropping the "+" also lets
     justify-content:center truly center the chip against Apply. */
  .st-key-roster_intake [data-testid="stFileChips"] { width: 100%; }
  .st-key-roster_intake [data-testid="stFileChips"] > div { width: 100%; }
  .st-key-roster_intake [data-testid="stFileChips"]
      button[data-testid="stBaseButton-borderlessIcon"] { display: none; }
  /* Headings: trim the generous default margins */
  h1 { font-size: 1.65rem; margin: 0 0 0.2rem 0; }
  h2, h3 { margin: 0.25rem 0 0.2rem 0; }
  .stSubheader, [data-testid="stHeading"] { margin-bottom: 0.2rem; }
  /* Expanders / popovers sit flush */
  div[data-testid="stExpander"] details { margin: 0.1rem 0; }
  div[data-testid="stExpander"] summary { padding: 0.25rem 0.5rem; }
  /* Metrics shrink to fit four-across inside the modal */
  [data-testid="stMetricValue"] { font-size: 1.25rem; }
  [data-testid="stMetricLabel"] { font-size: 0.75rem; }
  /* Captions a touch smaller for dense lists */
  .stCaption, small { font-size: 0.72rem; }
  /* Horizontal rules with minimal vertical footprint */
  hr { margin: 0.4rem 0; }
  /* Deliverables save/finalize status banner: cap its width so the green
     (or red) notification never stretches across the whole tray. */
  .st-key-save_status_bar { max-width: 520px; }
  .st-key-save_status_bar [data-testid="stAlert"] { max-width: 520px; }
  /* Window 3 email chip: shrink the click-to-copy code block into a compact
     pill pushed to the right of the name heading, hugging the baseline. */
  .st-key-w3_email_chip { margin-left: auto; width: fit-content; max-width: 100%; }
  .st-key-w3_email_chip [data-testid="stCode"] { margin: 0; }
  .st-key-w3_email_chip pre { margin: 0; padding: 2px 8px; }
  .st-key-w3_email_chip code { font-size: 0.72rem; }
</style>
"""

# Theme surfaces: .streamlit/config.toml defines the calm grey/red palettes for
# light and dark mode; this layer adds what config keys cannot express — a
# tinted fill on secondary buttons (so every clickable control reads as a
# button on the grey canvas) and a slightly elevated surface for modal dialogs
# and popovers. Values are per-theme so both modes keep the same character.
THEME_SURFACES = {
    "light": {
        "btn_bg": "#DFDCD4", "btn_border": "#BDB8AE", "btn_hover_bg": "#D5D1C7",
        "surface": "#F0EEE9", "grid": "#CBC7BE",
    },
    "dark": {
        "btn_bg": "#3A3835", "btn_border": "#57544F", "btn_hover_bg": "#454340",
        "surface": "#2E2D2B", "grid": "#45433F",
    },
}


def theme_mode() -> str:
    """Active theme type, normalised to a THEME_SURFACES key."""
    return "dark" if st.context.theme.type == "dark" else "light"


def style_chart(fig) -> None:
    """Give a Plotly figure theme-aware grid/axis lines.

    Plotly's gridlines assume a white canvas, so on the grey theme they
    disappear into the background; repaint them in a grey that contrasts
    with the active theme's surfaces."""
    grid = THEME_SURFACES[theme_mode()]["grid"]
    axis = dict(gridcolor=grid, linecolor=grid, zerolinecolor=grid)
    fig.update_xaxes(**axis)
    fig.update_yaxes(**axis)


def theme_css() -> str:
    """Build the theme-aware CSS block for the active (light/dark) theme."""
    mode = theme_mode()
    s = THEME_SURFACES[mode]
    accent = st.get_option(f"theme.{mode}.primaryColor")
    return f"""
<style>
  /* Secondary buttons (incl. form submits and popover triggers): visible
     tinted surface instead of a flat page-colored rectangle. */
  button[kind="secondary"], button[kind="secondaryFormSubmit"],
  [data-testid="stPopover"] > button {{
      background: {s['btn_bg']};
      border-color: {s['btn_border']};
  }}
  button[kind="secondary"]:hover, button[kind="secondaryFormSubmit"]:hover,
  [data-testid="stPopover"] > button:hover {{
      background: {s['btn_hover_bg']};
      border-color: {accent};
  }}
  /* Modal dialogs (Settings, Upload, LLM parameters, Archived, Add/Edit class,
     Assignment analytics) and popover bodies: sit on a slightly elevated
     surface so they separate from the page behind them. */
  div[data-testid="stDialog"] div[role="dialog"] {{ background: {s['surface']}; }}
  [data-testid="stPopoverBody"] {{ background: {s['surface']}; }}
  /* Read-only report-card grade chips (MYP Grade / School Grade): a rounded,
     bordered box on the TALL tier (2.5rem, same as the Effort selectbox in
     the same row) — deliberately big so the value is easy to select and
     copy/paste — but with NO fill so it reads as display-only. */
  .cam-grade-box {{
      display: inline-flex; align-items: center; justify-content: center;
      min-height: 2.5rem; min-width: 3.2rem; padding: 0.2rem 0.75rem;
      border: 1px solid {s['btn_border']}; border-radius: 0.5rem;
      background: transparent; font-weight: 600; font-size: 1.05rem;
  }}
</style>
"""


# --------------------------------------------------------------------------
# Session state
# --------------------------------------------------------------------------

def init_state() -> None:
    """Create every session_state key once, defensively."""
    defaults = {
        "gradebook": Gradebook(),
        "unit_plan": None,
        # "active" is a per-TERM alias: ensure_term_context() points it at
        # active_by_term[current term], so the On checkboxes select which
        # assignments count in the selected term's assessment.
        "active": {},            # assignment name -> bool (current term's map)
        "active_by_term": {},    # term -> {assignment name -> bool}
        "date_override": {},     # assignment name -> date
        "late_flags": {},        # "sid||assignment||crit" -> bool
        "late_flags_cleanup_v1": False,  # one-shot: old-dialog redundant-override purge done
        "excused_flags": {},     # "sid||assignment" -> bool (drop from math/AI)
        "roster": [],            # ordered list of {"key", "name"}
        "focus_sid": None,
        "sel_assignment": None,  # assignment row selected for analytics
        "edit_cell": None,       # (sid, assignment, crit) under edit
        "final_override": {},    # sid -> int band (professional judgment)
        "teacher_remarks": {},   # sid -> str
        "llm_prompt": "",
        # "llm_response" is a per-TERM alias: ensure_term_context() points it
        # at comments_by_term[current term], so every term's overall comments
        # are kept for the whole year and edited under their own term.
        "llm_response": {},      # sid -> overall comment (current term's map)
        "comments_by_term": {},  # term -> {sid -> overall comment}
        "effort_by_term": {},    # term -> {sid -> 0-5 Effort/English use score}
        "calc_method_by_term": {},  # term -> {sid -> pinned calc method}; unpinned
                                    # students follow the count-based auto default

        "llm_status": ("", ""),  # (kind, message) for last API attempt
        "ingested_sigs": set(),  # de-dupe committed files
        "ingested_files": {},    # abspath -> {hash,mtime,...} cloud-sync registry
        # Roster-aware identity routing (Sync/anonymous plan Phase 3):
        #   work_aliases[class]   = {csv_key -> roster_key} — the DURABLE manual
        #     (+ auto-recorded prefix) map that survives Sync's purge-replace, so
        #     a teacher's "this work is theirs" call is never re-derived.
        #   unmatched_works[class][assignment] = [pool-row dicts] — CSV rows that
        #     matched no roster student; REBUILT every time the assignment's CSV
        #     is (re)ingested (purge-replace, like scores). Never mints phantoms.
        "work_aliases": {},      # class -> {csv_key -> roster_key} (durable)
        "unmatched_works": {},   # class -> {assignment -> [pool-row dicts]}
        "session_sync_done": False,  # one-shot guard: session-start global sync
        # Post-🖌-launch marker (Phase 1): {class, assignment, folder_ref,
        # launched_at, last_probe} while a grading session is open, else None.
        # Drives the scoped post-session probe and the global scan's skip guard.
        # Session-only — never persisted (a restart clears it deliberately).
        "active_launch": None,
        "staging": {},           # sig -> staged-file record (uncommitted)
        "archived": set(),       # soft-deleted assignment names
        "db_loaded": False,      # one-shot guard for boot-time DB load
        "db_load_blocked": None,  # Phase-1 quarantine: {reason,path} when the
                                  # configured DB is unreadable/unavailable and
                                  # persist() must refuse to overwrite it.
        # Comment cloud-mirror (docs/COMMENT_CLOUD_MIRROR_PLAN.md, Phase 2):
        #   mirror_ready               one-shot gate — mirroring only runs after
        #                              the boot heal/restore (invariant 1).
        #   mirror_fingerprints        class -> fingerprint of its last mirrored
        #                              slice; skips churn-free rewrites (inv. 3).
        #   mirror_deletions_this_session  a comment/remark was cleared in-app,
        #                              so the shrink tripwire (inv. 2) must allow
        #                              a legitimate mass reduction to reach disk.
        "mirror_ready": False,
        "mirror_fingerprints": {},
        "mirror_deletions_this_session": False,
        "db_switch_pending": None,  # Phase-2 adopt-vs-overwrite: {path,counts,
                                    # old_custom} while the settings dialog asks
                                    # whether to load an existing DB at a newly
                                    # configured path or overwrite it.
        "save_status": ("", ""), # (kind, message) for last save attempt
        "prefs": DEFAULT_PREFS.copy(),  # device-local UI prefs (overwritten below)
        "classes": [],           # [{"name","grade","myp_year"}] taught this year
        "active_class": "",      # name of the class currently in focus
        "active_term": TERMS[0], # term newly-ingested assignments are tagged to
        "rosters": {},           # class name -> roster list (per-class students)
        "archived_students": {}, # class name -> [removed roster entries] (soft)
        "unit_plans": {},        # class name -> parsed UnitPlan (persisted)
        "dlg_settings": False,   # modal visibility flags (flag-driven dialogs)
        "dlg_upload": False,
        "dlg_llm": False,
        "dlg_archived": False,
        "dlg_addclass": False,
        "llm_cfg": {
            "mode": "Clipboard prompt",
            "provider": "Claude",
            "pronouns": "they/them",
            "name_mode": "first",   # full | first | none — name usage in comment
            "inc_strengths": True,
            "inc_growth": True,
            "inc_next": True,
            "inc_criteria": True,
            "word_limit": 100,
            "n_strengths": 1,
            "n_growth": 1,
            "tone_formal": False,      # nudge toward formal, academic language
            "tone_encouraging": False, # nudge toward warm, encouraging language
            "focus_scope": "current",  # "current" | "include_past" (multi-term)
            "inc_trend": True,         # inject the math-engine trajectory string
            "inc_late": True,          # inject the [SUBMISSION TIMELINESS] block
            "inc_missing": True,       # inject the [MISSING WORK] block
            "no_numbers": False,       # comment must not state numeric grades/scores
            "skip_existing": True,     # batch: skip students who already have a comment
            "model": "",
        },
        # School-specific report-card figures layered on top of the MYP
        # criterion grades. All OFF by default so a fresh/public install shows
        # only the criterion A-D grading; a school that uses them enables them
        # in ⚙ Settings. Round-trips in the shared DB (like llm_cfg) so the
        # choice follows the teacher across devices. The CAM master (Excel)
        # export always carries all three regardless of these toggles.
        "report_cfg": {
            "show_myp_grade": False,     # MYP Grade (1-7) banded lookup
            "show_effort": False,        # Effort / English-use score (editable)
            "show_school_grade": False,  # School Grade (1-10) banded lookup
            "effort_min": 0,             # inclusive Effort/English-use range
            "effort_max": 5,
        },
    }
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value

    # Device-local UI prefs load every boot (separate from the cloud database).
    st.session_state["prefs"] = load_prefs()

    # Year-long persistence: on first boot of a session, hydrate from disk so
    # student marks survive across terms / refreshes for the whole school year.
    # Reads from the resolved (possibly cloud) database path.
    # Phase-4 first-boot gate: on a machine that has not chosen a data home yet,
    # defer the hydrate entirely — main() renders the setup panel and returns, so
    # nothing (not even the sample DB) is loaded or persisted until the teacher
    # picks. db_loaded stays False; once a choice commits a path, the hydrate
    # runs on the next rerun with the gate cleared.
    if not st.session_state["db_loaded"] and not _needs_first_boot_setup():
        path = db_path()
        # Load-guard (Phase 1): never let the demo session run — and then
        # autosave — on top of a real database we merely failed to read. If the
        # configured path is unreadable, empty-but-heavy, or on missing storage,
        # quarantine instead of proceeding; persist() then refuses to write.
        blocked = diagnose_db_load(path)
        if blocked:
            st.session_state["db_load_blocked"] = blocked
        else:
            loaded = load_database(path)
            if loaded and (loaded["gradebook"].students or loaded["gradebook"].assignments):
                st.session_state["gradebook"] = loaded["gradebook"]
                restore_session(loaded.get("session", {}))
            # Cloud-mirror invariant 1 (heal before mirror): restore the durable
            # per-class twins into any blank session slot BEFORE the first mirror
            # write, so a DB that lost its teacher input (the 2026-07-10 incident)
            # is refilled from the class files instead of overwriting them with
            # its own emptiness. Session text always wins; only blanks are filled.
            _heal_from_class_mirrors()
            # Seed the no-churn fingerprints from the healed state: a class whose
            # twin already matches is marked so the first persist() won't rewrite
            # it, while a class whose twin is missing/staler is left unseeded so
            # that same persist() backfills it (the first-ever twin for the
            # restored comments). See _seed_mirror_fingerprints.
            _seed_mirror_fingerprints()
            st.session_state["mirror_ready"] = True
        st.session_state["db_loaded"] = True


def gb() -> Gradebook:
    return st.session_state["gradebook"]


# --------------------------------------------------------------------------
# Persistence (auto-save + manual save, year-long local retention)
# --------------------------------------------------------------------------

def build_session_payload() -> dict:
    """Collect the JSON-safe teacher-side state that travels with the gradebook."""
    ss = st.session_state
    return {
        "active": ss["active"],                    # legacy alias (current term)
        "active_by_term": ss["active_by_term"],
        "late_flags": ss["late_flags"],
        "late_flags_cleanup_v1": ss["late_flags_cleanup_v1"],
        "excused_flags": ss["excused_flags"],
        "final_override": ss["final_override"],
        "teacher_remarks": ss["teacher_remarks"],
        # ``llm_response`` (the current term's comment map) is NOT persisted: it
        # is a live in-memory alias of ``comments_by_term[current term]`` (see
        # ensure_term_context), so writing it duplicated ~11% of the DB for
        # nothing. The loader already prefers ``comments_by_term`` and keeps
        # ``llm_response`` only as a read-only fallback for pre-multi-term DBs.
        "comments_by_term": ss["comments_by_term"],
        "effort_by_term": ss["effort_by_term"],
        "calc_method_by_term": ss["calc_method_by_term"],
        "classes": ss["classes"],
        "active_class": ss["active_class"],
        "active_term": ss["active_term"],
        "rosters": ss["rosters"],
        "archived_students": ss["archived_students"],
        "archived": sorted(ss["archived"]),
        "date_override": {k: v.isoformat() for k, v in ss["date_override"].items()},
        # Cloud-sync dedup registry: which watch-folder files we've ingested.
        "ingested_files": ss["ingested_files"],
        # Roster-aware identity routing (Phase 3): the durable alias map and the
        # per-assignment unmatched-works pools (see init_state).
        "work_aliases": ss["work_aliases"],
        "unmatched_works": ss["unmatched_works"],
        # Parsed unit plans per class (previously session-only, now durable).
        "unit_plans": {cls: unit_plan_to_dict(p)
                       for cls, p in ss["unit_plans"].items() if p},
        # Tuned LLM parameters (never the API key — that stays memory-only).
        "llm_cfg": {k: v for k, v in ss["llm_cfg"].items() if k != "api_key"},
        # School-specific report-card figure toggles + Effort range.
        "report_cfg": dict(ss["report_cfg"]),
    }


def restore_session(session: dict) -> None:
    """Re-hydrate teacher-side state from a loaded payload (inverse of above)."""
    if not session:
        return
    ss = st.session_state
    for k in ("late_flags", "excused_flags", "final_override", "teacher_remarks"):
        if isinstance(session.get(k), dict):
            ss[k] = session[k]
    # One-shot cleanup marker (Phase C): restore so the redundant-override purge
    # never re-runs on a database that has already been cleaned.
    if isinstance(session.get("late_flags_cleanup_v1"), bool):
        ss["late_flags_cleanup_v1"] = session["late_flags_cleanup_v1"]
    if isinstance(session.get("classes"), list):
        ss["classes"] = session["classes"]
    if isinstance(session.get("active_class"), str):
        ss["active_class"] = session["active_class"]
    if session.get("active_term") in TERMS:
        ss["active_term"] = session["active_term"]
    if isinstance(session.get("rosters"), dict):
        ss["rosters"] = session["rosters"]
    if isinstance(session.get("archived_students"), dict):
        ss["archived_students"] = session["archived_students"]
    # Legacy single-roster DBs (pre-multi-class) -> stash for migration.
    if isinstance(session.get("roster"), list):
        ss["roster"] = session["roster"]
    if isinstance(session.get("archived"), list):
        ss["archived"] = set(session["archived"])
    if isinstance(session.get("ingested_files"), dict):
        ss["ingested_files"] = session["ingested_files"]
    # Roster-aware identity routing (Phase 3). Both are class-keyed dicts; parse
    # defensively so a malformed entry is skipped, not fatal. work_aliases values
    # are {csv_key -> roster_key}; unmatched_works values are
    # {assignment -> [pool-row dicts]}.
    wa = session.get("work_aliases")
    if isinstance(wa, dict):
        ss["work_aliases"] = {
            str(cls): {str(ck): str(rk) for ck, rk in m.items()}
            for cls, m in wa.items() if isinstance(m, dict)}
    uw = session.get("unmatched_works")
    if isinstance(uw, dict):
        ss["unmatched_works"] = {
            str(cls): {str(asg): [r for r in rows if isinstance(r, dict)]
                       for asg, rows in m.items() if isinstance(rows, list)}
            for cls, m in uw.items() if isinstance(m, dict)}
    # Parsed unit plans per class. Parse each entry defensively so one malformed
    # plan is skipped rather than aborting the whole restore; the class-context
    # pass (ensure_class_context) later repoints the ``unit_plan`` alias.
    up = session.get("unit_plans")
    if isinstance(up, dict):
        restored_plans = {}
        for cls, d in up.items():
            if not isinstance(d, dict):
                continue
            try:
                restored_plans[str(cls)] = unit_plan_from_dict(d)
            except (KeyError, ValueError, TypeError):
                continue
        ss["unit_plans"] = restored_plans
    # Tuned LLM parameters: merge saved values OVER the init_state defaults so a
    # key added in a future release keeps its default. The API key is never
    # persisted (it lives only in st.session_state["llm_api_key"]).
    saved_cfg = session.get("llm_cfg")
    if isinstance(saved_cfg, dict):
        merged_cfg = dict(ss["llm_cfg"])
        for k, v in saved_cfg.items():
            if k in merged_cfg and k != "api_key":
                merged_cfg[k] = v
        ss["llm_cfg"] = merged_cfg
    # Report-card figure toggles + Effort range: same merge-over-defaults, so a
    # key added in a future release keeps its default on an older database.
    saved_rc = session.get("report_cfg")
    if isinstance(saved_rc, dict):
        merged_rc = dict(ss["report_cfg"])
        for k, v in saved_rc.items():
            if k in merged_rc:
                merged_rc[k] = v
        ss["report_cfg"] = merged_rc
    # Legacy databases carry a top-level "calc_method" (the retired global method
    # dropdown) and "w_new" (the retired recency slider); both are ignored — the
    # calculation method is now per-student/per-term (calc_method_by_term, parsed
    # below) and unpinned students follow the count-based auto default, so a teacher
    # who had hand-picked a global method simply re-pins it per student once. No
    # migration runs against the file.
    restored_dates = {}
    for name, iso in (session.get("date_override") or {}).items():
        try:
            restored_dates[name] = date.fromisoformat(iso)
        except (TypeError, ValueError):
            pass
    ss["date_override"] = restored_dates

    # ---- Per-term maps (with one-shot migration from legacy flat maps) ----
    # Runs after active_term is restored so blank/legacy tags resolve sensibly.
    abt = session.get("active_by_term")
    if isinstance(abt, dict):
        ss["active_by_term"] = {str(t): {str(n): bool(v) for n, v in m.items()}
                                for t, m in abt.items() if isinstance(m, dict)}
    elif isinstance(session.get("active"), dict):
        # Legacy flat On-map: apply each assignment's flag to its tagged term.
        legacy = session["active"]
        per: dict = {}
        for a in ss["gradebook"].assignments:
            if a.name in legacy:
                tag = getattr(a, "term", "") or ss["active_term"]
                per.setdefault(tag, {})[a.name] = bool(legacy[a.name])
        ss["active_by_term"] = per
    cbt = session.get("comments_by_term")
    if isinstance(cbt, dict):
        ss["comments_by_term"] = {str(t): {str(s): str(c) for s, c in m.items()}
                                  for t, m in cbt.items() if isinstance(m, dict)}
    elif isinstance(session.get("llm_response"), dict):
        # Legacy single comment set: it belongs to the term that was active.
        ss["comments_by_term"] = {ss["active_term"]: dict(session["llm_response"])}
    ebt = session.get("effort_by_term")
    if isinstance(ebt, dict):
        ss["effort_by_term"] = {
            str(t): {str(s): int(v) for s, v in m.items()
                     if isinstance(v, (int, float))}
            for t, m in ebt.items() if isinstance(m, dict)}
    # Per-term calculation-method pins: keep only still-known methods, drop the
    # rest (a renamed/removed method silently reverts that student to auto).
    cmt = session.get("calc_method_by_term")
    if isinstance(cmt, dict):
        ss["calc_method_by_term"] = {
            str(t): {str(s): m for s, m in mp.items() if m in CALCULATION_METHODS}
            for t, mp in cmt.items() if isinstance(mp, dict)}


def persist(show: bool = False, allow_shrink: bool = False) -> None:
    """Write the gradebook + session state to the JSON database.

    Called after every mutation (auto-save) and from the manual Save button
    (``show=True`` surfaces a confirmation). Failures are captured rather than
    raised so a transient disk error never takes the dashboard down.

    ``allow_shrink`` bypasses the Phase-3 shrink tripwire for a deliberate,
    typed-confirmed mass reduction — the Danger-zone wipe and the explicit,
    backed-up Phase-2 Replace. Every other caller (autosave included) leaves it
    False so a catastrophic mass-loss write is refused.
    """
    # Phase-1 quarantine: the configured database could not be read at boot, so
    # the in-memory state is demo/quarantine state. Writing it would clobber the
    # real (merely unreadable) file — exactly wipe mechanism 2. Refuse; the
    # full-width banner already explains the situation and the fix. A shrink
    # tripwire that fired earlier this session sets the same flag (below).
    if st.session_state.get("db_load_blocked"):
        if show:
            st.session_state["save_status"] = (
                "error", "Save blocked: the database is in read-only "
                "quarantine (see the banner at the top). Fix the file or path "
                "and restart CAM.")
        return
    try:
        path = db_path()
        payload = build_session_payload()
        # Phase-3 shrink tripwire: refuse a catastrophic mass-loss overwrite —
        # a demo/quarantine session must never flatten a rich database on disk.
        # Park the outgoing payload for inspection and raise the same read-only
        # quarantine banner Phase 1 uses. Deliberate wipes set allow_shrink.
        if not allow_shrink and _shrink_would_lose(path, gb(), payload):
            parked = _park_blocked_payload(path, gb(), payload)
            st.session_state["db_load_blocked"] = {
                "reason": "shrink-blocked", "path": path, "parked": parked}
            note = ("Save blocked: this would have erased most of the database "
                    "on disk (see the banner at the top).")
            if parked:
                note += f" Your unsaved session was parked at {parked}."
            st.session_state["save_status"] = ("error", note)
            return
        # Rotating daily backup: snapshot the existing on-disk DB the first time
        # we persist on a new calendar day, before save_database overwrites it.
        _rotate_daily_backup(path)
        save_database(path, gb(), payload)
        if show:
            st.session_state["save_status"] = ("ok", f"Saved to {path}.")
    except Exception as exc:  # disk full / permissions / etc.
        st.session_state["save_status"] = ("error", f"Save failed: {exc}")
        return
    # Cloud-mirror the per-class teacher input AFTER the DB write succeeds, so a
    # mirror hiccup can never abort the primary save. Runs outside the try above:
    # it has its own never-raises contract, and a shrink-tripwire refusal here
    # deliberately overwrites the "Saved" status with its own warning.
    _mirror_classes_to_cloud()


def _db_file_counts(path: str) -> dict:
    """Assignment / roster-student / class counts for an on-disk DB file.

    Read for the Phase-2 settings switch panel so the teacher can recognise the
    database that already lives at a newly-configured path before choosing to
    adopt or overwrite it. Best-effort and never raises — an unreadable file
    reports zeros (the panel only shows for a file that already parsed ``ok``,
    so this is a belt-and-braces guard)."""
    loaded = load_database(path)
    if not loaded:
        return {"assignments": 0, "students": 0, "classes": 0}
    gbk = loaded["gradebook"]
    sess = loaded.get("session", {}) or {}
    rosters = sess.get("rosters", {})
    if isinstance(rosters, dict) and rosters:
        students = sum(len(v) for v in rosters.values() if isinstance(v, list))
    else:
        students = len(gbk.students)
    classes = sess.get("classes", [])
    n_classes = len(classes) if isinstance(classes, list) else 0
    return {"assignments": len(gbk.assignments), "students": students,
            "classes": n_classes}


def _backup_replaced_db(path: str) -> str:
    """Copy an existing DB aside before the current session overwrites it.

    Phase 2 (docs/CROSS_DEVICE_AND_DB_SAFETY_PLAN.md): the explicit "push my
    current session over that file" step must never be silent — the file that
    is about to be replaced is snapshotted to
    ``acm_database.json.bak-replaced-<YYYYMMDD-HHMMSS>`` beside it first (the
    ``.bak-*`` pattern is already git-ignored). Returns the backup path, or
    ``""`` when there was nothing at ``path`` to copy."""
    if not os.path.exists(path):
        return ""
    ts = datetime.now().strftime("%Y%m%d-%H%M%S")
    backup = f"{path}.bak-replaced-{ts}"
    shutil.copy2(path, backup)
    return backup


# ---- Phase 3: persist() shrink tripwire + rotating daily backups ---------
# The last line of defence in docs/CROSS_DEVICE_AND_DB_SAFETY_PLAN.md. persist()
# fires after every mutation and mirrors the whole in-memory session straight to
# disk, so a demo / quarantine session pointed at a rich database would flatten
# it. Before each write we compare a cheap structural "mass" of the outgoing
# payload against the file already on disk and refuse the write when a
# substantial database would collapse to a small fraction of its size. Generous
# on purpose: deleting one class of several still clears the bar; wiping every
# class to the demo gradebook does not. Danger-zone wipes and the explicit,
# backed-up Phase-2 Replace pass ``allow_shrink=True`` (after their own typed
# confirmation) to bypass it.
SHRINK_MIN_ASSIGNMENTS = 10   # below this the on-disk DB is too young to guard
SHRINK_KEEP_RATIO = 0.33      # outgoing mass must stay >= this fraction of disk
AUTO_BACKUP_KEEP = 7          # rotating .bak-auto-<date> snapshots retained


def _scored_student_count(gradebook: Gradebook) -> int:
    """Students carrying at least one criterion score (any bucket non-empty)."""
    return sum(1 for s in gradebook.students.values()
               if any(s.scores.get(c) for c in s.scores))


def _roster_entry_count(session: dict) -> int:
    """Total roster entries across every class in a session payload."""
    rosters = (session or {}).get("rosters", {})
    if not isinstance(rosters, dict):
        return 0
    return sum(len(v) for v in rosters.values() if isinstance(v, list))


def _outgoing_mass(gradebook: Gradebook, session: dict) -> int:
    """Structural mass of the in-memory state about to be written.

    ``assignments + roster entries + scored students`` — the same three
    dimensions :func:`_ondisk_mass` reads back from the file already on disk, so
    the two are directly comparable for the shrink tripwire."""
    return (len(gradebook.assignments) + _roster_entry_count(session)
            + _scored_student_count(gradebook))


def _ondisk_mass(path: str) -> Optional[tuple]:
    """``(n_assignments, total_mass)`` of the database already at ``path``.

    Counted straight from the raw JSON without rebuilding engine objects, so the
    shrink tripwire stays light enough to run on every persist. Returns ``None``
    when nothing readable is there — an absent or unreadable file has no mass to
    protect (the Phase-1 boot guard owns the unreadable-file case)."""
    try:
        with open(path, "r", encoding="utf-8") as fh:
            payload = json.load(fh)
    except (OSError, ValueError):
        return None
    if not isinstance(payload, dict):
        return None
    gbk = payload.get("gradebook", {}) or {}
    students = gbk.get("students", []) or []
    assignments = gbk.get("assignments", []) or []
    n_assign = len(assignments) if isinstance(assignments, list) else 0
    n_scored = sum(1 for s in students if isinstance(s, dict) and s.get("scores"))
    sess = payload.get("session", {}) or {}
    n_roster = _roster_entry_count(sess if isinstance(sess, dict) else {})
    return (n_assign, n_assign + n_roster + n_scored)


def _shrink_would_lose(path: str, gradebook: Gradebook, session: dict) -> bool:
    """True when writing ``gradebook``/``session`` over ``path`` would erase most
    of a substantial database already there (the shrink tripwire condition).

    Only guards a database with real substance (``>= SHRINK_MIN_ASSIGNMENTS`` on
    disk); a young file is expected to shrink as classes come and go. Trips when
    the outgoing structural mass falls below ``SHRINK_KEEP_RATIO`` of the
    on-disk mass."""
    disk = _ondisk_mass(path)
    if disk is None:
        return False
    disk_assign, disk_mass = disk
    if disk_assign < SHRINK_MIN_ASSIGNMENTS or disk_mass <= 0:
        return False
    return _outgoing_mass(gradebook, session) < SHRINK_KEEP_RATIO * disk_mass


def _park_blocked_payload(path: str, gradebook: Gradebook, session: dict) -> str:
    """Write the refused outgoing payload aside as ``<path>.blocked-<ts>``.

    So a blocked write is never simply lost: the session that would have been
    saved is parked for inspection (via the engine's atomic writer) beside the
    protected database. Returns the parked path, or ``""`` on failure."""
    ts = datetime.now().strftime("%Y%m%d-%H%M%S")
    blocked = f"{path}.blocked-{ts}"
    try:
        save_database(blocked, gradebook, session)
    except Exception:
        return ""
    return blocked


def _prune_auto_backups(path: str, keep: int = AUTO_BACKUP_KEEP) -> None:
    """Keep only the newest ``keep`` ``<path>.bak-auto-*`` snapshots.

    The ``-<YYYYMMDD>`` suffix sorts chronologically, so lexicographic order is
    date order. Only auto snapshots match the glob — manual ``.bak-replaced-*`` /
    ``.bak-<purpose>-*`` files are never touched (see the safety rules in
    CLAUDE.md: never prune existing ``.bak-*`` files)."""
    if keep <= 0:
        return
    try:
        snaps = sorted(glob.glob(glob.escape(path) + ".bak-auto-*"))
    except OSError:
        return
    for old in snaps[:-keep]:
        try:
            os.remove(old)
        except OSError:
            pass


def _rotate_daily_backup(path: str) -> None:
    """Snapshot the existing on-disk DB once per calendar day before it is
    overwritten, so any future incident is at most a one-day loss even without
    OneDrive version history.

    Keyed by today's date (``<path>.bak-auto-<YYYYMMDD>``) and skipped when that
    snapshot already exists, so it fires on the *first* persist of the day and
    survives restarts. Best-effort — a copy failure never blocks the save."""
    if not os.path.exists(path):
        return
    backup = f"{path}.bak-auto-{datetime.now().strftime('%Y%m%d')}"
    if os.path.exists(backup):
        return
    try:
        shutil.copy2(path, backup)
    except OSError:
        return
    _prune_auto_backups(path)


def calc_method_map() -> dict:
    """The CURRENT term's per-student calculation-method pins ({sid -> method}).

    A new term has no entry, so it opens empty and everyone falls to auto — the
    "reset to default at a term boundary" falls out of the data shape for free."""
    return st.session_state["calc_method_by_term"].setdefault(current_term(), {})


def qualifying_assignment_count() -> int:
    """How many of the current term's assignments feed the auto-default sizing.

    Qualifying = On in the current term (``current_term_assignment_names``),
    criteria-bearing (assignment-table ``criteria != "—"`` — formative events and
    still-unbanded exams carry "—" and don't count), and **not** a ``(Reflection)``
    task (an adjunct of its parent assignment). Banded exams count. Counted overall,
    never per criterion."""
    names = current_term_assignment_names()
    n = 0
    for r in assignment_table():
        if r["name"] not in names or r["criteria"] == "—":
            continue
        if r["name"].strip().lower().endswith("(reflection)"):
            continue
        n += 1
    return n


def auto_calc_method() -> str:
    """Auto default from the size of this term's assessment set: <= 15 qualifying
    assignments -> 60/40 Recency, > 15 -> Weighted Median. Recomputed live, so a
    mid-term flip when the 16th assignment lands is expected and acceptable."""
    return METHOD_60_40 if qualifying_assignment_count() <= 15 else METHOD_WEIGHTED_MEDIAN


def calculation_method(sid: str) -> str:
    """The student's grading method for the current term: their explicit per-term
    pin when one is stored (and still a known method), else ``auto_calc_method()``.

    ``sid`` is REQUIRED — there is deliberately no zero-argument form, so a missed
    call site fails loudly instead of silently reverting to a single global value."""
    pinned = calc_method_map().get(sid)
    return pinned if pinned in CALCULATION_METHODS else auto_calc_method()


# --------------------------------------------------------------------------
# Finalized term summaries (per-class files in the cloud folder)
# --------------------------------------------------------------------------

def class_term_summaries(cls: str) -> dict:
    """Finalized term summaries for a class: read from the class's own data
    folder first, falling back to the database root for legacy files."""
    by_term = load_term_summaries(class_data_dir(cls), cls)
    if by_term:
        return by_term
    return load_term_summaries(db_folder(), cls)


def past_term_context_for(student) -> list:
    """Ordered ``[(term, comment)]`` of prior-term context for a student.

    For each term preceding the current one, prefer the finalized cloud
    summary; fall back to the overall comment saved in-app for that term, so
    multi-term generation works even before a term is formally finalized."""
    cls = st.session_state["active_class"]
    by_term = class_term_summaries(cls)
    cur_idx = term_index(current_term())
    out = []
    for term in TERMS:
        if term_index(term) >= cur_idx:
            continue
        text = ((by_term.get(term, {}) or {}).get(student.student_id, "")
                or "").strip()
        if not text:
            text = ((st.session_state["comments_by_term"].get(term, {}) or {})
                    .get(student.student_id, "") or "").strip()
        if text:
            out.append((term, text))
    return out


# --------------------------------------------------------------------------
# Teacher-input cloud mirror (per-class files; docs/COMMENT_CLOUD_MIRROR_PLAN.md)
# --------------------------------------------------------------------------
#
# Every teacher-typed input the app cannot rebuild from the export CSVs — overall
# comments, teacher remarks, effort scores, final-grade overrides and CAM-typed
# score comments — lives only in the DB session. The 2026-07-10 wipe destroyed a
# term of it. persist() mirrors each class's slice of that state into the class's
# own cloud folder (``acm_term_summaries_<class>.json``, the v2 payload written
# by engine.save_class_mirror) so a wipe leaves a durable cloud twin for the
# heal-on-load pass (Phase 3) to restore. The DB session stays the runtime source
# of truth; the class files are its durable twin.
#
# Three safety invariants gate the write (all enforced in _mirror_classes_to_cloud):
#   1. Heal before mirror — never push a freshly-wiped/quarantined session's
#      emptiness over good class files (``mirror_ready`` / ``db_load_blocked``).
#   2. Shrink tripwire — refuse to collapse a term's comment count below half of
#      what the file holds, unless a comment/remark was deliberately cleared
#      in-app this session (``mirror_deletions_this_session``).
#   3. No churn — skip the write when the slice is byte-identical to the last one
#      mirrored (per-class fingerprint), so OneDrive isn't spammed on every edit.

# Shrink tripwire: a term must keep >= half its on-disk comments (invariant 2),
# but only once the file holds enough of them to be worth guarding.
MIRROR_SHRINK_MIN_COMMENTS = 4
MIRROR_SHRINK_KEEP_RATIO = 0.5


def _class_sids(cls: str) -> set:
    """Every student id owned by a class: its live roster plus its archived
    (departed-but-grades-kept) students, so a departed student's typed comment
    still earns a durable cloud twin instead of silently dropping out."""
    ss = st.session_state
    sids = {e["key"] for e in ss["rosters"].get(cls, []) if e.get("key")}
    for e in ss.get("archived_students", {}).get(cls, []):
        if e.get("key"):
            sids.add(e["key"])
    return sids


def build_class_mirror(cls: str) -> dict:
    """Assemble the v2 teacher-input mirror slice for one class from live session
    state + the gradebook. Every global teacher-input map is filtered to this
    class's own students so the per-class file never carries another class's
    content; blank leaves and empty containers are dropped by the engine on
    write (engine.save_class_mirror -> _clean_mirror)."""
    ss = st.session_state
    sids = _class_sids(cls)

    terms: dict = {}
    for term, by_sid in ss["comments_by_term"].items():
        slice_ = {sid: txt for sid, txt in by_sid.items()
                  if sid in sids and str(txt).strip()}
        if slice_:
            terms[term] = slice_

    remarks = {sid: txt for sid, txt in ss["teacher_remarks"].items()
               if sid in sids and str(txt).strip()}

    effort: dict = {}
    for term, by_sid in ss["effort_by_term"].items():
        slice_ = {sid: v for sid, v in by_sid.items() if sid in sids}
        if slice_:
            effort[term] = slice_

    final_override = {sid: dict(crits)
                      for sid, crits in ss["final_override"].items()
                      if sid in sids and crits}

    class_assignments = {a.name for a in gb().assignments
                         if getattr(a, "class_name", "") == cls}
    score_comments: dict = {}
    for sid in sids:
        student = gb().students.get(sid)
        if student is None:
            continue
        for crit, bucket in student.scores.items():
            for sc in bucket:
                text = (sc.comment or "").strip()
                if not text or sc.assignment not in class_assignments:
                    continue
                score_comments.setdefault(sc.assignment, {}).setdefault(
                    sid, {})[crit] = sc.comment

    return {
        "terms": terms,
        "remarks": remarks,
        "effort": effort,
        "final_override": final_override,
        "score_comments": score_comments,
    }


def _mirror_fingerprint(mirror: dict) -> str:
    """Stable content hash of a class slice, for the no-churn guard (invariant
    3). Rebuilt identically from unchanged session state, so an autosave that
    didn't touch this class's teacher input reproduces the same fingerprint."""
    blob = json.dumps(mirror, sort_keys=True, ensure_ascii=False)
    return hashlib.sha256(blob.encode("utf-8")).hexdigest()


def _mirror_shrink_would_lose(folder: str, cls: str, new_mirror: dict) -> bool:
    """True when writing ``new_mirror`` would drop some term's overall-comment
    count below half of what the class file already holds (invariant 2). Only
    the ``terms`` section is guarded — it is the human-typed content a wipe
    cannot rebuild. Best-effort: an unreadable/absent file never blocks."""
    try:
        existing = load_class_mirror(folder, cls)
    except Exception:
        return False
    old_terms = existing.get("terms", {}) or {}
    new_terms = new_mirror.get("terms", {}) or {}
    for term, old_map in old_terms.items():
        old_n = len(old_map)
        if old_n < MIRROR_SHRINK_MIN_COMMENTS:
            continue
        if len(new_terms.get(term, {})) < old_n * MIRROR_SHRINK_KEEP_RATIO:
            return True
    return False


def _mark_teacher_input_deleted() -> None:
    """Record that a comment/remark was deliberately cleared in-app this session,
    so the mirror shrink tripwire (invariant 2) lets the reduction reach disk
    instead of mistaking a real deletion for catastrophic mass loss."""
    st.session_state["mirror_deletions_this_session"] = True


def _mirror_classes_to_cloud() -> None:
    """Mirror every class's teacher-input slice to its cloud file. Called from
    persist() after a successful DB write.

    Never raises (invariant 4): one class's cloud hiccup degrades to a
    ``save_status`` note and never blocks the other classes or the autosave that
    already succeeded. The three write gates are invariants 1-3 above."""
    ss = st.session_state
    # Invariant 1: never mirror before the boot heal/restore, nor from a
    # quarantined boot (that state is demo/empty and would clobber good files).
    if ss.get("db_load_blocked") or not ss.get("mirror_ready"):
        return
    fingerprints = ss.setdefault("mirror_fingerprints", {})
    allow_shrink = ss.get("mirror_deletions_this_session", False)
    blocked: list = []
    for cls in class_names():
        try:
            mirror = build_class_mirror(cls)
            fp = _mirror_fingerprint(mirror)
            if fingerprints.get(cls) == fp:
                continue  # invariant 3: unchanged slice -> no churn
            folder = class_data_dir(cls, create=True)
            if not allow_shrink and _mirror_shrink_would_lose(folder, cls, mirror):
                blocked.append(cls)
                continue  # invariant 2: refuse a catastrophic comment-mass loss
            save_class_mirror(folder, cls, mirror)
            fingerprints[cls] = fp
        except Exception:
            continue  # invariant 4: degrade gracefully, never take autosave down
    if blocked:
        ss["save_status"] = (
            "error",
            "Cloud comment backup skipped for "
            + ", ".join(f"'{c}'" for c in blocked)
            + ": the class file holds many more comments than this session, so "
            "the backup was held back to avoid overwriting them. If you did "
            "clear comments on purpose, edit any comment to confirm and it will "
            "save.")


# --------------------------------------------------------------------------
# Heal on load (Phase 3) — the read direction of the cloud mirror
# --------------------------------------------------------------------------
#
# The write direction (above) mirrors every class's teacher input to its cloud
# twin on autosave. The read direction fills the session back from those twins
# when the DB lost them (the 2026-07-10 incident wiped comments_by_term and the
# effort/remark/override maps while grades self-healed from the export CSVs).
#
# Two heal passes, both fill blank slots only — a value still present in the
# session always wins over its twin, so a live edit is never clobbered:
#   * _heal_from_class_mirrors  — comments / remarks / effort / final_override,
#     at boot right after restore_session (before the first mirror write).
#   * _heal_score_comments_from_mirrors — sc.comment, after a Sync purge-replace
#     (which rebuilds scores from the CSVs and so drops any CAM-typed comment).
#
# All heal/seed helpers degrade silently (invariant 4): a cloud read hiccup must
# never take boot or Sync down.


def _heal_from_class_mirrors() -> None:
    """Fill blank teacher-input slots from each class's cloud twin (Phase 3).

    For every class in the session, load its mirror and backfill only the slots
    the session left blank in ``comments_by_term`` / ``teacher_remarks`` /
    ``effort_by_term`` / ``final_override``. Session content always wins where
    both are present, so this restores a wiped DB without ever overwriting a
    value the teacher still holds. Never raises."""
    ss = st.session_state
    cbt = ss["comments_by_term"]
    remarks = ss["teacher_remarks"]
    ebt = ss["effort_by_term"]
    fo = ss["final_override"]
    for cls in class_names():
        try:
            mirror = load_class_mirror(class_data_dir(cls), cls)
        except Exception:
            continue  # invariant 4: a cloud hiccup never blocks boot
        # Overall comments: {term: {sid: text}} — fill blank/absent text only.
        for term, by_sid in mirror.get("terms", {}).items():
            dst = cbt.setdefault(term, {})
            for sid, text in by_sid.items():
                if not str(dst.get(sid, "")).strip():
                    dst[sid] = text
        # Teacher remarks: flat {sid: text}.
        for sid, text in mirror.get("remarks", {}).items():
            if not str(remarks.get(sid, "")).strip():
                remarks[sid] = text
        # Effort: {term: {sid: int}} — presence, not truthiness (0 is a real
        # score), so a set effort of 0 is never re-healed away.
        for term, by_sid in mirror.get("effort", {}).items():
            dst = ebt.setdefault(term, {})
            for sid, val in by_sid.items():
                if sid not in dst:
                    dst[sid] = val
        # Final override: {sid: {crit: band}} — fill only the missing criteria.
        for sid, by_crit in mirror.get("final_override", {}).items():
            dst = fo.setdefault(sid, {})
            for crit, band in by_crit.items():
                if crit not in dst:
                    dst[crit] = band


def _heal_score_comments_from_mirrors() -> None:
    """Refill blank ``sc.comment`` slots from each class's cloud twin (Phase 3).

    Runs after a Sync purge-replace, which rebuilds a class's scores from its
    export CSVs and so drops any comment that was typed in CAM (not carried by
    the CSV). Fill blanks only — a comment the CSV still carries, or one typed
    this session, always wins. Never raises."""
    for cls in class_names():
        try:
            mirror = load_class_mirror(class_data_dir(cls), cls)
        except Exception:
            continue
        sc_map = mirror.get("score_comments", {})
        if not sc_map:
            continue
        for assignment, by_sid in sc_map.items():
            for sid, by_crit in by_sid.items():
                student = gb().students.get(sid)
                if student is None:
                    continue
                for crit, text in by_crit.items():
                    for sc in student.scores.get(crit, []):
                        if (sc.assignment == assignment
                                and not (sc.comment or "").strip()):
                            sc.comment = text


def _seed_mirror_fingerprints() -> None:
    """Seed the no-churn fingerprints (invariant 3) from the post-heal state.

    A class whose cloud twin already matches the healed session is fingerprinted
    now, so the first ``persist()`` recognises it as unchanged and does not
    rewrite an identical file. A class whose twin is missing or staler than the
    session is deliberately left UNSEEDED, so that first ``persist()`` backfills
    it — this is what gives the restored comments their (often first-ever) cloud
    twin. Never raises."""
    ss = st.session_state
    fingerprints = ss.setdefault("mirror_fingerprints", {})
    for cls in class_names():
        try:
            session_fp = _mirror_fingerprint(build_class_mirror(cls))
            disk_fp = _mirror_fingerprint(load_class_mirror(class_data_dir(cls), cls))
            if session_fp == disk_fp:
                fingerprints[cls] = session_fp
        except Exception:
            continue


# --------------------------------------------------------------------------
# Term backup & restore (docs/TERM_BACKUP_RESTORE_PLAN.md)
# --------------------------------------------------------------------------
#
# The THIRD line of defence behind the mirror-heal (above) and the rotating
# ``.bak-auto`` files: a deliberate, teacher-initiated snapshot of one whole
# term — every assignment, score, exam result, comment, effort, override and
# flag CAM holds for that term — written as one self-describing JSON to a folder
# the teacher chooses (may be outside OneDrive), plus a loader that can put that
# term back after a database disaster.
#
# Scope is by TERM TAG, never by dates. Every assignment carries its term, so a
# single stable predicate (``_term_of_assignment``) partitions the gradebook;
# the loader touches only rows tagged with the backup's term and leaves every
# other term byte-identical. Restore REPLACES the term's slice wholesale (it is
# a disaster tool, not an editing tool) behind a dry-run diff, a typed
# confirmation and an automatic pre-restore DB backup.


def _term_of_assignment(a) -> str:
    """The term an assignment belongs to for backup/restore scoping.

    Its own tag, or ``TERMS[0]`` for a blank/legacy tag. Deliberately STABLE —
    independent of the active term (unlike :func:`assignment_term`) — so the
    same assignment always maps to the same backup no matter what term the UI
    happens to be showing when a backup or a restore runs."""
    return getattr(a, "term", "") or TERMS[0]


def _term_backup_slug(term: str) -> str:
    """Filesystem-safe lowercase slug for a term (``"Term 1"`` -> ``"term-1"``)."""
    slug = "".join(ch.lower() if ch.isalnum() else "-" for ch in (term or ""))
    return "-".join(part for part in slug.split("-") if part) or "term"


def _sid_to_class_index() -> dict:
    """Reverse index ``{student_id: class_name}`` from every class's owned sids
    (roster + archived). A sid on no class falls to ``""`` so a term map entry is
    grouped under a catch-all rather than silently dropped from a backup."""
    idx: dict = {}
    for cls in class_names():
        for sid in _class_sids(cls):
            idx.setdefault(sid, cls)
    return idx


def build_term_backup(term: str) -> dict:
    """Assemble the self-describing backup dict for one whole term.

    Everything CAM knows for ``term``: its assignments (all classes), the scores
    and exam results those assignments produced, the term's overall comments,
    effort scores, active-map and per-student calc-method pins, plus the
    late/excused flags for its assignments. Teacher remarks and final overrides
    are NOT term-scoped in the session (a single flat map spans the year), so the
    whole non-blank map is captured for reference; restore fills only blank slots
    from them (see :func:`restore_term_backup`).

    Comments / effort / remarks / overrides / scores / exam results are grouped
    per class so a human (and the dry-run diff) can read the file class by class;
    active / calc-method / late / excused are the term's flat maps verbatim."""
    ss = st.session_state
    gbk = gb()

    term_asgs = [a for a in gbk.assignments if _term_of_assignment(a) == term]
    names_by_class: dict = {}
    for a in term_asgs:
        names_by_class.setdefault(getattr(a, "class_name", ""), set()).add(a.name)
    all_term_names = {a.name for a in term_asgs}

    sid_class = _sid_to_class_index()

    def _group(flat: dict) -> dict:
        """``{sid: value}`` -> ``{class: {sid: value}}`` via the reverse index."""
        out: dict = {}
        for sid, val in flat.items():
            out.setdefault(sid_class.get(sid, ""), {})[sid] = val
        return out

    comments = _group({sid: txt for sid, txt
                       in (ss["comments_by_term"].get(term, {}) or {}).items()
                       if str(txt).strip()})
    effort = _group(dict(ss["effort_by_term"].get(term, {}) or {}))
    remarks = _group({sid: txt for sid, txt in ss["teacher_remarks"].items()
                      if str(txt).strip()})
    final_override = _group({sid: dict(crits) for sid, crits
                             in ss["final_override"].items() if crits})

    # Scores + exam results, per class per student, filtered to the term's
    # assignment names. A class's students = its roster/archived set PLUS any
    # score-only student who has evidence for one of these names (so a student
    # not yet on the roster is never dropped from the snapshot).
    scores: dict = {}
    exam_results: dict = {}
    for cls, names in names_by_class.items():
        cls_sids = set(_class_sids(cls))
        for student in gbk.students.values():
            has_score = any(sc.assignment in names
                            for bucket in student.scores.values() for sc in bucket)
            has_exam = any(k in names for k in getattr(student, "exam_results", {}))
            if has_score or has_exam:
                cls_sids.add(student.student_id)
        for sid in cls_sids:
            student = gbk.students.get(sid)
            if student is None:
                continue
            sc_list = [score_to_dict(sc) for bucket in student.scores.values()
                       for sc in bucket if sc.assignment in names]
            if sc_list:
                scores.setdefault(cls, {})[sid] = sc_list
            ex_list = [exam_result_to_dict(r) for k, r
                       in getattr(student, "exam_results", {}).items() if k in names]
            if ex_list:
                exam_results.setdefault(cls, {})[sid] = ex_list

    active = {n: bool(v) for n, v in (ss["active_by_term"].get(term, {}) or {}).items()
              if n in all_term_names}
    calc_method = {sid: m for sid, m
                   in (ss["calc_method_by_term"].get(term, {}) or {}).items()
                   if m in CALCULATION_METHODS}

    late_flags = {k: bool(v) for k, v in ss["late_flags"].items()
                  if len(k.split("||")) == 3 and k.split("||")[1] in all_term_names}
    excused = {k: bool(v) for k, v in ss["excused_flags"].items()
               if len(k.split("||")) == 2 and k.split("||")[1] in all_term_names}

    assignments = [assignment_to_dict(a) for a in term_asgs]
    classes = sorted({getattr(a, "class_name", "") for a in term_asgs}
                     | set(comments) | set(effort) | set(scores))

    counts = {
        "students": len({sid for m in scores.values() for sid in m}
                        | {sid for m in comments.values() for sid in m}),
        "assignments": len(assignments),
        "comments": sum(len(m) for m in comments.values()),
        "scores": sum(len(lst) for m in scores.values() for lst in m.values()),
        "exam_results": sum(len(lst) for m in exam_results.values()
                            for lst in m.values()),
    }

    return {
        "version": TERM_BACKUP_VERSION,
        "kind": TERM_BACKUP_KIND,
        "term": term,
        "created_at": datetime.now().isoformat(),
        "db_path": db_path(),      # provenance only
        "classes": classes,
        "counts": counts,
        "payload": {
            "assignments": assignments,
            "scores": scores,
            "exam_results": exam_results,
            "comments": comments,
            "effort": effort,
            "active": active,
            "calc_method": calc_method,
            "late_flags": late_flags,
            "excused": excused,
            "remarks": remarks,
            "final_override": final_override,
        },
    }


def write_term_backup(term: str, folder: str, backup: Optional[dict] = None) -> str:
    """Serialize :func:`build_term_backup` for ``term`` into ``folder`` atomically.

    Filename ``cam_term_backup_<slug>_<YYYYMMDD-HHMMSS>.json``; written to a temp
    file in the same folder then renamed over the target, so a crash mid-write
    never leaves a truncated backup. Returns the path written. Pass a prebuilt
    ``backup`` dict to avoid rebuilding it when the caller also needs its counts.
    Raises on I/O failure — the caller (Settings) surfaces it on ``save_status``."""
    if backup is None:
        backup = build_term_backup(term)
    ts = datetime.now().strftime("%Y%m%d-%H%M%S")
    fname = f"cam_term_backup_{_term_backup_slug(term)}_{ts}.json"
    directory = os.path.abspath(folder or ".")
    os.makedirs(directory, exist_ok=True)
    path = os.path.join(directory, fname)
    fd, tmp = tempfile.mkstemp(suffix=".tmp", dir=directory)
    try:
        with os.fdopen(fd, "w", encoding="utf-8") as fh:
            json.dump(backup, fh, ensure_ascii=False, indent=2)
        os.replace(tmp, path)
    finally:
        if os.path.exists(tmp):
            os.remove(tmp)
    return path


def validate_term_backup(payload) -> tuple:
    """Cheap structural validation of a parsed backup file.

    Returns ``(term, "")`` when the file is a well-formed term backup, or
    ``(None, message)`` describing the first problem otherwise — checked BEFORE
    any dry-run diff or confirmation is offered, so a malformed / wrong-kind /
    wrong-version file is refused up front and never mutates anything."""
    if not isinstance(payload, dict):
        return None, "Not a CAM backup file (the JSON is not an object)."
    if payload.get("kind") != TERM_BACKUP_KIND:
        return None, ("This file is not a CAM term backup "
                      f"(kind = {payload.get('kind')!r}).")
    if payload.get("version") != TERM_BACKUP_VERSION:
        return None, (f"Unsupported backup version {payload.get('version')!r} "
                      f"(this CAM reads version {TERM_BACKUP_VERSION}).")
    term = payload.get("term")
    if term not in TERMS:
        return None, f"Backup names an unknown term {term!r}."
    if not isinstance(payload.get("payload"), dict):
        return None, "Backup file is missing its payload section."
    return term, ""


def _merge_class_maps(by_class: dict) -> dict:
    """Flatten a ``{class: {sid: value}}`` backup section back to ``{sid: value}``
    for a wholesale term-slice replace. Class grouping is a storage/display
    convenience only — the session maps it restores into are flat per term."""
    out: dict = {}
    for by_sid in (by_class or {}).values():
        if isinstance(by_sid, dict):
            out.update(by_sid)
    return out


def diff_term_backup(payload: dict) -> dict:
    """Compute the dry-run diff of a validated backup against the live session.

    Nothing is written. Returns a structure the Settings dialog renders so the
    teacher sees exactly what a restore would do BEFORE typing the confirmation:
    per class, which comments would be newly filled vs would overwrite a
    different current comment, which assignments the backup adds vs which live
    ones it would remove (they postdate the backup), and the backup-vs-live score
    counts; plus how many remarks / overrides the fill-blanks pass would add."""
    ss = st.session_state
    gbk = gb()
    term = payload["term"]
    p = payload.get("payload", {})

    backup_comments = p.get("comments", {}) or {}
    backup_scores = p.get("scores", {}) or {}
    backup_asgs = [assignment_from_dict(d) for d in p.get("assignments", [])]
    backup_names_by_class: dict = {}
    for a in backup_asgs:
        backup_names_by_class.setdefault(getattr(a, "class_name", ""), set()).add(a.name)
    live_names_by_class: dict = {}
    for a in gbk.assignments:
        if _term_of_assignment(a) == term:
            live_names_by_class.setdefault(getattr(a, "class_name", ""), set()).add(a.name)

    live_comments = ss["comments_by_term"].get(term, {}) or {}

    def _live_score_count(cls: str) -> int:
        names = live_names_by_class.get(cls, set())
        if not names:
            return 0
        sids = _class_sids(cls) | set(backup_scores.get(cls, {}))
        total = 0
        for sid in sids:
            student = gbk.students.get(sid)
            if student is None:
                continue
            total += sum(1 for bucket in student.scores.values()
                         for sc in bucket if sc.assignment in names)
        return total

    per_class: dict = {}
    all_classes = (set(backup_comments) | set(backup_scores)
                   | set(backup_names_by_class) | set(live_names_by_class))
    for cls in sorted(all_classes):
        bc = backup_comments.get(cls, {})
        new_comments = [sid for sid, txt in bc.items()
                        if not str(live_comments.get(sid, "")).strip()]
        changed_comments = [
            sid for sid, txt in bc.items()
            if str(live_comments.get(sid, "")).strip()
            and str(live_comments.get(sid, "")).strip() != str(txt).strip()]
        b_names = backup_names_by_class.get(cls, set())
        l_names = live_names_by_class.get(cls, set())
        per_class[cls] = {
            "new_comments": sorted(new_comments),
            "changed_comments": sorted(changed_comments),
            "assignments_added": sorted(b_names - l_names),
            "assignments_removed": sorted(l_names - b_names),
            "backup_scores": sum(len(v) for v in backup_scores.get(cls, {}).values()),
            "live_scores": _live_score_count(cls),
        }

    # Fill-blanks-only maps: count how many entries the restore would actually add.
    remarks = ss["teacher_remarks"]
    remarks_fill = sum(
        1 for by_sid in (p.get("remarks", {}) or {}).values()
        for sid in by_sid if not str(remarks.get(sid, "")).strip())
    fo = ss["final_override"]
    override_fill = 0
    for by_sid in (p.get("final_override", {}) or {}).values():
        for sid, crits in by_sid.items():
            live_crits = fo.get(sid, {})
            override_fill += sum(1 for crit in crits if crit not in live_crits)

    return {
        "term": term,
        "created_at": payload.get("created_at", ""),
        "counts": payload.get("counts", {}),
        "per_class": per_class,
        "remarks_fill": remarks_fill,
        "override_fill": override_fill,
    }


def _pre_restore_backup(path: str) -> str:
    """Copy the live DB aside as ``…bak-pre-term-restore-<stamp>`` before a
    restore mutates anything. Never pruned (only ``.bak-auto-*`` rotates), per
    CLAUDE.md. Returns the backup path, or ``""`` when there was nothing to copy
    or the copy failed (the wholesale slice replace is still guarded by the diff
    + typed confirmation, so a copy hiccup does not abort the restore)."""
    if not os.path.exists(path):
        return ""
    ts = datetime.now().strftime("%Y%m%d-%H%M%S")
    backup = f"{path}.bak-pre-term-restore-{ts}"
    try:
        shutil.copy2(path, backup)
    except OSError:
        return ""
    return backup


def restore_term_backup(payload: dict) -> str:
    """Replace the live term's slice wholesale with a validated backup's.

    Disaster-recovery semantics (docs/TERM_BACKUP_RESTORE_PLAN.md §4):

    * Term-tagged data — assignments, their scores/exam results, the term's
      comments/effort/active/calc-method maps and the late/excused flags for its
      assignments — is DELETED then REPLACED with the backup's, including
      removing term-tagged rows that exist live but not in the backup (they
      postdate the backup; the staleness warning covers this).
    * Non-term-scoped maps (teacher_remarks, final_override) are filled for
      students/criteria with no live entry only — a wholesale replace would
      clobber another term's remarks while restoring this one.
    * Data for every OTHER term is never touched.

    Writes the automatic pre-restore ``.bak`` first, then a single
    ``persist(allow_shrink=True)`` at the end (atomic tmp+replace) — so a crash
    mid-restore leaves the pre-restore backup intact. Seeds the mirror-deletion
    flag and clears the per-class fingerprints so the restored comments re-mirror
    to their cloud twins and the mirror tripwire does not mistake the restore for
    a mass deletion. Returns the pre-restore backup path (may be ``""``)."""
    ss = st.session_state
    gbk = gb()
    term = payload["term"]
    p = payload.get("payload", {})

    backup_path = _pre_restore_backup(db_path())

    # Term-assignment name sets to purge, per class: the LIVE term's names (to
    # drop live-only rows) UNION the backup's names.
    backup_asgs = [assignment_from_dict(d) for d in p.get("assignments", [])]
    backup_names_by_class: dict = {}
    for a in backup_asgs:
        backup_names_by_class.setdefault(getattr(a, "class_name", ""), set()).add(a.name)
    live_names_by_class: dict = {}
    for a in gbk.assignments:
        if _term_of_assignment(a) == term:
            live_names_by_class.setdefault(getattr(a, "class_name", ""), set()).add(a.name)
    scores_in = p.get("scores", {}) or {}
    exams_in = p.get("exam_results", {}) or {}

    # Purge the live term's scores/exam results per class. Roster-scoped like
    # delete_class so a unit name shared with another class is never touched
    # there; the backup's own sids extend the scope so a score-only student
    # present in the snapshot is re-placed cleanly rather than duplicated.
    for cls in set(backup_names_by_class) | set(live_names_by_class):
        purge = backup_names_by_class.get(cls, set()) | live_names_by_class.get(cls, set())
        if not purge:
            continue
        roster_keys = set(_class_sids(cls)) | set(scores_in.get(cls, {})) \
            | set(exams_in.get(cls, {}))
        for student in gbk.students.values():
            if roster_keys and student.student_id not in roster_keys:
                continue
            for crit, bucket in list(student.scores.items()):
                student.scores[crit] = [sc for sc in bucket if sc.assignment not in purge]
            results = getattr(student, "exam_results", {})
            for k in [n for n in results if n in purge]:
                results.pop(k, None)

    # Remove every live assignment tagged to this term, then register the
    # backup's (their tags already resolve to this term by construction).
    gbk.assignments = [a for a in gbk.assignments if _term_of_assignment(a) != term]
    for a in backup_asgs:
        gbk.register_assignment(a)

    # Re-add the backup's scores + exam results.
    for cls, by_sid in scores_in.items():
        for sid, sc_list in by_sid.items():
            student = gbk.get_or_create(sid)
            for d in sc_list:
                try:
                    student.add_score(score_from_dict(d))
                except (KeyError, ValueError, TypeError):
                    continue
    for cls, by_sid in exams_in.items():
        for sid, ex_list in by_sid.items():
            student = gbk.get_or_create(sid)
            for d in ex_list:
                try:
                    r = exam_result_from_dict(d)
                    student.exam_results[r.assignment] = r
                except (KeyError, ValueError, TypeError):
                    continue

    # Wholesale-replace the term slice of the term-partitioned maps.
    ss["comments_by_term"][term] = {str(s): str(t) for s, t
                                    in _merge_class_maps(p.get("comments", {})).items()}
    ss["effort_by_term"][term] = {str(s): int(v) for s, v
                                  in _merge_class_maps(p.get("effort", {})).items()
                                  if isinstance(v, (int, float))}
    ss["active_by_term"][term] = {str(n): bool(v)
                                  for n, v in (p.get("active", {}) or {}).items()}
    ss["calc_method_by_term"][term] = {str(s): m for s, m
                                       in (p.get("calc_method", {}) or {}).items()
                                       if m in CALCULATION_METHODS}

    # late/excused: global maps keyed by assignment name — drop this term's live
    # entries, then add the backup's.
    term_names = ({a.name for a in backup_asgs}
                  | {n for names in live_names_by_class.values() for n in names})
    lf = ss["late_flags"]
    for k in [k for k in lf
              if len(k.split("||")) == 3 and k.split("||")[1] in term_names]:
        lf.pop(k, None)
    for k, v in (p.get("late_flags", {}) or {}).items():
        lf[str(k)] = bool(v)
    ef = ss["excused_flags"]
    for k in [k for k in ef
              if len(k.split("||")) == 2 and k.split("||")[1] in term_names]:
        ef.pop(k, None)
    for k, v in (p.get("excused", {}) or {}).items():
        ef[str(k)] = bool(v)

    # Non-term-scoped maps: fill blank slots only (never clobber another term).
    remarks = ss["teacher_remarks"]
    for sid, text in _merge_class_maps(p.get("remarks", {})).items():
        if not str(remarks.get(sid, "")).strip():
            remarks[sid] = str(text)
    fo = ss["final_override"]
    for sid, crits in _merge_class_maps(p.get("final_override", {})).items():
        dst = fo.setdefault(sid, {})
        for crit, band in (crits or {}).items():
            if crit not in dst:
                dst[crit] = band

    # Re-point the term/class aliases, then seed the mirror flags so the restored
    # comments reach their cloud twins on this persist and the shrink tripwire
    # treats the restore as an intended change, not catastrophic mass loss.
    ensure_class_context()
    ensure_term_context()
    _mark_teacher_input_deleted()
    ss["mirror_fingerprints"] = {}
    # Deliberate, typed-confirmed, already-backed-up write: the shrink tripwire's
    # exempt path (like the Danger-zone wipe and the Phase-2 Replace).
    persist(allow_shrink=True)
    return backup_path


def sync_roster_into_students() -> None:
    """Reconcile the roster into the gradebook.

    Every roster entry must have a matching ``Student`` record, because Focus,
    the cockpit (Window 3), the exports and ``students_for_active_class`` all
    read ``gb().students`` — a roster-only student (added by CSV upload, with no
    grades yet) is otherwise invisible to the whole app. So we CREATE the record
    when it is missing (not just update it), then mirror the roster's display
    name + gender onto it so exports, the cockpit and LLM pronouns match and it
    persists in the shared database."""
    for entry in st.session_state["roster"]:
        key = entry.get("key")
        if not key:
            continue
        student = gb().get_or_create(key, entry.get("name", "") or "")
        if entry.get("name"):
            student.name = entry["name"]
        if entry.get("gender"):
            student.gender = entry["gender"]


def roster_gender(sid: str) -> str:
    for entry in st.session_state["roster"]:
        if entry.get("key") == sid:
            return entry.get("gender", "") or ""
    return ""


def pronouns_for(student) -> str:
    """Pronouns derived from the student's gender (Window 2 selection)."""
    gender = getattr(student, "gender", "") or roster_gender(student.student_id)
    return GENDER_PRONOUNS.get(gender, "they/them")


def first_name_for(student) -> str:
    """The student's first name (from the roster's First Name column), falling
    back to the trailing token of their "Surname First" display name, then ID."""
    for entry in st.session_state["roster"]:
        if entry.get("key") == student.student_id and entry.get("first"):
            return entry["first"]
    name = getattr(student, "name", "") or ""
    if name:
        return name.split()[-1]   # "Surname First" -> first
    return student.student_id


def student_email_for(student) -> str:
    """The student's school email address (from the roster), or "" if unknown.

    The email is not stored on the ``Student`` record itself — only the
    email-derived numeric id is — so it is looked up from the roster entry."""
    for entry in st.session_state["roster"]:
        if entry.get("key") == student.student_id:
            return entry.get("email", "") or ""
    return ""


# --------------------------------------------------------------------------
# Class / level context (one database, many classes across the year)
# --------------------------------------------------------------------------

def class_names() -> list:
    return [c["name"] for c in st.session_state["classes"]]


def active_class_dict() -> dict:
    name = st.session_state["active_class"]
    for c in st.session_state["classes"]:
        if c["name"] == name:
            return c
    return {}


def active_myp_year() -> str:
    return str(active_class_dict().get("myp_year", "") or "")


def active_subject() -> str:
    """The subject taught to the active class (e.g. "Visual Arts")."""
    return str(active_class_dict().get("subject", "") or "").strip()


def subject_label() -> str:
    """Subject for display in reports/exports ("—" when not set)."""
    return active_subject() or "—"


def class_label(c: dict) -> str:
    """Compact label, e.g. "1-4 · Year 7 · MYP Y2 · Visual Arts"."""
    bits = [c.get("name", "")]
    if c.get("grade"):
        bits.append(c["grade"])
    if c.get("myp_year"):
        bits.append(f"MYP Y{c['myp_year']}")
    if c.get("subject"):
        bits.append(c["subject"])
    return "  ·  ".join(b for b in bits if b)


def current_phase() -> str:
    """Rubric phase for the active class: prefer its MYP year, else unit plan."""
    yr = active_myp_year()
    if yr:
        return rubrics.phase_for_year(yr)
    plan = st.session_state.get("unit_plan")
    return rubrics.phase_for_year(plan.myp_year if plan else None)


def set_active_roster(entries: list) -> None:
    """Replace the active class's roster (and the live alias) atomically.

    Reconciles immediately so the gradebook gains a ``Student`` record for each
    roster entry — Focus, Window 3 and the exports (which all read
    ``gb().students``) then work the instant a roster is loaded, even before any
    grade exists. An uploaded Classroom CSV now behaves like a manual student
    add."""
    cls = st.session_state["active_class"]
    st.session_state["rosters"][cls] = entries
    st.session_state["roster"] = entries
    sync_roster_into_students()


def archived_for_active() -> list:
    """Archived assignment names that belong to the active class."""
    active = st.session_state["active_class"]
    names = st.session_state["archived"]
    return sorted({a.name for a in gb().assignments
                   if a.name in names and getattr(a, "class_name", "") == active})


def ensure_class_context() -> None:
    """Guarantee a valid active class, migrating any legacy single-class data,
    then alias the active class's roster + unit plan onto the flat session keys
    the rest of the app reads (``roster`` / ``unit_plan``)."""
    ss = st.session_state
    if not ss["classes"]:
        # First run or a pre-multi-class database: adopt whatever exists into a
        # single default class so nothing is lost and the app stays usable.
        default = {"name": "Class 1", "grade": "", "myp_year": "",
                   "subject": "", "master_dir": ""}
        ss["classes"] = [default]
        ss["active_class"] = "Class 1"
        for a in gb().assignments:
            if not getattr(a, "class_name", ""):
                a.class_name = "Class 1"
        legacy = ss.get("roster") or []
        if legacy and "Class 1" not in ss["rosters"]:
            ss["rosters"]["Class 1"] = legacy
    # Make sure the selected class is real.
    if ss["active_class"] not in class_names():
        ss["active_class"] = class_names()[0]
    # Alias the active class's roster + unit plan onto the flat keys.
    ss["roster"] = ss["rosters"].setdefault(ss["active_class"], [])
    ss["unit_plan"] = ss["unit_plans"].get(ss["active_class"])


def create_class(name: str, grade: str, myp_year: str, subject: str = "",
                 master_dir: str = "") -> bool:
    name = (name or "").strip()
    if not name:
        st.session_state["save_status"] = ("error", "Class name is required.")
        return False
    if name in class_names():
        st.session_state["save_status"] = ("error", f"Class '{name}' already exists.")
        return False
    st.session_state["classes"].append(
        {"name": name, "grade": (grade or "").strip(),
         "myp_year": str(myp_year or ""),
         "subject": (subject or "").strip(),
         "master_dir": (master_dir or "").strip()}
    )
    st.session_state["active_class"] = name
    st.session_state["rosters"].setdefault(name, [])
    # Every class keeps its own consolidated data folder inside the database
    # folder (grading exports, exam scans, caches, term summaries).
    class_data_dir(name, create=True)
    persist()
    return True


def rename_class(old: str, new: str) -> bool:
    """Rename a class/level, moving every reference keyed by its name.

    A class name is the key for its roster, archived students, unit plan, each
    assignment's ``class_name``, the cloud-sync registry rows, and its on-disk
    data folder — so all of them have to move together or something is
    orphaned. Returns True on success (or when the name is unchanged), False on
    a blank/duplicate name or an unknown class.
    """
    ss = st.session_state
    new = (new or "").strip()
    if old not in class_names():
        ss["save_status"] = ("error", f"Class '{old}' not found.")
        return False
    if not new:
        ss["save_status"] = ("error", "Class name is required.")
        return False
    if new == old:
        return True
    if new in class_names():
        ss["save_status"] = ("error", f"Class '{new}' already exists.")
        return False

    # Move the class's on-disk data folder (grading exports, exam scans, term
    # summaries) so the renamed class keeps its history. Best-effort: if the
    # move fails (folder open/locked, or on a cloud mirror), the data folder is
    # simply recreated empty under the new name on next access.
    try:
        old_dir, new_dir = class_data_dir(old), class_data_dir(new)
        if os.path.isdir(old_dir) and not os.path.exists(new_dir):
            os.rename(old_dir, new_dir)
    except OSError:
        pass

    # The class entry itself.
    for c in ss["classes"]:
        if c["name"] == old:
            c["name"] = new
    # Containers keyed by class name.
    for store in ("rosters", "archived_students", "unit_plans"):
        if old in ss[store]:
            ss[store][new] = ss[store].pop(old)
    # Assignments carry their owning class's name.
    for a in gb().assignments:
        if getattr(a, "class_name", "") == old:
            a.class_name = new
    # Cloud-sync registry rows tag each file with its class.
    for rec in ss["ingested_files"].values():
        if rec.get("class") == old:
            rec["class"] = new
    # Follow the active pointer, then re-alias roster/unit-plan onto the flat keys.
    if ss["active_class"] == old:
        ss["active_class"] = new
    ensure_class_context()
    persist()
    ss["save_status"] = ("ok", f"Renamed class '{old}' → '{new}'.")
    return True


def update_class(old: str, name: str, grade: str, myp_year: str,
                 subject: str = "", master_dir: str = "") -> bool:
    """Edit an existing class in place: rename (via rename_class, which moves
    every name-keyed store together) plus the descriptive fields. Returns
    False when the rename is rejected (blank/duplicate name)."""
    if not rename_class(old, name):
        return False
    name = (name or "").strip()
    for c in st.session_state["classes"]:
        if c["name"] == name:
            c.update({"grade": (grade or "").strip(),
                      "myp_year": str(myp_year or ""),
                      "subject": (subject or "").strip(),
                      "master_dir": (master_dir or "").strip()})
    persist()
    st.session_state["save_status"] = ("ok", f"Updated class '{name}'.")
    return True


# --------------------------------------------------------------------------
# Class master directory & Watch (subfolders -> Assignment/Exam rows)
# --------------------------------------------------------------------------

def class_master_dir() -> str:
    """The active class's master directory: a local path or a Drive folder ID."""
    return str(active_class_dict().get("master_dir", "") or "").strip()


def _master_is_local(ref: str) -> bool:
    """Local path when it exists on disk or carries path separators; anything
    else is treated as a Google Drive folder ID."""
    if os.path.isdir(ref):
        return True
    return any(sep in ref for sep in ("\\", "/", ":"))


def _drive_list_subfolders(folder_id: str) -> list:
    """[(name, id)] subfolders of a Drive folder, via the grading workspace's
    saved OAuth token. Raises RuntimeError with a teacher-readable message when
    credentials or the Google packages are unavailable."""
    token = os.path.join(GRADING_WORKSPACE_DIR, "token.json")
    if not os.path.exists(token):
        raise RuntimeError(
            "No Google token found — click 🔗 Connect Google Drive in the "
            "✎ Add / Edit class dialog (top bar) to run the one-time "
            "sign-in, then Watch again.")
    try:
        from google.oauth2.credentials import Credentials
        from google.auth.transport.requests import Request
        from googleapiclient.discovery import build
    except ImportError:
        raise RuntimeError(
            "Google API packages are not installed in this environment — "
            "run: pip install google-api-python-client google-auth-oauthlib")
    creds = Credentials.from_authorized_user_file(
        token, ["https://www.googleapis.com/auth/drive.readonly"])
    if not creds.valid and creds.expired and creds.refresh_token:
        creds.refresh(Request())
    service = build("drive", "v3", credentials=creds, cache_discovery=False)
    subs, page_token = [], None
    query = (f"'{folder_id}' in parents and trashed = false "
             f"and mimeType = 'application/vnd.google-apps.folder'")
    while True:
        resp = service.files().list(
            q=query, spaces="drive",
            fields="nextPageToken, files(id, name)",
            pageSize=200, pageToken=page_token,
            includeItemsFromAllDrives=True, supportsAllDrives=True,
            orderBy="name_natural",
        ).execute()
        subs.extend((f.get("name", "(unnamed)"), f["id"])
                    for f in resp.get("files", []))
        page_token = resp.get("nextPageToken")
        if not page_token:
            break
    return subs


def _watch_class_master(cls_name: str, ref: str) -> tuple:
    """Scan one class's master directory; every subfolder found becomes an
    Assignment/Exam row in Window 1's assignment list.

    Idempotent: each discovered folder is pinned to its assignment via
    ``folder_ref``, so re-watching never duplicates a row — even after the
    teacher renames the assignment in the Manage menu (the physical folder is
    never renamed). Returns ``(n_found, n_created)``; raises on scan failure."""
    if _master_is_local(ref):
        if not os.path.isdir(ref):
            raise RuntimeError(f"Master directory not found: {ref}")
        subs = [(e.name, os.path.abspath(e.path))
                for e in sorted(os.scandir(ref), key=lambda e: e.name)
                if e.is_dir()]
    else:
        subs = _drive_list_subfolders(ref)

    known_refs = {getattr(a, "folder_ref", "") for a in gb().assignments
                  if getattr(a, "class_name", "") == cls_name}
    existing_names = {a.name for a in gb().assignments}
    created = 0
    adopted = 0
    for sub_name, sub_ref in subs:
        if sub_ref in known_refs:
            continue   # already watched (possibly renamed since) — no duplicate
        # Adopt an existing same-name assignment in this class that isn't yet
        # pinned to a folder — typically one created by CSV ingest (which sets
        # no folder_ref) or a manual add. Stamping the folder_ref onto it keeps
        # a graded assignment and its watched source folder as ONE record,
        # instead of spawning a parallel "X (2)" placeholder. That duplication
        # is what let a re-sync purge (keyed by name) orphan the name-keyed
        # scores, so adoption is the durable fix, not just cosmetics.
        adopt = next((a for a in gb().assignments
                      if a.name == sub_name
                      and getattr(a, "class_name", "") == cls_name
                      and not getattr(a, "folder_ref", "")), None)
        if adopt is not None:
            adopt.folder_ref = sub_ref
            known_refs.add(sub_ref)
            term_active_map(current_term()).setdefault(sub_name, True)
            adopted += 1
            continue
        display = sub_name
        n = 2
        while display in existing_names:
            display = f"{sub_name} ({n})"
            n += 1
        gb().register_assignment(Assignment(
            name=display, criteria=[], ingested_at=datetime.now(),
            note="watched from class master directory",
            class_name=cls_name, term=current_term(), folder_ref=sub_ref))
        existing_names.add(display)
        term_active_map(current_term()).setdefault(display, True)
        created += 1
    if created or adopted:
        persist()
    return len(subs), created


def watch_master_directory() -> None:
    """Watch button: scan the ACTIVE class's master directory."""
    ref = class_master_dir()
    if not ref:
        st.session_state["save_status"] = (
            "error", "Set this class's master directory first "
                     "(✎ Add / Edit class in the top bar).")
        return
    try:
        found, created = _watch_class_master(st.session_state["active_class"], ref)
    except Exception as exc:
        st.session_state["save_status"] = ("error", f"Watch failed: {exc}")
        return
    if not found:
        msg = "Watch: no subfolders found in the master directory."
    else:
        msg = (f"Watch: {found} folder(s) found — {created} new "
               f"assignment(s)/exam(s) added to the list.")
    st.session_state["save_status"] = ("ok", msg)


# --------------------------------------------------------------------------
# App bridge: CAM -> Grading Workspace (Flask sub-app on its own port)
# --------------------------------------------------------------------------

def _port_open(port: int) -> bool:
    with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
        s.settimeout(0.3)
        return s.connect_ex(("127.0.0.1", port)) == 0


# The CGW child CAM spawned, so CAM can stop it on exit instead of leaving an
# orphaned python.exe holding port 5001 (CGW has no window and Streamlit's
# Ctrl-C does not clean up children). ``job`` is a Windows Job Object handle
# kept open for CAM's lifetime — see _bind_workspace_to_cam.
_WORKSPACE = {"proc": None, "job": None, "atexit": False}


def _terminate_workspace() -> None:
    """Stop the CGW child if CAM started one and it's still alive (atexit)."""
    proc = _WORKSPACE.get("proc")
    if proc is not None and proc.poll() is None:
        try:
            proc.terminate()
            try:
                proc.wait(timeout=5)
            except subprocess.TimeoutExpired:
                proc.kill()
        except Exception:
            pass


def _bind_workspace_to_cam(proc) -> None:
    """Tie the CGW child's lifetime to CAM's so it can never outlive it.

    Two layers, because neither alone is sufficient:
      * ``atexit`` — graceful shutdown (Ctrl-C, normal interpreter exit). Does
        NOT run on a hard kill (taskkill /F, closing the terminal window).
      * a Windows Job Object with KILL_ON_JOB_CLOSE — the OS kills every process
        in the job the instant CAM's handle to it closes, which happens however
        CAM dies. This is the layer that actually prevents orphans on a hard
        kill. The job handle is stashed in ``_WORKSPACE`` and deliberately kept
        open for CAM's whole run: closing it early would kill CGW immediately.
    """
    _WORKSPACE["proc"] = proc
    if not _WORKSPACE["atexit"]:
        atexit.register(_terminate_workspace)
        _WORKSPACE["atexit"] = True
    if os.name != "nt":
        return
    try:
        import ctypes
        from ctypes import wintypes

        ULONG_PTR = ctypes.c_size_t

        class JOBOBJECT_BASIC_LIMIT_INFORMATION(ctypes.Structure):
            _fields_ = [
                ("PerProcessUserTimeLimit", ctypes.c_int64),
                ("PerJobUserTimeLimit", ctypes.c_int64),
                ("LimitFlags", wintypes.DWORD),
                ("MinimumWorkingSetSize", ctypes.c_size_t),
                ("MaximumWorkingSetSize", ctypes.c_size_t),
                ("ActiveProcessLimit", wintypes.DWORD),
                ("Affinity", ULONG_PTR),
                ("PriorityClass", wintypes.DWORD),
                ("SchedulingClass", wintypes.DWORD),
            ]

        class IO_COUNTERS(ctypes.Structure):
            _fields_ = [
                ("ReadOperationCount", ctypes.c_uint64),
                ("WriteOperationCount", ctypes.c_uint64),
                ("OtherOperationCount", ctypes.c_uint64),
                ("ReadTransferCount", ctypes.c_uint64),
                ("WriteTransferCount", ctypes.c_uint64),
                ("OtherTransferCount", ctypes.c_uint64),
            ]

        class JOBOBJECT_EXTENDED_LIMIT_INFORMATION(ctypes.Structure):
            _fields_ = [
                ("BasicLimitInformation", JOBOBJECT_BASIC_LIMIT_INFORMATION),
                ("IoInfo", IO_COUNTERS),
                ("ProcessMemoryLimit", ctypes.c_size_t),
                ("JobMemoryLimit", ctypes.c_size_t),
                ("PeakProcessMemoryUsed", ctypes.c_size_t),
                ("PeakJobMemoryUsed", ctypes.c_size_t),
            ]

        k32 = ctypes.WinDLL("kernel32", use_last_error=True)
        k32.CreateJobObjectW.restype = wintypes.HANDLE
        k32.CreateJobObjectW.argtypes = [wintypes.LPVOID, wintypes.LPCWSTR]
        k32.SetInformationJobObject.restype = wintypes.BOOL
        k32.SetInformationJobObject.argtypes = [
            wintypes.HANDLE, ctypes.c_int, wintypes.LPVOID, wintypes.DWORD]
        k32.AssignProcessToJobObject.restype = wintypes.BOOL
        k32.AssignProcessToJobObject.argtypes = [wintypes.HANDLE, wintypes.HANDLE]

        job = k32.CreateJobObjectW(None, None)
        if not job:
            return
        info = JOBOBJECT_EXTENDED_LIMIT_INFORMATION()
        JOB_OBJECT_LIMIT_KILL_ON_JOB_CLOSE = 0x00002000
        info.BasicLimitInformation.LimitFlags = JOB_OBJECT_LIMIT_KILL_ON_JOB_CLOSE
        JobObjectExtendedLimitInformation = 9
        if not k32.SetInformationJobObject(
                job, JobObjectExtendedLimitInformation,
                ctypes.byref(info), ctypes.sizeof(info)):
            return
        # proc._handle is the child's process HANDLE on Windows.
        k32.AssignProcessToJobObject(job, int(proc._handle))
        _WORKSPACE["job"] = job     # keep open for CAM's lifetime
    except Exception:
        pass    # best-effort; atexit still covers the graceful shutdown path


def _ensure_workspace_running() -> bool:
    """Spawn the Flask grading workspace if it isn't already listening.

    The sub-app is started as a detached process so Streamlit is never
    blocked. Returns True once the port answers, False on failure."""
    app_path = os.path.join(GRADING_WORKSPACE_DIR, "app.py")
    if not os.path.exists(app_path):
        st.session_state["save_status"] = (
            "error", f"Grading workspace not found at {app_path}.")
        return False
    if not _port_open(GRADING_PORT):
        # Capture the sub-app's stdout/stderr to a log file (not DEVNULL) so a
        # crash-on-startup — e.g. a missing dependency like PyMuPDF — leaves a
        # readable trace we can surface, instead of silently returning True and
        # letting the caller open a browser tab that just says "connection
        # refused".
        log_path = os.path.join(GRADING_WORKSPACE_DIR, "workspace_startup.log")
        try:
            flags = getattr(subprocess, "CREATE_NO_WINDOW", 0)
            log = open(log_path, "w", encoding="utf-8", errors="replace")
            proc = subprocess.Popen(
                [sys.executable, app_path, "--port", str(GRADING_PORT)],
                cwd=GRADING_WORKSPACE_DIR, creationflags=flags,
                stdout=log, stderr=subprocess.STDOUT)
            # Bind it to CAM's lifetime so closing CAM doesn't leave an orphan.
            _bind_workspace_to_cam(proc)
        except Exception as exc:
            st.session_state["save_status"] = (
                "error", f"Could not launch the grading workspace: {exc}")
            return False
        for _ in range(40):            # wait for Flask to boot (max ~10s)
            if _port_open(GRADING_PORT):
                return True
            time.sleep(0.25)
        # Spawned but never started listening — read the log's last non-blank
        # line (usually the exception) so the banner names the actual fix.
        detail = ""
        try:
            lines = [ln.strip() for ln in open(log_path, encoding="utf-8",
                                               errors="replace") if ln.strip()]
            if lines:
                detail = " — " + lines[-1]
        except OSError:
            pass
        st.session_state["save_status"] = (
            "error",
            f"Grading workspace did not start on port {GRADING_PORT}{detail} "
            f"(see cam_grading_workspace/workspace_startup.log)")
        return False
    return True


def _seed_workspace_class() -> str:
    """Align the grading workspace's settings with CAM before a handoff.

    Two things get pushed through the workspace's own /api/config endpoint
    (call only after _ensure_workspace_running()), which updates the running
    app's in-memory SETTINGS *and* persists gcg_settings.json (root + cloud
    mirror) exactly like a manual Settings-dialog save:

      * ``cloud_dir`` -> CAM's database folder (:func:`db_folder`), so the
        workspace routes its grade/exam CSV exports into the per-class
        subfolder that CAM's Sync scans. Without this the workspace falls back
        to its own app root (a browser download), the CSV never reaches CAM's
        database folder, and Sync reports "nothing new" forever. Only set when
        a custom database path is configured — Sync is inert otherwise.
      * the active class's name -> Drive-folder-ID entry, so the CAM-bridge
        autoloader can resolve the class without the teacher re-typing it.
        Other classes in the map are left untouched.

    Returns '' when the workspace is aligned (or there is nothing to seed yet).
    Since PDF/local-mode plan Phase 4 the workspace's ``/api/class`` enumerates a
    local master directory off disk too, so both a Drive folder ID and a
    local-path master are live entries in the class map — folder grading is no
    longer Drive-only and this never refuses a local class."""
    cls = st.session_state["active_class"]
    ref = class_master_dir()
    want_cloud = (db_folder()
                  if (st.session_state.get("prefs", {}).get("db_custom_path")
                      or "").strip() else "")

    api = f"http://127.0.0.1:{GRADING_PORT}/api/config"
    try:
        with urlopen(api, timeout=5) as resp:
            cfg = json.load(resp) or {}
        classes = cfg.get("classes") or {}
        payload = {}
        # Point the workspace's exports at CAM's database folder.
        cur_cloud = (cfg.get("cloud_dir") or "").strip()
        if want_cloud and os.path.normcase(os.path.normpath(cur_cloud or ".")) \
                != os.path.normcase(os.path.normpath(want_cloud)):
            payload["cloud_dir"] = want_cloud
        # Seed the class -> master map. The reference is a Drive folder ID or a
        # local class-master path; since Phase 4 the workspace resolves both
        # (its /api/class enumerates a local master off disk), so a local path
        # is a live entry, not a dead one.
        if ref and classes.get(cls) != ref:
            classes[cls] = ref
            payload["classes"] = classes
        if payload:
            req = Request(api, data=json.dumps(payload).encode("utf-8"),
                          headers={"Content-Type": "application/json"})
            with urlopen(req, timeout=5):
                pass
    except Exception as exc:
        return f"Could not sync the workspace settings for class '{cls}': {exc}"
    return ""


def _workspace_state_key(folder_ref: str) -> str:
    """The durable persistence key the grading workspace derives for a marking
    target, so CAM's handoff file (``cam_grades_<key>.json``) lands under the
    exact name the workspace reads.

    A Drive folder ID is already stable and *is* the key. A local assignment
    folder path is reduced to the **same** ``local-<hash>`` slug the workspace's
    ``LocalProvider.state_key`` produces (sha1 of the normcased absolute path,
    first 16 hex) — the two apps must agree or the round-trip silently breaks.
    Mirrors CGW ``app.py`` ``LocalProvider.state_key``."""
    if _master_is_local(folder_ref):
        norm = os.path.normcase(os.path.abspath(folder_ref))
        return "local-" + hashlib.sha1(norm.encode("utf-8")).hexdigest()[:16]
    return folder_ref


def _publish_workspace_grades(assignment_name: str, folder_ref: str) -> None:
    """Publish CAM's current grades for one folder-backed assignment so the
    grading workspace can start from them (the CAM-is-source-of-truth loop).

    Written at handoff, right after :func:`_seed_workspace_class`, into the
    class's data folder — the same ``[db folder]/[class]/`` the workspace
    reads and exports into — as ``cam_grades_<folderId>.json``: one entry per
    student (CAM's email-derived student id) holding the current 0-8 band per
    criterion plus the CAM-side comment. On its next load of that folder the
    workspace consumes the file, adopts these values, flags any band that
    changed since its last export as MODIFIED, and carries everything forward
    into its next export. Every export is therefore a superset of CAM's
    latest, which is what keeps Sync's whole-assignment purge-replace from
    wiping marks edited in CAM. Raises ``OSError`` when the file can't be
    written (the caller aborts the handoff — a lossy launch is worse than no
    launch). Skipped when no custom database path is set: without ``cloud_dir``
    routing the export never reaches Sync, so there is no round-trip."""
    if not (st.session_state.get("prefs", {}).get("db_custom_path") or "").strip():
        return
    cls = st.session_state["active_class"]
    students = {}
    for student in students_for_active_class():
        grades = {}
        comment = ""
        for c in CRIT_ORDER:
            sc = find_score(student, assignment_name, c)
            if sc is None or not sc.is_valid:
                continue
            grades[c] = int(sc.value)
            comment = comment or (sc.comment or "").strip()
        if grades:
            students[student.student_id] = {"grades": grades,
                                            "comment": comment}
    # Publish reverse-map (Phase 3): a student whose scores arrived through an
    # alias is graded in CAM under the ROSTER id, but CGW still knows that work
    # only by its anonymous csv_key (CGW computes the work's key as
    # student_id_from_email(email) or name — for a local/unmatched work that is
    # the csv_key, never the roster id). So mirror each aliased student's entry
    # under the csv_key too, so CGW's api_load reconcile lands on the right work.
    # The roster-id key is kept as well (harmless — CGW just routes it to
    # cam_extra if no work matches, exactly like a files-less CAM student).
    aliases = st.session_state["work_aliases"].get(cls, {})
    for csv_key, roster_key in aliases.items():
        if roster_key in students and csv_key not in students:
            students[csv_key] = students[roster_key]
    payload = {
        "assignment": assignment_name,
        "class": cls,
        "folder_ref": folder_ref,
        "published": datetime.now().isoformat(timespec="seconds"),
        "students": students,
    }
    import re
    # Key the handoff file by the workspace's DURABLE key, not the raw ref: a
    # Drive ID passes through unchanged, but a local path becomes the same
    # local-<hash> slug the workspace reads (a raw path would sanitise to a
    # different name and the workspace would never find it).
    safe_fid = re.sub(r"[^A-Za-z0-9_-]", "_", _workspace_state_key(folder_ref))
    dest = os.path.join(class_data_dir(cls, create=True),
                        f"cam_grades_{safe_fid}.json")
    tmp = dest + ".tmp"
    with open(tmp, "w", encoding="utf-8") as fh:
        json.dump(payload, fh, ensure_ascii=False, indent=2)
    os.replace(tmp, dest)   # atomic — the workspace never sees a half-write


def launch_grading_workspace(assignment_name: str) -> str:
    """Open the grading workspace targeting one assignment. Returns the URL
    opened ('' on failure).

    The target assignment travels as URL query parameters that the
    workspace's CAM-bridge autoloader resolves (Drive folder ID first, then
    exam::<name>, then display-name match)."""
    cls = st.session_state["active_class"]
    asg = next((a for a in gb().assignments
                if a.name == assignment_name
                and getattr(a, "class_name", "") == cls), None)
    if not _ensure_workspace_running():
        return ""
    ref = getattr(asg, "folder_ref", "") if asg else ""
    is_exam = bool(getattr(asg, "is_exam", False)) if asg else False
    target = (f"exam::{assignment_name}" if is_exam and not ref
              else (ref or assignment_name))
    # Seed the workspace (class map + cloud_dir). Both Drive and local masters
    # are seedable since Phase 4, so seed_err is empty for a valid class; it is
    # only non-empty when the /api/config round-trip itself fails, which is a
    # genuine dead end for folder grading (exams are exempt — they live in the
    # workspace by class name, no Drive and no class map needed).
    seed_err = _seed_workspace_class()
    if seed_err and not target.startswith("exam::"):
        st.session_state["save_status"] = ("error", seed_err)
        return ""
    # Folder grading round-trips through the workspace: hand it CAM's current
    # grades so its next export carries them forward (exams don't round-trip).
    if ref and not target.startswith("exam::"):
        # Close the stale-handoff race (Terrain §T4): if a previous session's
        # export was never synced, CAM's values are OLDER than the CSV on disk;
        # publishing them now would make CGW's reconcile adopt the stale value
        # and lose the teacher's newer marks. Scoped-sync this assignment FIRST
        # so we publish CAM's latest. A duplicate-dated group or a parse failure
        # cancels the launch — a session started on an ambiguous/stale baseline
        # clobbers marks (same philosophy as the publish-failure cancel below).
        ssum = sync_assignment(cls, assignment_name)
        if ssum["duplicates"]:
            _report_sync(ssum)   # renders the duplicate-refusal banner
            return ""
        if ssum["errors"]:
            prev = st.session_state.get("save_status", ("", ""))
            detail = prev[1] if prev and prev[0] == "error" else ""
            st.session_state["save_status"] = (
                "error",
                (f"Launch cancelled — could not sync the latest export for "
                 f"'{assignment_name}' before publishing (a stale baseline "
                 f"would overwrite newer marks). {detail}").strip())
            return ""
        try:
            _publish_workspace_grades(assignment_name, ref)
        except OSError as exc:
            st.session_state["save_status"] = (
                "error",
                f"Could not publish CAM's current grades for "
                f"'{assignment_name}' ({exc}) — launch cancelled, otherwise "
                "the workspace's next export could overwrite marks edited "
                "in CAM.")
            return ""
    params = {"class": cls, "assignment": target, "aname": assignment_name}
    url = f"http://127.0.0.1:{GRADING_PORT}/?{urlencode(params)}"
    try:
        webbrowser.open(url)
    except Exception:
        pass
    # Post-session scoped probe (Phase 1): for a round-tripping folder
    # assignment, remember what we just launched so the next reruns can cheaply
    # os.stat its export CSV(s) and auto-ingest a fresh export the moment the
    # teacher returns (see _run_active_launch_probe), and so the global scan
    # keeps its hands off that file meanwhile (active_skip in sync_from_cloud).
    # Exams don't round-trip, so they get no marker. Session-only — a restart
    # forgets it and the session-start global pass catches up.
    if ref and not target.startswith("exam::"):
        st.session_state["active_launch"] = {
            "class": cls,
            "assignment": assignment_name,
            "folder_ref": ref,
            "launched_at": time.time(),
            "last_probe": time.time(),
        }
    st.session_state["save_status"] = ("ok", f"Grading workspace launched → {url}")
    return url


def launch_exam_setup(exam_name: str) -> str:
    """Open the grading workspace's Exam Setup page for the active class,
    pre-filled with ``exam_name``. Returns the URL opened ('' on failure)."""
    if not _ensure_workspace_running():
        return ""
    # Best-effort seeding only: exams are keyed by class name and run on local
    # folders (no OAuth), so a local or unset master directory never blocks
    # Exam Setup — the Drive mapping is just a bonus for Drive-backed classes.
    _seed_workspace_class()
    params = {"class": st.session_state["active_class"], "exam": exam_name}
    url = f"http://127.0.0.1:{GRADING_PORT}/exam_setup?{urlencode(params)}"
    try:
        webbrowser.open(url)
    except Exception:
        pass
    st.session_state["save_status"] = ("ok", f"Exam setup launched → {url}")
    return url


def launch_drive_signin() -> str:
    """Open the grading workspace's /signin route, which runs the one-time
    Google OAuth flow and writes token.json — the token Watch needs to list
    Drive subfolders. Returns the URL opened ('' on failure)."""
    if not _ensure_workspace_running():
        return ""
    url = f"http://127.0.0.1:{GRADING_PORT}/signin"
    try:
        webbrowser.open(url)
    except Exception:
        pass
    st.session_state["save_status"] = ("ok", f"Google sign-in launched → {url}")
    return url


# --------------------------------------------------------------------------
# Term context (multi-term compression for the prompt engine)
# --------------------------------------------------------------------------

def current_term() -> str:
    """The term newly-ingested work is tagged to and the cockpit treats as live."""
    t = st.session_state.get("active_term", TERMS[0])
    return t if t in TERMS else TERMS[0]


def term_index(term: str) -> int:
    """Position of a term in the school year (unknown terms sort last)."""
    return TERMS.index(term) if term in TERMS else len(TERMS)


def assignment_term(asg) -> str:
    """The term an assignment was ingested into (its default-On term).

    A blank tag (legacy data ingested before terms existed) is treated as the
    active term so nothing silently drops out of the live cockpit."""
    return getattr(asg, "term", "") or current_term()


def ensure_term_context() -> None:
    """Point the flat ``active`` / ``llm_response`` aliases at the current
    term's slice of the per-term maps, so every existing read/write of those
    keys is automatically scoped to the term selected in the top bar."""
    ss = st.session_state
    t = current_term()
    ss["active"] = ss["active_by_term"].setdefault(t, {})
    ss["llm_response"] = ss["comments_by_term"].setdefault(t, {})


def term_active_map(term: str = "") -> dict:
    """The On-map ({assignment name -> bool}) for a term (default: current)."""
    t = term if term in TERMS else current_term()
    return st.session_state["active_by_term"].setdefault(t, {})


def _term_tag(name: str) -> str:
    """The ingest term tag of an assignment (blank when unknown/legacy)."""
    cls = st.session_state.get("active_class", "")
    fallback = ""
    for a in gb().assignments:
        if a.name != name:
            continue
        if getattr(a, "class_name", "") == cls:
            return getattr(a, "term", "")
        fallback = fallback or getattr(a, "term", "")
    return fallback


def assignment_on(name: str, term: str = "") -> bool:
    """Whether an assignment counts in a term's assessment (default: current).

    An explicit checkbox choice wins; otherwise an assignment defaults to On
    only in the term it was ingested into (blank legacy tags -> current term),
    so switching terms starts each term with just its own work selected."""
    t = term if term in TERMS else current_term()
    m = term_active_map(t)
    if name in m:
        return bool(m[name])
    return (_term_tag(name) or current_term()) == t


def current_term_assignment_names() -> set:
    """Active-class, non-archived assignment names that are On in the current
    term — i.e. the current term's assessment set."""
    cls = st.session_state.get("active_class", "")
    archived = st.session_state.get("archived", set())
    return {a.name for a in gb().assignments
            if getattr(a, "class_name", "") == cls
            and a.name not in archived
            and assignment_on(a.name)}


def term_assignment_names(term: str) -> set:
    """Active-class, non-archived assignment names that are On in ``term``."""
    cls = st.session_state.get("active_class", "")
    archived = st.session_state.get("archived", set())
    return {a.name for a in gb().assignments
            if getattr(a, "class_name", "") == cls
            and a.name not in archived
            and assignment_on(a.name, term)}


def term_has_class_data(term: str) -> bool:
    """True when the active class shows any trace of activity in ``term`` —
    assignments selected into it, saved comments, or a finalized summary."""
    if term_assignment_names(term):
        return True
    if st.session_state["comments_by_term"].get(term):
        return True
    cls = st.session_state.get("active_class", "")
    return bool(class_term_summaries(cls).get(term))


def student_worked_in_term(student, term: str) -> bool:
    """Whether the student has any recorded score in ``term``'s assessment set
    (False for a student who transferred in after that term)."""
    names = term_assignment_names(term)
    return any(sc.assignment in names
               for bucket in student.scores.values() for sc in bucket)


def missing_past_terms(student) -> list:
    """Prior terms that hold class data but no saved comment (finalized or
    in-app) for this student — the soft-alert list for multi-term generation."""
    cur_idx = term_index(current_term())
    have = {t for t, _ in past_term_context_for(student)}
    return [t for t in TERMS[:cur_idx]
            if t not in have and term_has_class_data(t)]


# --------------------------------------------------------------------------
# Data helpers
# --------------------------------------------------------------------------

def all_scores():
    """Yield every CriterionScore in the gradebook."""
    for student in gb():
        for bucket in student.scores.values():
            for score in bucket:
                yield student, score


def scores_for_assignment(name: str):
    return [(s, sc) for s, sc in all_scores() if sc.assignment == name]


def late_key(sid: str, assignment: str, crit: str) -> str:
    return f"{sid}||{assignment}||{crit}"


def is_late(sid: str, assignment: str, crit: str, score=None) -> bool:
    """Two-layer lateness read.

    A manual override in ``late_flags`` **wins when its key is present** — this
    is the teacher's CAM-side waive/force layer, set from the edit dialog's Late
    checkbox and persisted, so it survives re-syncs. When no override exists,
    fall back to the synced score's own ``late`` field (populated from the
    grading workspace's Late column at ingest). ``score`` is passed by call
    sites that already hold the CriterionScore; otherwise only the manual layer
    is consulted."""
    lf = st.session_state["late_flags"]
    key = late_key(sid, assignment, crit)
    if key in lf:
        return bool(lf[key])
    return bool(getattr(score, "late", False))


# ---- Missing = 0 / Excused / Awaiting Grade policy --------------------------
# A selected (On) assignment a student never submitted injects a mathematical
# 0 into every criterion it assesses: plotted on the Window 3 trend graph and
# fed to whichever calculation method is selected. The per-student "Excused"
# flag (Window 3 edit panel) removes an assignment from the trend, the grade
# calculation and the AI prompt entirely — for missing AND submitted work.
# EXCEPTION — Awaiting Grade (a TWO-STATE rule on folder-backed work):
# a folder-backed assignment (``folder_ref`` set) whose student has no score is
# read either as "still being graded" or as "grading finished, this student had
# no work", decided by the assignment's ``grading_complete`` flag — computed
# read-only from its most recent synced CSV in ``sync_from_cloud`` (every
# submitted row is graded):
#   * NOT complete  -> Awaiting Grade. A read-only "⏳ Awaiting Grade" row in
#     Window 3; contributes NOTHING to the trend, the grade math or the AI
#     prompt — no invented 0. Grades arrive via the workspace round-trip (§8).
#   * complete      -> the student falls through to the standard Missing = 0
#     policy exactly like a non-folder task: an editable "0 (missing)" row per
#     criterion and a real mathematical 0 in the trend / grade math / AI prompt
#     until the teacher enters a mark or ticks Excused. The physical 0 is a
#     deliberate product decision so unsubmitted folder work is noticed fast.
# A later export with a new ungraded submission flips the assignment back to
# not-complete on the next Sync, so the Awaiting pill returns (self-correcting).
# ``awaiting_grade(row)`` is the SINGLE predicate for this gate — every call
# site (missing math, Window 3, report export, Window 2 popover) uses it so the
# rule can never drift.

def awaiting_grade(row) -> bool:
    """True when a folder-backed assignment is still being graded — its pill is
    locked and it stays out of the math. Once ``grading_complete`` flips True a
    scoreless student falls through to the Missing = 0 policy instead."""
    return bool(row.get("folder_ref")) and not row.get("grading_complete")


def excused_key(sid: str, assignment: str) -> str:
    return f"{sid}||{assignment}"


def is_excused(sid: str, assignment: str) -> bool:
    return bool(st.session_state["excused_flags"].get(excused_key(sid, assignment)))


def set_excused(sid: str, assignment: str, value: bool) -> None:
    if value:
        st.session_state["excused_flags"][excused_key(sid, assignment)] = True
    else:
        st.session_state["excused_flags"].pop(excused_key(sid, assignment), None)


def excused_assignments_for(sid: str) -> set:
    """Assignment names excused for one student."""
    prefix = f"{sid}||"
    return {k[len(prefix):] for k, v in st.session_state["excused_flags"].items()
            if v and k.startswith(prefix)}


def missing_assignment_rows(student, names=None) -> list:
    """Assignment-table rows the student has no scores in, among the selected
    criteria-bearing assignments — the Missing = 0 candidates. Excused
    assignments are excluded (they leave the pipeline entirely), as are
    formative events and exams that haven't been banded yet (their criteria
    column is '—', so there is no 0-8 scale to zero)."""
    rows = assignment_table()
    if names is None:
        names = {r["name"] for r in rows if assignment_on(r["name"])}
    student_asgs = {sc.assignment for b in student.scores.values() for sc in b}
    out = []
    for r in rows:
        if r["name"] not in names or r["criteria"] == "—":
            continue
        # Folder-backed work still being graded is Awaiting Grade, not missing —
        # it must never invent a 0 (see the policy note above). Once its folder
        # grading is complete a scoreless student falls through here into the
        # normal Missing = 0 handling below.
        if awaiting_grade(r):
            continue
        # An exam with no banded 0-8 scores yet (e.g. manually created with a
        # target criterion) stays out of the math until banding happens.
        if r["is_exam"] and r["avg"] is None:
            continue
        if r["name"] in student_asgs:
            continue
        if is_excused(student.student_id, r["name"]):
            continue
        out.append(r)
    return out


def missing_zero_points(student, crit_letter: str, names=None) -> list:
    """Synthetic (timestamp, 0) points a criterion earns from missing work."""
    return [(r["date"], 0) for r in missing_assignment_rows(student, names)
            if crit_letter in [c.strip() for c in r["criteria"].split(",")]]


def aggregate_with_policy(student, crit_letter: str, names=None):
    """One criterion's aggregation under the Missing=0 / Excused policy, using
    this student's per-term calculation method (their pin, else the auto default)."""
    return aggregate_student_criterion(
        student, Criterion(crit_letter), method=calculation_method(student.student_id),
        include_assignments=names,
        extra_scores=missing_zero_points(student, crit_letter, names),
        exclude_assignments=excused_assignments_for(student.student_id) or None,
    )


def report_cfg() -> dict:
    """School-specific report-card figure toggles + Effort range (shared cfg)."""
    return st.session_state["report_cfg"]


def effort_bounds() -> tuple[int, int]:
    """The configured inclusive (min, max) for the Effort/English-use score."""
    rc = report_cfg()
    lo, hi = int(rc.get("effort_min", 0)), int(rc.get("effort_max", 5))
    return (hi, lo) if hi < lo else (lo, hi)


def effort_map() -> dict:
    """The CURRENT term's Effort/English-use map ({sid -> score})."""
    return st.session_state["effort_by_term"].setdefault(current_term(), {})


def student_effort(sid: str) -> int:
    """This student's current-term Effort score, clamped to the configured
    range (defaults to EFFORT_DEFAULT when unset, then clamped)."""
    v = effort_map().get(sid, EFFORT_DEFAULT)
    if not isinstance(v, int):
        v = EFFORT_DEFAULT
    lo, hi = effort_bounds()
    return max(lo, min(hi, v))


def student_term_grades(student):
    """One student's report-card grade bundle for the current term.

    The single source of truth for Window 3 AND every export, so they can
    never disagree. Each criterion's final grade resolves exactly as the
    grade panel does — locked override if present, else the auto band; a
    criterion with neither is N/A and simply doesn't count this term.

    Returns ``(n_criteria, crit_total, effort, myp, school)`` where
    ``n_criteria`` is how many criteria carry a grade, ``crit_total`` the
    integer sum of those grades (each defensively rounded — the School
    lookup tables index on whole numbers only), and ``myp``/``school`` are
    the school-lookup grades (``None`` -> shown as N/A)."""
    final = st.session_state["final_override"].get(student.student_id, {})
    n_criteria = 0
    crit_total = 0
    for c in CRIT_ORDER:
        if c in final:
            grade = final[c]
        else:
            res = aggregate_with_policy(student, c)
            grade = res.rounded_band if res else None
        if grade is None:
            continue   # criterion not assessed this term
        n_criteria += 1
        crit_total += int(round(grade))
    effort = student_effort(student.student_id)
    return (n_criteria, crit_total, effort,
            myp_grade(crit_total, n_criteria),
            school_grade(crit_total, effort, n_criteria))


def sync_active_into_scores() -> None:
    """Mirror the CURRENT term's per-assignment On checkboxes onto
    include_in_report, so the math/trend/prompt all follow the selected term."""
    on = {a.name: assignment_on(a.name) for a in gb().assignments}
    for _, score in all_scores():
        score.include_in_report = on.get(score.assignment, True)


def apply_date_override(name: str, new_date: date) -> None:
    """Push a teacher-chosen date onto every score of an assignment."""
    st.session_state["date_override"][name] = new_date
    for _, sc in scores_for_assignment(name):
        keep_time = sc.timestamp.time()
        sc.timestamp = datetime.combine(new_date, keep_time)


def academic_sort_key(dt: datetime):
    """Sort key that orders months April..March within the school year."""
    pos = ACADEMIC_MONTHS.index(dt.month) if dt.month in ACADEMIC_MONTHS else 99
    return (dt.year if dt.month >= 4 else dt.year - 1, pos, dt)


def assignment_table():
    """Return assignment rows enriched with analytics, academic-year sorted."""
    rows = []
    archived = st.session_state.get("archived", set())
    active_class = st.session_state.get("active_class", "")
    for asg in gb().assignments:
        if getattr(asg, "class_name", "") != active_class:
            continue  # belongs to a different class/level
        if asg.name in archived:
            continue  # soft-deleted: hidden from the active table & math
        pairs = scores_for_assignment(asg.name)
        valid = [sc.value for _, sc in pairs if sc.is_valid]
        lates = sum(
            1 for s, sc in pairs
            if is_late(s.student_id, asg.name, sc.criterion.value, sc)
        )
        # Display the assignment's actual date — the per-score timestamp parsed
        # from the CSV "Due Date" column — not ``ingested_at`` (merely when the
        # file was imported, i.e. today). Use the dominant date across scores,
        # keeping a real timestamp on it; fall back to the ingest time only for
        # formative events that carry no scores.
        score_times = [sc.timestamp for _, sc in pairs]
        if score_times:
            common_date = Counter(
                t.date() for t in score_times
            ).most_common(1)[0][0]
            when = min(t for t in score_times if t.date() == common_date)
        else:
            when = asg.ingested_at or datetime.now()
        if asg.name in st.session_state["date_override"]:
            ov = st.session_state["date_override"][asg.name]
            when = datetime.combine(ov, datetime.min.time())
        # Exams carry raw item-level results off the 0-8 scale; surface the
        # raw average (and count raw submissions) alongside any banded scores.
        is_exam = bool(getattr(asg, "is_exam", False))
        raw_totals = []
        if is_exam:
            raw_totals = [
                s.exam_results[asg.name].total for s in gb()
                if asg.name in getattr(s, "exam_results", {})
            ]
        rows.append({
            "name": asg.name,
            "criteria": ",".join(asg.criteria) if asg.criteria else "—",
            "is_formative": asg.is_formative,
            "is_exam": is_exam,
            "folder_ref": getattr(asg, "folder_ref", "") or "",
            "grading_complete": bool(getattr(asg, "grading_complete", False)),
            "max_total": int(getattr(asg, "max_total", 0)),
            "raw_avg": round(mean(raw_totals), 1) if raw_totals else None,
            "term": assignment_term(asg),
            "date": when,
            "submissions": max(len({s.student_id for s, _ in pairs}),
                               len(raw_totals)),
            "avg": round(mean(valid), 2) if valid else None,
            "spread": round(pstdev(valid), 2) if len(valid) > 1 else 0.0,
            "lates": lates,
            "n_scores": len(pairs),
        })
    rows.sort(key=lambda r: academic_sort_key(r["date"]))
    return rows


def student_label(student) -> str:
    return student.name or student.student_id


def submitter_keys(assignment: str) -> set:
    keys = set()
    for s, sc in scores_for_assignment(assignment):
        keys.add(s.student_id)
        if s.name:
            keys.add(s.name)
    return keys


def find_student(sid):
    return gb().students.get(sid)


# --------------------------------------------------------------------------
# Ingestion (from uploaded files -> temp path -> engine)
# --------------------------------------------------------------------------

def _save_temp(uploaded) -> str:
    suffix = os.path.splitext(uploaded.name)[1]
    fd, path = tempfile.mkstemp(suffix=suffix)
    with os.fdopen(fd, "wb") as fh:
        fh.write(uploaded.getbuffer())
    return path


def ingest_unit_plan(uploaded) -> None:
    path = _save_temp(uploaded)
    plan = parse_unit_plan(path)
    st.session_state["unit_plan"] = plan
    # Bind the unit plan to the active class so each class keeps its own.
    st.session_state["unit_plans"][st.session_state["active_class"]] = plan
    # Unit plans are now durable — persist immediately so an upload survives a
    # restart even without a later unrelated mutation to trigger auto-save.
    persist()


def commit_staged(sig: str, name: str, crit_letter, override_date) -> int:
    """Ingest one staged file into the gradebook and clear it from staging.

    ``crit_letter`` is a letter A-D for a manually-mapped generic Grade column,
    or None for auto-detected / formative files. ``override_date`` is an
    optional :class:`date` chosen on the staging calendar. Returns the count of
    scores created (0 for a formative event)."""
    item = st.session_state["staging"].get(sig)
    if not item:
        return 0
    fd, path = tempfile.mkstemp(suffix=".csv")
    with os.fdopen(fd, "wb") as fh:
        fh.write(item["data"])
    target = Criterion(crit_letter) if crit_letter else None
    gov = (datetime.combine(override_date, datetime.min.time())
           if override_date else None)
    try:
        if item.get("is_exam"):
            # Item-level exam export: raw marks live off the 0-8 scale until
            # the teacher assigns bands in Window 1's exam-banding panel.
            created = IngestionPipeline(gb()).ingest_exam_csv(
                path, assignment=name)
        else:
            created = IngestionPipeline(gb()).ingest_csv(
                path, assignment=name,
                manual_criterion_target=target,
                global_override_date=gov,
            )
    except ValueError as exc:
        st.session_state["save_status"] = ("error", f"{item['name']}: {exc}")
        return 0
    # Tag the freshly-registered assignment with the active class + term so it
    # shows only under that class/level (ingest_csv appends it last).
    if gb().assignments:
        gb().assignments[-1].class_name = st.session_state["active_class"]
        gb().assignments[-1].term = current_term()
    st.session_state["active"].setdefault(name, True)
    st.session_state["archived"].discard(name)
    st.session_state["ingested_sigs"].add(sig)
    del st.session_state["staging"][sig]
    persist()
    return len(created)


def _ensure_class(name: str) -> None:
    """Make sure a class/level exists so synced files have a home.

    Cloud-sync subfolders are named by class; if the teacher hasn't created
    that class yet we register it (with a blank roster) rather than dropping
    the file. Does not change the active class."""
    name = (name or "").strip()
    if not name or name in class_names():
        return
    st.session_state["classes"].append(
        {"name": name, "grade": "", "myp_year": "", "subject": ""}
    )
    st.session_state["rosters"].setdefault(name, [])


def _file_fingerprint(path: str) -> tuple:
    """Return ``(md5_hex, mtime)`` for a file, content-hashing its bytes.

    The hash detects modified-since-last-sync files even when the OS mtime is
    unreliable (cloud-drive downloads often reset it); mtime is kept alongside
    for display/debugging."""
    h = hashlib.md5()
    with open(path, "rb") as fh:
        for chunk in iter(lambda: fh.read(65536), b""):
            h.update(chunk)
    return h.hexdigest(), os.path.getmtime(path)


def _csv_grading_complete(path: str) -> Optional[bool]:
    """Read-only completeness check on one grading CSV — is every submitted work
    in the folder graded?

    Returns ``True``/``False`` for a criterion grading CSV, or ``None`` when the
    file must not gate the Awaiting pill (an exam CSV, or an unreadable file).
    This NEVER ingests or mutates the gradebook — it only opens and parses rows,
    preserving the sync invariant that an unchanged CSV is never re-ingested (a
    re-ingest would purge-replace the assignment and destroy CAM-side edits).

    Completeness rule: a row is *submitted* when its ``File Count`` cell parses
    to an int > 0; if that column is absent, fall back to a non-empty ``Files
    (newest first)`` cell; if both columns are absent (legacy CSV), every row is
    treated as submitted. A submitted row is *graded* when at least one cell is
    non-blank among columns whose header starts with ``Grade``. Complete = every
    submitted row is graded."""
    try:
        with open(path, "r", encoding="utf-8-sig", newline="") as fh:
            reader = csv.DictReader(fh)
            fieldnames = reader.fieldnames or []
            if is_exam_csv([(h or "").strip() for h in fieldnames]):
                return None  # exams never gate this pill
            file_count_col = files_col = None
            grade_cols = []
            for h in fieldnames:
                low = (h or "").strip().lower()
                if low == "file count":
                    file_count_col = h
                elif low == "files (newest first)":
                    files_col = h
                if low.startswith("grade"):
                    grade_cols.append(h)
            for row in reader:
                if file_count_col is not None:
                    raw = (row.get(file_count_col) or "").strip()
                    try:
                        submitted = int(raw) > 0
                    except ValueError:
                        submitted = False
                elif files_col is not None:
                    submitted = bool((row.get(files_col) or "").strip())
                else:
                    submitted = True  # legacy CSV: no submission columns
                if not submitted:
                    continue
                graded = any((row.get(c) or "").strip() for c in grade_cols)
                if not graded:
                    return False
            return True
    except (OSError, csv.Error):
        return None


def _rebind_import_name(incoming: str, class_name: str) -> str:
    """Re-map a round-tripped grading-CSV name back to the CAM assignment it
    was exported from.

    The grading workspace names every export after CAM's assignment name, but
    first collapses filesystem-illegal characters — ``/ \\ : * ? " < > |`` — to
    ``_`` (mirrored here by :func:`_safe_dirname`). So an assignment named
    ``"Maquette / Mock Up"`` returns from the round-trip as the file
    ``"Maquette _ Mock Up_Grades_<date>.csv"`` and its cleaned stem no longer
    equals the original name. Left unmapped, Sync treats it as a brand-new
    assignment: the original folder-backed row keeps its stale marks while a
    ``_``-mangled duplicate appears, and grades never sync back.

    Given the cleaned name read off the file, return the real name of an
    existing assignment in the same class whose sanitized name matches it, so
    the caller binds to that row instead. A folder-backed original outranks a
    ``_``-mangled orphan a past sync left behind (so marks reach the real row
    even before the orphan is cleaned up), and an exact-name match breaks any
    remaining tie (an assignment literally named ``"Maquette _ Mock Up"`` keeps
    its own row). Returns ``incoming`` unchanged when nothing matches."""
    if not incoming:
        return incoming
    peers = [a for a in gb().assignments
             if getattr(a, "class_name", "") == class_name]
    # Compare against each peer run through the SAME sanitize→clean round-trip
    # the export+sync applies, so trailing-separator trimming stays symmetric.
    matches = [a for a in peers
               if clean_assignment_name(_safe_dirname(a.name)) == incoming]
    if not matches:
        return incoming
    matches.sort(
        key=lambda a: (bool(getattr(a, "folder_ref", "")), a.name == incoming),
        reverse=True)
    return matches[0].name


def _sync_stamp_completeness(path: str, fname: str, class_name: str) -> bool:
    """Compute one CSV's grading completeness and stamp it onto the matching
    folder-backed assignment (name derived from the filename exactly as
    ``_ingest_cloud_file`` does, located by name + class). Only records with a
    ``folder_ref`` are touched. Returns ``True`` when a record's
    ``grading_complete`` value actually changed, so the caller knows to persist.

    Read-only on the CSV (see :func:`_csv_grading_complete`); an exam or
    unreadable file stamps nothing."""
    complete = _csv_grading_complete(path)
    if complete is None:
        return False
    name = _rebind_import_name(
        clean_assignment_name(os.path.splitext(fname)[0]), class_name)
    changed = False
    for a in gb().assignments:
        if (a.name == name
                and getattr(a, "class_name", "") == class_name
                and getattr(a, "folder_ref", "")):
            if bool(getattr(a, "grading_complete", False)) != complete:
                a.grading_complete = complete
                changed = True
    return changed


def _sync_reconcile_late(path: str, fname: str, class_name: str) -> int:
    """Read-only Late reconciliation for an export CSV Sync skipped as unchanged.

    Re-reads ONLY the ``Late`` column of a CSV whose bytes match the
    ``ingested_files`` registry hash and updates the synced ``late`` field of
    the matching ``csv:``-sourced scores in place — in **both** directions — so
    a flag the running ingest code dropped (the stale-ingest hole; see
    docs/LATE_FLAG_INTEGRITY_PLAN_V2.md §1) self-heals on the next Sync without
    a byte change and without re-ingesting (which would purge-replace grades the
    teacher edited in CAM since the export).

    Returns the number of scores whose ``late`` actually changed. A CSV with no
    ``Late`` header (a legacy export), an exam CSV, or an unreadable file
    reconciles nothing and returns 0 — never zero an existing flag. The manual
    ``late_flags`` override layer is NEVER touched by this pass."""
    try:
        with open(path, "r", encoding="utf-8-sig", newline="") as fh:
            reader = csv.DictReader(fh)
            fieldnames = reader.fieldnames or []
            if is_exam_csv([(h or "").strip() for h in fieldnames]):
                return 0
            # Resolve the Late column exactly first, then case/whitespace-
            # insensitively — mirroring ingest_csv (engine/ingestion.py:446).
            # No Late column -> legacy CSV: reconcile nothing.
            late_col = None
            for cand in fieldnames:
                if cand and cand.strip().lower() == "late":
                    late_col = cand
                    break
            if not late_col:
                return 0
            # Id column resolved exactly as ingest_csv does ("Student Name").
            # Resolve each CSV id through the durable alias map (Phase 3) so an
            # aliased row's Late heals the score living under the ROSTER id, not
            # the anonymous csv_key (which owns no score). Exact rows pass
            # through unchanged (alias miss → sid).
            aliases = st.session_state["work_aliases"].get(class_name, {})
            late_by_sid = {}
            for row in reader:
                sid = (row.get("Student Name") or "").strip()
                if not sid:
                    continue
                key = aliases.get(sid, sid)
                late_by_sid[key] = (row.get(late_col) or "").strip().lower() in {
                    "1", "true", "yes"}
    except OSError:
        return 0

    name = _rebind_import_name(
        clean_assignment_name(os.path.splitext(fname)[0]), class_name)
    changed = 0
    for student, sc in scores_for_assignment(name):
        sid = str(student.student_id)
        if sid not in late_by_sid:
            continue
        # Only the synced layer's own rows — never manual:missing-entry etc.
        if not str(getattr(sc, "source", "")).startswith("csv:"):
            continue
        want = late_by_sid[sid]
        if bool(getattr(sc, "late", False)) != want:
            sc.late = want
            changed += 1
    return changed


def _peek_csv(path: str):
    """Cheaply read a CSV's header row and first data row.

    Returns ``(fieldnames, first_row_or_None)`` — ``first_row`` is a dict keyed
    by header, or ``None`` when the file has no data rows. Raises ``OSError`` on
    a read failure so callers can decide how to degrade."""
    with open(path, "r", encoding="utf-8-sig", newline="") as fh:
        reader = csv.DictReader(fh)
        fieldnames = reader.fieldnames or []
        first = next(reader, None)
    return fieldnames, first


def _duplicate_file_kind(path: str, fname: str) -> str:
    """Classify one member of a duplicate-export group.

    ``"canonical"`` — the filename date matches the file's own ``Due Date``
    content (the durable, due-date-named export); ``"fallback"`` — they differ
    (an export-date fallback, written while the deadline was missing — likely
    stale); ``"unknown"`` — no data row, or no parseable Due Date / filename
    date, so we never guess."""
    try:
        fieldnames, first = _peek_csv(path)
    except OSError:
        return "unknown"
    if not first:
        return "unknown"
    due_col = next((h for h in (fieldnames or [])
                    if (h or "").strip().lower() == "due date"), None)
    due_raw = (first.get(due_col) or "").strip() if due_col else ""
    due_dt = parse_iso_date(due_raw) or parse_date_from_filename(due_raw)
    fname_dt = parse_date_from_filename(os.path.splitext(fname)[0])
    if not due_dt or not fname_dt:
        return "unknown"
    return "canonical" if due_dt.date() == fname_dt.date() else "fallback"


def _scan_class_duplicate_groups(class_path: str) -> dict:
    """Group a class subtree's grading CSVs by cleaned assignment name.

    Returns only the groups with two or more members (the duplicate exports),
    as ``{assignment_name: [abspath, ...]}``. Exam CSVs are excluded from
    grouping — they keep their existing handling and may legitimately share a
    stem with a grading CSV. Unreadable files are left for the main loop to
    report; they never join a group."""
    groups: dict = {}
    for dirpath, _dirs, files in os.walk(class_path):
        for fname in sorted(files):
            if not fname.lower().endswith(".csv"):
                continue
            path = os.path.abspath(os.path.join(dirpath, fname))
            try:
                fieldnames, _first = _peek_csv(path)
            except OSError:
                continue
            if is_exam_csv([(h or "").strip() for h in fieldnames]):
                continue
            name = clean_assignment_name(os.path.splitext(fname)[0])
            groups.setdefault(name, []).append(path)
    return {name: paths for name, paths in groups.items() if len(paths) > 1}


def _format_duplicate_message(name: str, class_name: str, paths: list,
                              registry: dict) -> str:
    """Build the teacher-facing alert for one duplicate-export group.

    Identifies the canonical/original file (filename date == its own Due Date)
    versus export-date fallbacks, and shows each file's registry sync status and
    disk mtime. Never guesses when the filename dates don't single out exactly
    one canonical export."""
    kinds = {p: _duplicate_file_kind(p, os.path.basename(p)) for p in paths}
    canon = [p for p, k in kinds.items() if k == "canonical"]
    lines = [f"⚠ Duplicate exports for **{name}** in *{class_name}* — nothing "
             f"was synced for this assignment:"]
    for p in sorted(paths, key=lambda x: os.path.basename(x)):
        base = os.path.basename(p)
        kind = kinds[p]
        if kind == "canonical":
            tag = "filename matches its Due Date (**canonical export**)"
        elif kind == "fallback":
            tag = "filename is an export-date fallback (**likely stale**)"
        else:
            tag = "filename could not be matched against its Due Date (**unverified**)"
        prior = registry.get(p) or {}
        synced = prior.get("ingested_at")
        synced_txt = f"last synced {synced}" if synced else "never synced"
        try:
            mtime_txt = datetime.fromtimestamp(
                os.path.getmtime(p)).isoformat(timespec="seconds")
        except OSError:
            mtime_txt = "unknown"
        lines.append(f"• `{base}` — {tag}; {synced_txt}; disk mtime {mtime_txt}.")
    if len(canon) == 1:
        lines.append("Open both, verify which holds the correct grades/Late "
                     "flags, delete the other, then Sync again.")
    else:
        lines.append("Filename dates don't single out one canonical export "
                     "here — open each, verify which holds the correct "
                     "grades/Late flags, delete the rest, then Sync again.")
    return "\n".join(lines)


# How often the post-session scoped probe (see _run_active_launch_probe) is
# allowed to os.stat the active assignment's export CSV(s). One cheap check per
# ~30s of reruns keeps a mid-grading return from re-scanning on every widget
# interaction while still catching a fresh export within half a minute.
ACTIVE_LAUNCH_PROBE_INTERVAL = 30.0


def _assignment_csv_paths(class_name: str, assignment_name: str) -> list:
    """Every grading-CSV abspath in ``class_name``'s data folder that ingest
    would key to ``assignment_name``.

    Uses the same cleaned-name + rebind round-trip ``_ingest_cloud_file`` uses,
    so a filesystem-mangled name (``Maquette _ Mock Up``) still resolves to its
    real row. Exam CSVs are excluded (they never key to a grading assignment).
    Returns ``[]`` when the class folder or its CSVs are absent/unreadable."""
    class_path = class_data_dir(class_name)
    if not os.path.isdir(class_path):
        return []
    out = []
    for dirpath, _dirs, files in os.walk(class_path):
        for fname in sorted(files):
            if not fname.lower().endswith(".csv"):
                continue
            path = os.path.abspath(os.path.join(dirpath, fname))
            try:
                fieldnames, _first = _peek_csv(path)
            except OSError:
                continue
            if is_exam_csv([(h or "").strip() for h in fieldnames]):
                continue
            name = _rebind_import_name(
                clean_assignment_name(os.path.splitext(fname)[0]), class_name)
            if name == assignment_name:
                out.append(path)
    return out


def _sync_one_csv(path: str, fname: str, class_name: str,
                  registry: dict, summary: dict) -> tuple:
    """Fingerprint one grading CSV against the registry and act on it.

    The per-file core shared by the global scan (:func:`sync_from_cloud`) and
    the scoped scan (:func:`sync_assignment`) so both inherit — identically, not
    by re-implementation — the graceful parse-failure branch, the read-only
    completeness stamp + Late reconcile on an unchanged file, the purge-replace
    ingest on a changed one, and the Late-count tripwire. Mutates ``summary``
    and ``registry`` in place and increments ``summary["found"]`` for this file.

    Returns ``(ingested, changed)``: ``ingested`` is True when a real ingest
    happened (a new file or a changed re-ingest); ``changed`` is True when
    anything needs persisting (an ingest, a completeness flip, or a Late
    reconcile). The caller ORs these across files to decide whether to persist.

    Torn-read hardening: an unreadable file (mid-transfer / mid-export) counts
    as an error and leaves the registry untouched, so the next pass retries —
    never a traceback. A parse failure inside :func:`_ingest_cloud_file` returns
    ``None`` there and is likewise counted as an error here."""
    summary["found"] += 1
    try:
        file_hash, mtime = _file_fingerprint(path)
    except OSError as exc:
        summary["errors"] += 1
        summary["messages"].append(f"Unreadable: {fname} ({exc})")
        return (False, False)

    changed = False
    prior = registry.get(path)
    if prior and prior.get("hash") == file_hash:
        summary["skipped"] += 1          # already in the database
        # Completeness pass runs for EVERY csv, including files skipped as
        # unchanged — this unlocks assignments whose CSVs were synced before
        # this feature existed. It is strictly READ-ONLY on the CSV: never
        # re-ingest an unchanged file (that would purge-replace the assignment
        # and destroy grades the teacher edited in CAM since the last export).
        if _sync_stamp_completeness(path, fname, class_name):
            changed = True
        # Read-only Late reconciliation: re-read just the Late column of this
        # unchanged CSV and heal any synced `late` field an older ingest
        # dropped. Never re-ingests, never touches the manual late_flags layer.
        n_rec = _sync_reconcile_late(path, fname, class_name)
        if n_rec:
            changed = True
            summary["reconciled"] += n_rec
            rec_name = _rebind_import_name(
                clean_assignment_name(os.path.splitext(fname)[0]), class_name)
            summary["messages"].append(
                f"🩹 {rec_name}: reconciled {n_rec} synced Late "
                f"flag(s) from the unchanged export CSV.")
        return (False, changed)

    is_update = prior is not None

    # Late-count tripwire (advisory): a changed file re-ingesting an existing
    # assignment purge-replaces its scores. Count the synced-layer Late flags
    # before the ingest so we can tell if this sync silently zeroed them (the
    # incident's signature). Use the score's own `late` field ONLY — NOT
    # is_late(), which mixes in the teacher's manual CAM overrides.
    tripwire_name = ""
    late_before = 0
    if is_update:
        tripwire_name = _rebind_import_name(
            clean_assignment_name(os.path.splitext(fname)[0]), class_name)
        late_before = sum(
            1 for _, sc in scores_for_assignment(tripwire_name)
            if getattr(sc, "late", False))

    n = _ingest_cloud_file(path, fname, class_name, replace_prior=is_update,
                           summary=summary)
    if n is None:
        summary["errors"] += 1
        return (False, changed)
    # Stamp AFTER the ingest so the freshly-created record (which carries the
    # folder_ref forward) picks up the flag.
    if _sync_stamp_completeness(path, fname, class_name):
        changed = True

    registry[path] = {
        "hash": file_hash,
        "mtime": mtime,
        "class": class_name,
        "assignment": clean_assignment_name(os.path.splitext(fname)[0]),
        "ingested_at": datetime.now().isoformat(),
    }
    summary["scores"] += n
    summary["classes"].add(class_name)
    if is_update:
        summary["updated"] += 1
        late_after = sum(
            1 for _, sc in scores_for_assignment(tripwire_name)
            if getattr(sc, "late", False))
        if late_before > 0 and late_after < late_before:
            summary["messages"].append(
                f"⚠ {tripwire_name}: synced Late flags dropped "
                f"{late_before} → {late_after} in this sync. If "
                f"unexpected, check the export CSV in the class "
                f"folder before trusting it.")
    else:
        summary["ingested"] += 1
    return (True, changed)


def sync_assignment(class_name: str, assignment_name: str) -> dict:
    """Scoped Sync of one assignment's export CSV(s) in one class.

    The assignment-granular counterpart to :func:`sync_from_cloud`: it runs the
    identical per-CSV machinery (via the shared :func:`_sync_one_csv`) but only
    over the CSV(s) that ingest keys to ``assignment_name`` in ``class_name``'s
    data folder, and inherits the same duplicate pre-pass — 2+ dated exports of
    this assignment are refused as a whole group, nothing ingested, the same
    alert raised. Returns the same summary shape as ``sync_from_cloud`` so
    :func:`_report_sync` renders it.

    Used at 🖌 launch time to close the stale-handoff race (Terrain §T4 — CAM
    ingests any fresher export before it publishes its own values) and by the
    post-session probe (:func:`_run_active_launch_probe`). A caller cancels a
    launch by inspecting ``summary["duplicates"]`` / ``summary["errors"]``."""
    summary = {"found": 0, "ingested": 0, "updated": 0, "skipped": 0,
               "errors": 0, "scores": 0, "duplicates": 0, "reconciled": 0,
               "classes": set(), "messages": [], "duplicate_messages": []}
    custom = (st.session_state.get("prefs", {}).get("db_custom_path") or "").strip()
    if not custom:
        return summary   # no cloud_dir routing -> nothing to sync (quiet no-op)
    class_path = class_data_dir(class_name)
    if not os.path.isdir(class_path):
        return summary

    registry = st.session_state["ingested_files"]
    # Same duplicate pre-pass as the global scan, but only THIS assignment's
    # group matters. The pre-pass keys groups by the raw cleaned filename, so
    # match the target directly and — for a filesystem-mangled name — through
    # the same rebind round-trip ingest uses.
    dup_groups = _scan_class_duplicate_groups(class_path)
    my_dupes = dup_groups.get(assignment_name)
    if my_dupes is None:
        for gname, paths in dup_groups.items():
            if _rebind_import_name(gname, class_name) == assignment_name:
                my_dupes = paths
                break
    if my_dupes:
        summary["found"] += len(my_dupes)
        summary["duplicates"] += len(my_dupes)
        summary["duplicate_messages"].append(
            _format_duplicate_message(
                assignment_name, class_name, my_dupes, registry))
        return summary

    any_ingested = False
    any_change = False
    for path in _assignment_csv_paths(class_name, assignment_name):
        ingested, changed = _sync_one_csv(
            path, os.path.basename(path), class_name, registry, summary)
        any_ingested = any_ingested or ingested
        any_change = any_change or changed

    if any_ingested or any_change:
        if st.session_state["active_class"] not in class_names() and class_names():
            st.session_state["active_class"] = class_names()[0]
        ensure_class_context()
        # A purge-replace rebuilt scores from the CSV(s) and so may have dropped
        # CAM-typed score comments — refill them from the cloud twin before we
        # persist (and re-mirror), so they survive the sync and reach disk again.
        _heal_score_comments_from_mirrors()
        persist()
    return summary


def sync_from_cloud() -> dict:
    """Scan the database folder's class subfolders and ingest new/modified CSVs.

    The database folder (Settings → Custom Database Path) holds one subfolder
    per class; every ``.csv`` inside a subfolder is a grading export for that
    class. For each file we:

      * compute a content hash and consult the ``ingested_files`` registry
        (stored in ``acm_database.json``) — files already ingested with an
        unchanged hash are skipped, so re-running sync is idempotent;
      * ingest new/modified files, assigning them to the class named by their
        subfolder (creating the class if needed) and dating each row from its
        durable ``Due Date`` (deadline) column (legacy ``Assessed Date`` and
        filename dates fall back gracefully);
      * record the file's hash so the next sync ignores it.

    Returns a summary dict ``{found, ingested, updated, skipped, errors,
    scores, classes, messages}`` for the UI to surface.
    """
    summary = {"found": 0, "ingested": 0, "updated": 0, "skipped": 0,
               "errors": 0, "scores": 0, "duplicates": 0, "reconciled": 0,
               "classes": set(), "messages": [], "duplicate_messages": []}
    custom = (st.session_state.get("prefs", {}).get("db_custom_path") or "").strip()
    if not custom:
        summary["messages"].append(
            "Set a Custom Database Path first (⚙ Settings) — Sync scans that "
            "folder's class subfolders for grading CSVs.")
        return summary
    root = db_folder()
    if not os.path.isdir(root):
        summary["messages"].append(f"Database folder not found: {root}")
        return summary

    registry = st.session_state["ingested_files"]
    any_ingested = False
    any_change = False   # a completeness stamp changed a record -> persist

    # Active-launch guard (Phase 1): while a 🖌 grading session is open, its
    # scoped probe (_run_active_launch_probe) owns that assignment's export
    # CSV(s). The global scan must NOT touch them mid-grading — a re-ingest of
    # an in-flight export would race the probe. Skip them here; the probe
    # reclaims them on the next return, and the marker clears once the newer
    # export is ingested. (A *different* machine's live session is unknowable —
    # accepted residual risk, with the read-only Late reconcile as backstop.)
    marker = st.session_state.get("active_launch")
    active_skip = set()
    if marker:
        active_skip = set(_assignment_csv_paths(
            marker.get("class", ""), marker.get("assignment", "")))

    # One-time cleanup of materialised overrides (Phase C — see
    # docs/LATE_FLAG_INTEGRITY_PLAN_V2.md §5). The old edit dialog wrote
    # late_flags[key] on EVERY Save, so hundreds of "overrides" merely echo the
    # synced value; each redundant False silently suppresses a future CSV-synced
    # Late. Delete every key whose bool equals the matching score's synced
    # `late` (missing score -> False); keep every DIFFERING key (the deliberate
    # waives/forces). Guarded by a persisted marker so it runs exactly once.
    # It MUST run BEFORE the reconciliation below (§5.3): a redundant False
    # sitting on a still-stale-False cell whose CSV says Late=1 has to be
    # deleted while it is still redundant — after reconciliation it would differ
    # from the healed True and survive forever as a phantom waive.
    if not st.session_state.get("late_flags_cleanup_v1", False):
        lf = st.session_state["late_flags"]
        synced_late = {}
        for student, sc in all_scores():
            synced_late[late_key(str(student.student_id), sc.assignment,
                                 sc.criterion.value)] = bool(
                                     getattr(sc, "late", False))
        removed = 0
        for k in list(lf.keys()):
            if bool(lf[k]) == synced_late.get(k, False):
                del lf[k]
                removed += 1
        kept = len(lf)
        st.session_state["late_flags_cleanup_v1"] = True
        any_change = True   # persist the marker even if nothing else changes
        summary["messages"].append(
            f"🧹 One-time cleanup: removed {removed} redundant Late "
            f"override(s) left by the old edit dialog; {kept} deliberate "
            f"waive(s)/force(s) kept.")

    # Each immediate subfolder is a class; scan it (recursively) for CSVs.
    for entry in sorted(os.scandir(root), key=lambda e: e.name):
        if not entry.is_dir():
            continue
        class_name = entry.name
        # Pre-pass: one assignment must map to exactly one CSV. When a class
        # folder holds two differently-dated exports of the same assignment
        # (Flaw B — the incident), whichever is (re)ingested last silently
        # wins, so we refuse the WHOLE group: no ingest, no registry update, no
        # completeness stamp for any member. Alert the teacher and let them
        # verify + delete one; we never auto-tiebreak or auto-delete (D2).
        dup_groups = _scan_class_duplicate_groups(entry.path)
        dup_skip = {p for paths in dup_groups.values() for p in paths}
        for name in sorted(dup_groups):
            summary["duplicate_messages"].append(
                _format_duplicate_message(
                    name, class_name, dup_groups[name], registry))
        for dirpath, _dirs, files in os.walk(entry.path):
            for fname in sorted(files):
                if not fname.lower().endswith(".csv"):
                    continue
                path = os.path.abspath(os.path.join(dirpath, fname))
                # A live scoped probe owns this file mid-grading — skip it
                # entirely (invisible to this pass) so the two never race.
                if path in active_skip:
                    continue
                if path in dup_skip:
                    summary["found"] += 1
                    summary["duplicates"] += 1
                    continue
                # Shared per-file core: fingerprint, ingest-or-read-only-passes,
                # tripwire, graceful parse failure (see _sync_one_csv).
                ingested, changed = _sync_one_csv(
                    path, fname, class_name, registry, summary)
                any_ingested = any_ingested or ingested
                any_change = any_change or changed

    if any_ingested or any_change:
        # Keep the active class valid and mirror everything to disk (the
        # registry now lives inside acm_database.json's session payload).
        if st.session_state["active_class"] not in class_names() and class_names():
            st.session_state["active_class"] = class_names()[0]
        ensure_class_context()
        # Refill any CAM-typed score comment the purge-replace above dropped,
        # from the cloud twin, before persisting (and re-mirroring) the rebuild.
        _heal_score_comments_from_mirrors()
        persist()
    return summary


def sync_all() -> dict:
    """The Sync button: scan the database folder AND every class's master
    directory in one pass.

    1. The database folder's class subfolders are scanned for new/modified
       grading CSVs (see :func:`sync_from_cloud`).
    2. Every class with a master directory set gets a Watch pass, so new
       assignment/exam folders appear in the list without visiting each class.
    """
    summary = sync_from_cloud()
    summary["watch_found"] = 0
    summary["watch_created"] = 0
    for c in st.session_state["classes"]:
        ref = str(c.get("master_dir", "") or "").strip()
        if not ref:
            continue
        try:
            found, created = _watch_class_master(c["name"], ref)
        except Exception as exc:
            summary["errors"] += 1
            summary["messages"].append(f"Watch failed for {c['name']}: {exc}")
            continue
        summary["watch_found"] += found
        summary["watch_created"] += created
    return summary


def _report_sync(summary: dict) -> None:
    """Turn a sync summary into a clean status banner for the UI."""
    new_total = summary["ingested"] + summary["updated"]
    errors = summary["errors"]
    dups = summary.get("duplicate_messages") or []
    watch_note = ""
    if summary.get("watch_created"):
        watch_note = (f" 👁 {summary['watch_created']} new assignment/exam "
                      f"folder(s) added from class master directories.")
    # Pre-flight problems (no folder / missing folder) carry a message and no
    # files found -> surface them directly as an error.
    if summary["found"] == 0 and summary["messages"]:
        st.session_state["save_status"] = ("error", summary["messages"][0])
        return
    if summary["found"] == 0:
        st.session_state["save_status"] = (
            "ok", "Database folder scanned — no .csv files found in any class "
                  "subfolder." + watch_note)
        return
    if new_total == 0:
        if summary["skipped"]:
            msg = (f"All {summary['skipped']} file(s) already in the database — "
                   f"nothing new to sync.")
        else:
            msg = "No new grading data synced."
        kind = "ok"
    else:
        parts = [f"Found {new_total} new file(s), synced successfully"]
        detail = []
        if summary["updated"]:
            detail.append(f"{summary['updated']} updated")
        detail.append(f"{summary['scores']} score(s)")
        if summary["classes"]:
            detail.append(f"{len(summary['classes'])} class(es): "
                          + ", ".join(sorted(summary["classes"])))
        if summary["skipped"]:
            detail.append(f"{summary['skipped']} unchanged skipped")
        msg = parts[0] + " — " + "; ".join(detail) + "."
        kind = "ok"
    msg += watch_note
    if errors:
        msg += f" ⚠ {errors} step(s) failed: " + " | ".join(summary["messages"][-2:])
        kind = "error"
    elif summary["messages"]:
        # Advisory notes with no hard error — e.g. the Late-count tripwire.
        # Surface them without downgrading a successful sync to error state.
        msg += "  " + "  ".join(summary["messages"])
    # Duplicate exports are never a clean "ok": one assignment mapping to two
    # CSVs means Sync skipped it entirely, so surface it prominently (error
    # styling) with the per-group detail even when everything else succeeded.
    if dups:
        kind = "error"
        header = (f"⚠ {len(dups)} assignment(s) have duplicate exports — "
                  f"{summary['duplicates']} file(s) were NOT synced. Resolve "
                  f"these before trusting this sync:")
        msg = msg + "\n\n" + header + "\n\n" + "\n\n".join(dups)
    st.session_state["save_status"] = (kind, msg)


def _run_session_start_sync() -> None:
    """One automatic global catch-up per session, right after the DB loads.

    The db folder lives in OneDrive: CSVs appear/change with no local user
    action (another machine, hand-edits in Excel), so an event-driven-only sync
    would miss them — a global scan must still run *sometime*. This is that
    sometime. Guarded to run exactly once per session; a full ``sync_all`` so it
    also picks up new assignment/exam folders from every class master directory.

    For an unconfigured install (no Custom Database Path) it is a quiet no-op —
    no banner at all — so a first-run user isn't greeted by a sync error. When a
    path IS set, the result is surfaced through :func:`_report_sync` exactly as
    the manual re-scan is, so a duplicate-refusal or Late-tripwire is as loud at
    startup as it was from the old button."""
    if st.session_state.get("session_sync_done"):
        return
    st.session_state["session_sync_done"] = True
    custom = (st.session_state.get("prefs", {}).get("db_custom_path") or "").strip()
    if not custom:
        return   # unconfigured install -> stay silent (no error banner)
    summary = sync_all()
    _report_sync(summary)


def _run_active_launch_probe() -> None:
    """Post-🖌-session scoped probe: catch a fresh export when the teacher returns.

    CGW cannot push to CAM (separate processes), and Streamlit only reruns on
    interaction, so "sync on export" is approximated here: while an
    ``active_launch`` marker is held, each rerun — throttled to
    ~``ACTIVE_LAUNCH_PROBE_INTERVAL`` seconds — cheaply ``os.stat``s just that
    assignment's export CSV(s) against the registry mtimes (one stat per file,
    not a tree walk). On any change it runs the scoped :func:`sync_assignment`
    and surfaces the banner, then clears the marker once the newer export is
    actually ingested — handing the file back to the global scan.

    A duplicate-dated group keeps the marker (the teacher must resolve it; we
    keep re-nagging every probe). A torn read mid-export raises ``OSError`` on
    the stat and is skipped — the next probe retries — never a traceback."""
    marker = st.session_state.get("active_launch")
    if not marker:
        return
    now = time.time()
    if now - marker.get("last_probe", 0.0) < ACTIVE_LAUNCH_PROBE_INTERVAL:
        return
    marker["last_probe"] = now
    cls = marker.get("class", "")
    assignment = marker.get("assignment", "")
    registry = st.session_state["ingested_files"]
    # Cheap staleness probe: a brand-new export creates a path absent from the
    # registry; a re-export of an already-synced assignment bumps its mtime.
    # Either way sync_assignment re-checks by content hash before ingesting, so
    # a spurious mtime bump costs at most one hash and ingests nothing.
    stale = False
    for path in _assignment_csv_paths(cls, assignment):
        prior = registry.get(path)
        if not prior:
            stale = True
            break
        try:
            if abs(float(prior.get("mtime", 0)) - os.path.getmtime(path)) > 1e-6:
                stale = True
                break
        except OSError:
            continue   # mid-export torn read; next probe retries
    if not stale:
        return
    summary = sync_assignment(cls, assignment)
    _report_sync(summary)
    if (summary["ingested"] or summary["updated"]) and not summary["duplicates"]:
        st.session_state["active_launch"] = None   # global scan reclaims it


def _assignment_richness(a) -> tuple:
    """Rank an assignment record by how much real signal it carries.

    Used to pick the survivor when collapsing duplicate metadata records for the
    same name+class: an exam, one with scores/criteria, or one backed by a
    source file outranks a bare timeline placeholder."""
    return (
        bool(getattr(a, "is_exam", False)),
        int(getattr(a, "score_count", 0) or 0),
        len(getattr(a, "criteria", []) or []),
        bool(getattr(a, "source_file", "")),
    )


def _dedupe_assignments() -> int:
    """Collapse duplicate assignment records sharing (name, class_name).

    Scores live on students keyed by assignment *name*, so two Assignment
    records for the same name+class never double-count grades — but they DO
    collide on the timeline's per-name widget keys and crash the render with
    ``StreamlitDuplicateElementKey``. This heals a database that already picked
    up such a duplicate (e.g. a sync that appended a graded import next to the
    handoff's placeholder before the always-purge fix landed): keep the richest
    record per name+class, preserving timeline order. Returns how many were
    removed."""
    seen: dict = {}
    order: list = []
    for a in gb().assignments:
        key = (a.name, getattr(a, "class_name", ""))
        if key not in seen:
            seen[key] = a
            order.append(key)
        elif _assignment_richness(a) > _assignment_richness(seen[key]):
            seen[key] = a
    deduped = [seen[k] for k in order]
    removed = len(gb().assignments) - len(deduped)
    if removed:
        gb().assignments = deduped
    return removed


def _purge_assignment_in_class(name: str, class_name: str) -> None:
    """Remove a prior import of ``name`` in ``class_name`` (record + scores).

    Used when a watch-folder file is re-synced after being modified, so the
    refreshed grades replace the old ones instead of double-counting. Scoped to
    the assignment's class so a same-named unit in another class is untouched."""
    gbk = gb()
    gbk.assignments = [
        a for a in gbk.assignments
        if not (a.name == name and getattr(a, "class_name", "") == class_name)
    ]
    # If this assignment name still exists under another class, leave every
    # student's scores alone; otherwise purge the now-orphaned scores.
    still_used = any(a.name == name for a in gbk.assignments)
    if not still_used:
        for student in gbk:
            for crit, bucket in list(student.scores.items()):
                student.scores[crit] = [
                    sc for sc in bucket if sc.assignment != name
                ]
            # Raw exam results ride along with their assignment.
            getattr(student, "exam_results", {}).pop(name, None)


def _ingest_cloud_file(path: str, fname: str, class_name: str,
                       replace_prior: bool = False, summary: dict = None):
    """Ingest one watch-folder CSV into ``class_name``; return score count.

    Returns the number of scores created (>=0) on success, or ``None`` if the
    file could not be ingested (e.g. a generic 'Grade' column whose criterion
    can't be inferred from the filename). The criterion target is inferred from
    the filename ("... (Crit B) ...") for legacy generic-grade exports; durable
    multi-criterion files map themselves and ignore the hint. When
    ``replace_prior`` is accepted for call-site compatibility; the prior copy of
    this assignment in this class is now always purged before ingesting (see
    below), whether the flag is set or not.

    **Roster-aware identity routing (Phase 3).** When the target class has a
    roster, the grading CSV's rows are routed through :func:`resolve_identity`
    (exact → durable alias → unambiguous prefix → pool) instead of minting a
    phantom student for every unmatched ``Student Name`` cell. Unmatched rows are
    rebuilt into ``unmatched_works[class][assignment]`` (purge-replace, like the
    scores); fast-path prefix matches are recorded into the durable
    ``work_aliases[class]`` and announced on ``summary`` when one is given. A
    rosterless class skips all of this (legacy behaviour — the score-only
    folder-graded-before-roster path stays intact). Exam CSVs never route."""
    # The export names files "<assignment>_Grades_<deadline>.csv"; trim the
    # Grades tag and trailing date so the timeline shows just the clean name
    # (an embedded "(Crit X)" hint is preserved for criterion detection).
    assignment = clean_assignment_name(os.path.splitext(fname)[0])
    # A '/', '\', ':' etc. in the name can't survive as a filename — the
    # workspace sanitized it to '_' when naming the export — so re-bind to the
    # assignment it really came from before the purge/ingest below. Skipped,
    # the graded import lands on a '_'-mangled orphan and the marks never reach
    # the original folder-backed row (which keeps its stale grades).
    assignment = _rebind_import_name(assignment, class_name)
    target = detect_criterion_in_header(assignment)  # works on any string
    # Always replace any prior copy of this assignment in this class before
    # ingesting — a manual timeline placeholder (the CAM→CGW handoff creates
    # one, source_file="") OR an earlier import of a since-modified file. Left
    # in place, a first-time sync would APPEND a second assignment of the same
    # name; the timeline keys its per-row widgets on the name, so the duplicate
    # crashes the render (StreamlitDuplicateElementKey). Purging first makes the
    # graded import update the placeholder in place.
    #
    # Preserve the folder pin across the purge: the prior copy may be a
    # folder-backed assignment (Watch-created or CAM-renamed), and a re-sync
    # would otherwise strip its ``folder_ref`` — leaving it "not a folder
    # assignment" and inviting Watch to spawn a parallel row. Carry the ref
    # onto the freshly-ingested record so a graded assignment stays welded to
    # its source folder through every round-trip.
    prior_ref = next((getattr(a, "folder_ref", "") for a in gb().assignments
                      if a.name == assignment
                      and getattr(a, "class_name", "") == class_name
                      and getattr(a, "folder_ref", "")), "")
    _purge_assignment_in_class(assignment, class_name)
    # Roster-aware routing inputs (Phase 3): the class roster's keys and the
    # durable alias map. An empty roster leaves both falsy → ingest_csv skips
    # routing entirely (legacy behaviour). Aliases are copied so the engine can't
    # mutate the persisted map.
    roster_keys = {e.get("key") for e in st.session_state["rosters"].get(
        class_name, []) if e.get("key")}
    aliases = dict(st.session_state["work_aliases"].get(class_name, {}))
    unmatched: list = []
    auto_aliases: dict = {}
    try:
        # Exam exports (item-level, raw marks) route to the exam pipeline; the
        # header decides — a "Total Score" column marks a CAM exam CSV.
        with open(path, "r", encoding="utf-8-sig", newline="") as fh:
            header = [h.strip() for h in (next(csv.reader(fh), []) or [])]
        if is_exam_csv(header):
            created = IngestionPipeline(gb()).ingest_exam_csv(
                path, assignment=assignment)
        else:
            created = IngestionPipeline(gb()).ingest_csv(
                path, assignment=assignment, manual_criterion_target=target,
                roster_keys=roster_keys or None,
                aliases=aliases or None,
                unmatched_out=unmatched,
                auto_aliases_out=auto_aliases,
            )
    except (ValueError, OSError) as exc:
        st.session_state["save_status"] = (
            "error", f"Skipped '{fname}': {exc}")
        return None

    # Rebuild this assignment's unmatched-works pool (purge-replace: the CSV is
    # the durable source, so re-ingesting re-derives the pool from scratch). A
    # newly-recorded prefix auto-alias is merged into the DURABLE map (never
    # rebuilt) and surfaced on the sync banner so silent routing stays visible.
    if roster_keys:
        pool_map = st.session_state["unmatched_works"].setdefault(class_name, {})
        if unmatched:
            pool_map[assignment] = unmatched
        else:
            pool_map.pop(assignment, None)
        if auto_aliases:
            st.session_state["work_aliases"].setdefault(
                class_name, {}).update(auto_aliases)
            if summary is not None:
                for ck, rk in sorted(auto_aliases.items()):
                    summary["messages"].append(
                        f"🔗 {assignment}: matched `{ck}` → {rk} by prefix "
                        f"(recorded as an alias).")

    _ensure_class(class_name)
    if gb().assignments:
        gb().assignments[-1].class_name = class_name
        gb().assignments[-1].term = current_term()
        if prior_ref and not getattr(gb().assignments[-1], "folder_ref", ""):
            gb().assignments[-1].folder_ref = prior_ref
    st.session_state["active"].setdefault(assignment, True)
    st.session_state["archived"].discard(assignment)
    return len(created)


def assign_work(class_name: str, assignment: str, csv_key: str,
                roster_key: str) -> bool:
    """Resolve one pooled unmatched work to a roster student (Phase 4's action).

    Writes the durable alias ``csv_key → roster_key`` — so every later re-sync of
    this assignment routes the work the same way, surviving Sync's purge-replace
    — then immediately re-materializes the pooled row's score(s) under
    ``roster_key`` through the SAME engine path a routed ingest uses
    (:meth:`IngestionPipeline.materialize_row`), removes the row from the pool,
    and persists. Returns ``True`` on success, ``False`` when the pooled row is
    gone (already assigned, or the pool was rebuilt by a fresh sync).

    The grade lives in the CSV/pool, not the source file, so a work whose file
    was deleted after export (a sanctioned workflow) still assigns."""
    pool_map = st.session_state["unmatched_works"].get(class_name, {})
    rows = pool_map.get(assignment, [])
    row = next((r for r in rows if r.get("csv_key") == csv_key), None)
    if row is None:
        return False
    # 1. Durable alias first — recorded even if anything below is retried, so a
    #    re-sync always routes this csv_key to the roster student.
    st.session_state["work_aliases"].setdefault(
        class_name, {})[csv_key] = roster_key
    # 2. Re-materialize the pooled score(s) under the roster student now.
    IngestionPipeline(gb()).materialize_row(assignment, roster_key, row)
    # 3. Drop it from the pool; prune an emptied assignment / class.
    pool_map[assignment] = [r for r in rows if r.get("csv_key") != csv_key]
    if not pool_map[assignment]:
        pool_map.pop(assignment, None)
    if not pool_map:
        st.session_state["unmatched_works"].pop(class_name, None)
    persist()
    return True


def rename_assignment(old: str, new: str) -> bool:
    """Rename an assignment everywhere it is referenced (record + scores + state)."""
    new = new.strip()
    if not new or new == old:
        return False
    if new in {a.name for a in gb().assignments}:
        st.session_state["save_status"] = ("error", f"'{new}' already exists.")
        return False
    for a in gb().assignments:
        if a.name == old:
            a.name = new
    for _, sc in all_scores():
        if sc.assignment == old:
            sc.assignment = new
            if sc.source.startswith("csv:"):
                sc.source = f"csv:{new}"
            if sc.source == f"exam:{old}":
                sc.source = f"exam:{new}"
    for student in gb():
        results = getattr(student, "exam_results", {})
        if old in results:
            r = results.pop(old)
            r.assignment = new
            results[new] = r
    if old in st.session_state["date_override"]:
        st.session_state["date_override"][new] = \
            st.session_state["date_override"].pop(old)
    for m in st.session_state["active_by_term"].values():
        if old in m:
            m[new] = m.pop(old)
    if old in st.session_state["archived"]:
        st.session_state["archived"].discard(old)
        st.session_state["archived"].add(new)
    lf = st.session_state["late_flags"]
    for k in list(lf.keys()):
        parts = k.split("||")
        if len(parts) == 3 and parts[1] == old:
            lf["||".join([parts[0], new, parts[2]])] = lf.pop(k)
    ef = st.session_state["excused_flags"]
    for k in list(ef.keys()):
        parts = k.split("||")
        if len(parts) == 2 and parts[1] == old:
            ef["||".join([parts[0], new])] = ef.pop(k)
    if st.session_state["sel_assignment"] == old:
        st.session_state["sel_assignment"] = new
    persist()
    return True


def archive_assignment(name: str) -> None:
    """Soft-delete: hide from the table and exclude from math, but keep data."""
    st.session_state["archived"].add(name)
    for term in TERMS:
        term_active_map(term)[name] = False
    persist()


def restore_assignment(name: str) -> None:
    st.session_state["archived"].discard(name)
    st.session_state["active"][name] = True
    persist()


def delete_assignment_permanent(name: str) -> None:
    """Hard-delete: purge the assignment record AND every score it produced."""
    # Remember which classes held this name before the records go, so the
    # cloud-sync registry rows below can be scoped to them.
    classes = {getattr(a, "class_name", "") for a in gb().assignments
               if a.name == name}
    gb().assignments = [a for a in gb().assignments if a.name != name]
    for student in gb():
        for crit, bucket in list(student.scores.items()):
            student.scores[crit] = [sc for sc in bucket if sc.assignment != name]
        getattr(student, "exam_results", {}).pop(name, None)
    st.session_state["archived"].discard(name)
    for m in st.session_state["active_by_term"].values():
        m.pop(name, None)
    st.session_state["date_override"].pop(name, None)
    lf = st.session_state["late_flags"]
    for k in list(lf.keys()):
        parts = k.split("||")
        if len(parts) == 3 and parts[1] == name:
            lf.pop(k)
    ef = st.session_state["excused_flags"]
    for k in list(ef.keys()):
        parts = k.split("||")
        if len(parts) == 2 and parts[1] == name:
            ef.pop(k)
    if st.session_state["sel_assignment"] == name:
        st.session_state["sel_assignment"] = None
    # Forget the cloud-sync registry rows whose source file feeds this
    # assignment, so a later re-export/re-sync re-ingests it instead of skipping
    # it as "unchanged" — otherwise the file's grades sit stranded on disk with
    # no row to land on (deleting the row alone doesn't invalidate the hash the
    # importer dedups on). A '/'→'_' mangled filename is matched via the same
    # sanitize→clean round-trip the importer uses, so deleting "Maquette / Mock
    # Up" also releases its "Maquette _ Mock Up_Grades…csv".
    reg = st.session_state["ingested_files"]
    key = clean_assignment_name(_safe_dirname(name))
    for path in [p for p, rec in reg.items()
                 if rec.get("assignment") in (name, key)
                 and (not classes or rec.get("class") in classes)]:
        reg.pop(path, None)
    persist()


def delete_class(name: str) -> bool:
    """Permanently delete one class/level and everything scoped to it.

    Removes the class's assignments, the scores those assignments produced
    (restricted to the class's roster when one is known, so a unit name shared
    with another class is never touched there), its roster, its unit plan, and
    the class entry itself. Every other class is left intact. The active class
    is moved to a surviving class afterwards (or a fresh default is recreated
    via :func:`ensure_class_context` when the last class is removed). Returns
    True on success, False if the name isn't a real class.
    """
    ss = st.session_state
    if name not in class_names():
        ss["save_status"] = ("error", f"Class '{name}' not found.")
        return False

    names = {a.name for a in gb().assignments
             if getattr(a, "class_name", "") == name}
    roster_keys = {e["key"] for e in ss["rosters"].get(name, [])}

    # Drop this class's scores. Scope to its roster when known so a unit name
    # shared with another class isn't wiped from those students.
    for student in gb():
        if roster_keys and student.student_id not in roster_keys:
            continue
        for crit, bucket in list(student.scores.items()):
            student.scores[crit] = [sc for sc in bucket if sc.assignment not in names]
        results = getattr(student, "exam_results", {})
        for nm in [n for n in results if n in names]:
            results.pop(nm, None)

    # Drop this class's assignment records.
    gb().assignments = [a for a in gb().assignments
                        if getattr(a, "class_name", "") != name]

    # Forget per-assignment UI state for the removed units.
    for nm in names:
        for m in ss["active_by_term"].values():
            m.pop(nm, None)
        ss["date_override"].pop(nm, None)
        ss["archived"].discard(nm)
    lf = ss["late_flags"]
    for k in list(lf.keys()):
        parts = k.split("||")
        if len(parts) == 3 and parts[1] in names:
            lf.pop(k)
    ef = ss["excused_flags"]
    for k in list(ef.keys()):
        parts = k.split("||")
        if len(parts) == 2 and parts[1] in names:
            ef.pop(k)

    # Forget cloud-sync registry rows for this class so a future re-sync of the
    # same files re-ingests cleanly rather than being skipped as "unchanged".
    reg = ss["ingested_files"]
    for path in [p for p, rec in reg.items() if rec.get("class") == name]:
        reg.pop(path, None)

    # Remove the class container itself (roster, unit plan, class entry).
    ss["rosters"].pop(name, None)
    ss["unit_plans"].pop(name, None)
    ss["classes"] = [c for c in ss["classes"] if c["name"] != name]

    # If the deleted class was active, fall back to a surviving one.
    if ss["active_class"] == name:
        ss["active_class"] = class_names()[0] if class_names() else ""
        ss["focus_sid"] = None
        ss["sel_assignment"] = None
        ss["edit_cell"] = None
    ensure_class_context()  # recreate a default class if none remain
    # Deleting a class deliberately drops its students' teacher input from the
    # session: flag it so the mirror shrink tripwire treats the reduction as
    # intended rather than as catastrophic mass loss (invariant 2).
    _mark_teacher_input_deleted()
    persist()
    return True


def wipe_database_full() -> None:
    """Clear ALL teaching data across every class (device/window prefs kept)."""
    ss = st.session_state
    ss["gradebook"] = Gradebook()
    ss["classes"] = []
    ss["active_class"] = ""
    ss["rosters"] = {}
    ss["unit_plans"] = {}
    ss["roster"] = []
    ss["unit_plan"] = None
    for k in ("late_flags", "excused_flags", "final_override",
              "teacher_remarks", "date_override"):
        ss[k] = {}
    ss["active_by_term"] = {}
    ss["comments_by_term"] = {}
    ss["effort_by_term"] = {}
    ensure_term_context()   # re-point the active/llm_response aliases
    ss["archived"] = set()
    ss["ingested_sigs"] = set()
    ss["ingested_files"] = {}  # forget the cloud-sync registry -> full re-sync
    ss["work_aliases"] = {}    # Phase 3: durable alias map + unmatched pools
    ss["unmatched_works"] = {}
    ss["staging"] = {}
    ss["focus_sid"] = None
    ss["sel_assignment"] = None
    ss["edit_cell"] = None
    ensure_class_context()  # recreate the default empty class
    # A checkbox-confirmed wipe is a deliberate mass deletion: tell the cloud
    # mirror so its shrink tripwire doesn't fight the reduction (the old classes'
    # files aren't rewritten anyway — those classes are gone from the session).
    _mark_teacher_input_deleted()
    # Deliberate, checkbox-confirmed wipe: bypass the Phase-3 shrink tripwire
    # (that is exactly the mass loss it exists to catch). The daily .bak-auto
    # snapshot inside persist() still captures the pre-wipe DB.
    persist(allow_shrink=True)


# --------------------------------------------------------------------------
# Exports
# --------------------------------------------------------------------------

def students_for_active_class() -> list:
    """Students belonging to the currently selected Class/Level.

    A student counts as part of the active class if they are on its roster, or
    — when no roster has been loaded yet — if they have at least one score in an
    assignment tagged to this class. Ordering follows the roster first, then any
    remaining score-only students by ID. Exports use this so they never dump the
    whole gradebook (every class) into one file.

    Archived students are excluded: archiving removes them from the roster but
    deliberately keeps their grades, so without this they would re-enter through
    the score-only path and leak into every export (and the sync / comment
    passes) as a departed student with no email."""
    cls = st.session_state.get("active_class", "")
    roster_keys = [e["key"] for e in st.session_state["rosters"].get(cls, [])]
    archived_keys = {e.get("key") for e
                     in st.session_state.get("archived_students", {}).get(cls, [])}
    class_assignments = {a.name for a in gb().assignments
                         if getattr(a, "class_name", "") == cls}
    score_keys = {s.student_id for s, sc in all_scores()
                  if sc.assignment in class_assignments}
    ordered, seen = [], set()
    for key in roster_keys:
        student = gb().students.get(key)
        if student is not None and key not in seen:
            ordered.append(student)
            seen.add(key)
    for sid in sorted(score_keys):
        if sid not in seen and sid not in archived_keys:
            student = gb().students.get(sid)
            if student is not None:
                ordered.append(student)
                seen.add(sid)
    return ordered


def _classroom_folder_assignment_names() -> list:
    """Names of the active class's Google-Classroom *folder* assignments, in
    timeline (date) order.

    A folder assignment is one CAM ingested FROM Google Classroom: it carries a
    grading-workspace CSV (``source_file``) or a watched-folder pin
    (``folder_ref``), it is not an exam, and it produced at least one criterion
    score. Crit D ``(Reflection)`` tasks are deliberately excluded — they are
    graded separately, hold no per-student artwork comment, and are not the
    folder submissions the teacher keys back into Classroom."""
    active = st.session_state.get("active_class", "")
    by_name = {a.name: a for a in gb().assignments
               if getattr(a, "class_name", "") == active}
    names = []
    for row in assignment_table():        # date-sorted, active class, unarchived
        asg = by_name.get(row["name"])
        if asg is None or row["is_exam"] or row["is_formative"]:
            continue
        if not (row["folder_ref"] or getattr(asg, "source_file", "")):
            continue                      # not from a Classroom folder / CGW CSV
        if row["name"].strip().lower().endswith("(reflection)"):
            continue                      # reflections graded separately
        names.append(row["name"])
    return names


def _append_classroom_entry_sheet(wb, students) -> None:
    """Add the 'Classroom Entry' tab: one row per student with a paired
    Mark/Comment column per folder assignment, so the teacher copies each
    column straight back into Google Classroom without re-matching students.

    Rows are ordered in **Latin name order (first name, then surname)** to match
    how Google Classroom lists students — deliberately NOT the gojūon order the
    on-screen roster uses, so a pasted column lines up row-for-row. This is the
    one place we re-sort away from roster order; it is purely cosmetic, because
    the mark/comment lookup matches on ``student_id`` (below), never position."""
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    def _latin_key(student):
        first = first_name_for(student)
        name = getattr(student, "name", "") or ""
        if first and name.lower().endswith(first.lower()) \
                and len(name) > len(first):
            last = name[:len(name) - len(first)].strip()
        else:
            parts = name.split()
            last = " ".join(parts[:-1]) if len(parts) > 1 else name
        return (first.casefold(), last.casefold())
    students = sorted(students, key=_latin_key)

    names = _classroom_folder_assignment_names()

    # (student_id, assignment) -> (band, comment, is_valid); prefer a valid score
    # (a folder assignment is single-criterion, so this is one row per student).
    marks: dict = {}
    wanted = set(names)
    for student, sc in all_scores():
        if sc.assignment not in wanted:
            continue
        key = (student.student_id, sc.assignment)
        prev = marks.get(key)
        if prev is None or (not prev[2] and sc.is_valid):
            marks[key] = (sc.value, sc.comment or "", sc.is_valid)

    ws = wb.create_sheet("Classroom Entry")
    FONT = "Arial"
    hdr_fill = PatternFill("solid", fgColor="1F4E78")
    sub_fill = PatternFill("solid", fgColor="D9E1F2")
    name_fill = PatternFill("solid", fgColor="F2F2F2")
    white_bold = Font(name=FONT, bold=True, color="FFFFFF")
    blue_bold = Font(name=FONT, bold=True, color="1F4E78", size=10)
    base = Font(name=FONT, size=10)
    thin = Side(style="thin", color="BFBFBF")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    center = Alignment(horizontal="center", vertical="center")
    left = Alignment(horizontal="left", vertical="center")
    wrap = Alignment(horizontal="left", vertical="top", wrap_text=True)

    fixed = ["Name", "Student ID"]
    nfix = len(fixed)
    for i, h in enumerate(fixed, start=1):
        ws.merge_cells(start_row=1, end_row=2, start_column=i, end_column=i)
        c = ws.cell(row=1, column=i, value=h)
        c.font, c.fill, c.alignment, c.border = white_bold, hdr_fill, center, border
        low = ws.cell(row=2, column=i)
        low.fill, low.border = hdr_fill, border

    col = nfix + 1
    for name in names:
        ws.merge_cells(start_row=1, start_column=col, end_row=1, end_column=col + 1)
        c = ws.cell(row=1, column=col, value=name)
        c.font, c.fill, c.alignment, c.border = white_bold, hdr_fill, center, border
        pad = ws.cell(row=1, column=col + 1)
        pad.fill, pad.border = hdr_fill, border
        for off, label in ((0, "Mark"), (1, "Comment")):
            sc = ws.cell(row=2, column=col + off, value=label)
            sc.font, sc.fill, sc.alignment, sc.border = blue_bold, sub_fill, center, border
        col += 2

    for ri, student in enumerate(students, start=3):
        nm = ws.cell(row=ri, column=1, value=student_label(student))
        sid = ws.cell(row=ri, column=2, value=student.student_id)
        for cc in (nm, sid):
            cc.font, cc.fill, cc.alignment, cc.border = base, name_fill, left, border
        col = nfix + 1
        for name in names:
            band, comment, _ = marks.get(
                (student.student_id, name), (None, "", False))
            m = ws.cell(row=ri, column=col, value=band)
            m.font, m.alignment, m.border = base, center, border
            cm = ws.cell(row=ri, column=col + 1, value=comment or None)
            cm.font, cm.alignment, cm.border = base, wrap, border
            col += 2

    ws.column_dimensions["A"].width = 22
    ws.column_dimensions["B"].width = 11
    col = nfix + 1
    for _ in names:
        ws.column_dimensions[get_column_letter(col)].width = 6
        ws.column_dimensions[get_column_letter(col + 1)].width = 46
        col += 2
    ws.freeze_panes = "C3"
    ws.sheet_view.showGridLines = False


def build_excel_bytes() -> bytes:
    from openpyxl import Workbook

    wb = Workbook()

    # Export is scoped to the currently selected Class/Level only.
    students = students_for_active_class()
    class_ids = {s.student_id for s in students}

    # Tab 1: Suggested final scores (student x criterion).
    ws = wb.active
    ws.title = "Final Suggestions"
    ws.append(["Student ID", "Name"] + [f"Crit {c}" for c in CRIT_ORDER]
              + ["Effort", "MYP Grade", "School Grade"])
    for student in students:
        row = [student.student_id, student.name]
        for c in CRIT_ORDER:
            res = aggregate_with_policy(student, c)
            row.append(f"{res.rounded_band}*" if res else "")
        _n, _total, effort, myp, gyo = student_term_grades(student)
        row += [effort,
                myp if myp is not None else "N/A",
                gyo if gyo is not None else "N/A"]
        ws.append(row)

    # Tab 2: Raw scores ledger (active class students only).
    ws2 = wb.create_sheet("Raw Scores")
    ws2.append(["Student ID", "Assignment", "Criterion", "Grade",
                "Date", "Valid", "In report", "Late", "Comment"])
    for student, sc in all_scores():
        if student.student_id not in class_ids:
            continue
        ws2.append([
            student.student_id, sc.assignment, sc.criterion.value, sc.value,
            sc.timestamp.strftime("%Y-%m-%d"), sc.is_valid, sc.include_in_report,
            is_late(student.student_id, sc.assignment, sc.criterion.value, sc),
            sc.comment,
        ])

    # Tab 3: Assignment analytics, headed by the class/subject/term metadata
    # so the exported workbook is self-describing.
    ws3 = wb.create_sheet("Assignments")
    ac = active_class_dict()
    ws3.append(["Name", ac.get("name", "")])
    ws3.append(["Class", class_label(ac)])
    ws3.append(["Subject", active_subject()])
    ws3.append(["Term", current_term()])
    ws3.append([])
    ws3.append(["Assignment", "Criteria", "Date", "Submissions",
                "Avg grade", "Spread", "Lates"] + [f"In {t}" for t in TERMS])
    for r in assignment_table():
        ws3.append([
            r["name"], r["criteria"], r["date"].strftime("%Y-%m-%d"),
            r["submissions"], r["avg"], r["spread"], r["lates"],
        ] + [assignment_on(r["name"], t) for t in TERMS])

    # Tab 4: Classroom Entry — grades + comments per student for each folder
    # (Google Classroom) assignment. This tab re-sorts to Latin name order
    # (first name, surname) to match Classroom's own student order, so the marks
    # can be copy-pasted straight back into Classroom without re-matching by
    # email/ID (the on-screen roster stays in gojūon order).
    _append_classroom_entry_sheet(wb, students)

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _trend_png(student):
    """Render the progression trend as a PNG via matplotlib.

    Replaces the old Plotly ``fig.to_image`` path: that needed kaleido, which
    (in its Chromium-based 1.x releases) hangs for minutes and crops the
    bottom of the chart. Matplotlib's Agg backend renders in milliseconds
    with no browser engine. Mirrors the cockpit's trend exactly (Missing = 0
    plotted, Excused removed). Returns PNG bytes, or None when there is no
    scored data (or matplotlib isn't installed)."""
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
    except ImportError:
        return None

    excused = excused_assignments_for(student.student_id)
    zero_by_crit = {c: missing_zero_points(student, c) for c in CRIT_ORDER}
    label_order = sorted(
        {sc.timestamp
         for c in CRIT_ORDER
         for sc in student.criterion_scores(Criterion(c), valid_only=True)
         if sc.include_in_report and sc.assignment not in excused}
        | {ts for pts in zero_by_crit.values() for ts, _ in pts},
        key=academic_sort_key,
    )
    categories = []
    for ts in label_order:            # de-dupe same-day labels, keep order
        lbl = f"{ts:%b %d}"
        if lbl not in categories:
            categories.append(lbl)
    if not categories:
        return None
    x_of = {lbl: i for i, lbl in enumerate(categories)}

    fig, ax = plt.subplots(figsize=(8.4, 3.4), dpi=150)
    has = False
    for c in CRIT_ORDER:
        pts = [(sc.timestamp, sc.value)
               for sc in student.criterion_scores(Criterion(c), valid_only=True)
               if sc.include_in_report and sc.assignment not in excused]
        pts += zero_by_crit[c]
        pts.sort(key=lambda p: academic_sort_key(p[0]))
        if not pts:
            continue
        has = True
        ax.plot([x_of[f"{p[0]:%b %d}"] for p in pts], [p[1] for p in pts],
                marker="o", markersize=4, linewidth=1.6,
                color=CRIT_COLORS[c], label=f"{c} {Criterion(c).label}")
    if not has:
        plt.close(fig)
        return None
    ax.set_xticks(range(len(categories)))
    ax.set_xticklabels(categories, fontsize=8)
    ax.set_ylim(-0.3, 8.4)
    ax.set_yticks(range(0, 9))
    ax.set_ylabel("Grade (0-8)", fontsize=9)
    ax.tick_params(axis="y", labelsize=8)
    ax.grid(True, alpha=0.3)
    ax.legend(fontsize=8, loc="lower left", framealpha=0.6)
    fig.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png")
    plt.close(fig)
    return buf.getvalue()


def _new_report_document():
    """A blank Word document with the app-wide page setup: A4 paper (21 x
    29.7 cm) and 2 cm margins on every side. Every export the app builds goes
    through here so the layout stays consistent."""
    from docx import Document
    from docx.shared import Cm

    document = Document()
    for section in document.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
    return document


def _student_docx(document, student, method, include_effort_school=True):
    """One student's full report page: class/subject/term header, individual
    marks, the progression graph, final criterion grades, the term's overall
    comment, and a generation timestamp.

    Which report-card grade rows appear at all is set in ⚙ Settings →
    Report-card grades (Effort / MYP Grade / School Grade — all off by default).
    On top of that, ``include_effort_school`` gates the Effort/English Use and
    School Grade rows: the mail-merge pack passes ``False`` — those reports go to
    students before their official report cards, so those two figures are
    withheld until then (the MYP Grade, which students do see, stays)."""
    document.add_heading(f"Report Card - {student_label(student)}", level=1)
    email = student_email_for(student)
    if email:
        document.add_paragraph(email)
    ac = active_class_dict()
    head = document.add_paragraph()
    head.add_run(f"Class: {ac.get('name', '')}     "
                 f"Subject: {subject_label()}     "
                 f"Term: {current_term()}").bold = True
    plan = st.session_state["unit_plan"]
    if plan:
        document.add_paragraph(f"Unit: {plan.unit_title}")
        document.add_paragraph(f"Statement of Inquiry: {plan.statement_of_inquiry}")

    # ---- Individual marks (selected tasks, incl. Missing = 0 / Excused) ----
    # Student-facing wording: the 0-8 mark is labelled "Grade", never "band".
    rows = assignment_table()
    active_names = [r["name"] for r in rows if assignment_on(r["name"])]
    rows_by_name = {r["name"]: r for r in rows}
    marks = document.add_table(rows=1, cols=5)
    marks.style = "Table Grid"
    for cell, cap in zip(marks.rows[0].cells,
                         ("Assignment", "Criterion", "Grade", "Date", "Status")):
        cell.text = cap
    for asg in active_names:
        r = rows_by_name.get(asg, {})
        excused_now = is_excused(student.student_id, asg)
        scs = [sc for b in student.scores.values() for sc in b
               if sc.assignment == asg]
        if scs:
            for sc in scs:
                crit = sc.criterion.value
                status = ("excused" if excused_now else
                          "late" if is_late(student.student_id, asg, crit, sc)
                          else "")
                cells = marks.add_row().cells
                cells[0].text = asg
                cells[1].text = crit
                cells[2].text = str(sc.value) if sc.is_valid else "·"
                cells[3].text = f"{sc.timestamp:%b %d}"
                cells[4].text = status
        elif r and awaiting_grade(r):
            # Awaiting Grade — folder-backed and still being graded: no invented
            # 0, mirrored from Window 3 (and excluded from the math the same
            # way). Once the folder's grading completes this row falls through to
            # the missing-0 branch below instead.
            cells = marks.add_row().cells
            cells[0].text = asg
            cells[1].text = r.get("criteria", "—")
            cells[2].text = "—"
            cells[3].text = f"{r['date']:%b %d}"
            cells[4].text = "awaiting grade"
        elif r and r.get("criteria", "—") != "—":
            for crit in [c.strip() for c in r["criteria"].split(",")
                         if c.strip() in CRIT_ORDER]:
                cells = marks.add_row().cells
                cells[0].text = asg
                cells[1].text = crit
                cells[2].text = "—" if excused_now else "0"
                cells[3].text = f"{r['date']:%b %d}"
                cells[4].text = "excused" if excused_now else "missing = 0"

    # ---- Progression trend graph (matplotlib — mirrors the cockpit) ----
    png = _trend_png(student)
    if png:
        from docx.shared import Inches
        document.add_heading("Progression trend", level=2)
        document.add_picture(io.BytesIO(png), width=Inches(6.3))

    # ---- Final criterion grades (Final only — no working columns) ----
    document.add_heading("Final criterion grades", level=2)
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text = "Criterion", "Final"
    final = st.session_state["final_override"].get(student.student_id, {})
    for c in CRIT_ORDER:
        cells = table.add_row().cells
        cells[0].text = f"{c} - {Criterion(c).label}"
        if c in final:
            cells[1].text = str(final[c])
        else:
            res = aggregate_with_policy(student, c)
            cells[1].text = str(res.rounded_band) if res else "-"

    # ---- School report-card grades (shared helper — matches Window 3) ----
    # Only the figures enabled in ⚙ Settings → Report-card grades appear here
    # (all off by default). The CAM master Excel export keeps all three
    # regardless — this gating is for the student-facing reports only.
    rc = report_cfg()
    _n, _total, effort, myp, gyo = student_term_grades(student)
    grade_rows = []
    if rc.get("show_effort", False):
        grade_rows.append(("Effort / English Use", str(effort)))
    if rc.get("show_myp_grade", False):
        grade_rows.append(("MYP Grade", str(myp) if myp is not None else "N/A"))
    if rc.get("show_school_grade", False):
        grade_rows.append(
            ("School Grade", str(gyo) if gyo is not None else "N/A"))
    if not include_effort_school:
        # Withheld until report cards are issued; School goes too because it
        # folds effort in and would otherwise leak it. MYP Grade stays (only if
        # it is enabled at all).
        grade_rows = [r for r in grade_rows if r[0] == "MYP Grade"]
    for label, value in grade_rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value

    # ---- Term overall comment ----
    resp = st.session_state["llm_response"].get(student.student_id, "")
    if resp:
        document.add_heading(f"{current_term()} — overall comment", level=2)
        document.add_paragraph(resp)

    # ---- Generation timestamp ----
    stamp = document.add_paragraph()
    run = stamp.add_run(f"Report generated: {datetime.now():%Y-%m-%d %H:%M}")
    run.italic = True


def build_reportcards_docx(students) -> bytes:
    """The report-card pack: every student's individual report (marks, graph,
    final grades, comments) as their own page(s) in one document."""
    document = _new_report_document()
    for i, student in enumerate(students):
        # Resolve the method per student, inside the loop — two students in the
        # same class may now legitimately compute under different methods.
        _student_docx(document, student, calculation_method(student.student_id))
        if i < len(students) - 1:
            document.add_page_break()
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def build_reportcards_zip(students):
    """Mail-merge pack: each student's report as its OWN ``.docx`` (named
    ``<email>.docx``), all bundled into one ZIP.

    Distinct from ``build_reportcards_docx`` (which is ONE combined document):
    here the filename IS the student's school email, so a batch-send script can
    mail each file straight back to the address in its own name with no roster
    sheet to keep in sync. Conversion to PDF is left to send time, so the files
    stay ``.docx`` in the ZIP.

    Returns ``(zip_bytes, skipped)`` where ``skipped`` is a list of
    ``(student_label, reason)`` for students with no email, an email that would
    be illegal as a filename, or a duplicate email. The email is never
    sanitized — altering it would change who the file mails itself to — so such
    students are left out and surfaced to the UI rather than silently mis-sent.
    """
    import re
    import zipfile

    buf = io.BytesIO()
    skipped = []
    seen = set()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for student in students:
            label = student_label(student)
            email = student_email_for(student).strip()
            if not email:
                skipped.append((label, "no email on roster"))
                continue
            if re.search(r'[\\/:*?"<>|]', email):
                skipped.append((label, f"email unusable as filename: {email}"))
                continue
            if email.lower() in seen:
                skipped.append((label, f"duplicate email: {email}"))
                continue
            seen.add(email.lower())
            document = _new_report_document()
            _student_docx(document, student, calculation_method(student.student_id),
                          include_effort_school=False)
            doc_buf = io.BytesIO()
            document.save(doc_buf)
            zf.writestr(f"{email}.docx", doc_buf.getvalue())
    return buf.getvalue(), skipped


def build_class_comments_docx(students) -> bytes:
    """One document compiling every student's saved comments for the whole
    class — each term's overall comment plus any teacher remarks."""
    document = _new_report_document()
    cls = st.session_state["active_class"]
    document.add_heading(f"Class comments — {cls}", level=0)
    head = document.add_paragraph()
    head.add_run(f"Class: {cls}     Subject: {subject_label()}     "
                 f"Term: {current_term()}").bold = True
    by_term = st.session_state["comments_by_term"]
    for student in students:
        document.add_heading(student_label(student), level=1)
        wrote = False
        for term in TERMS:
            text = ((by_term.get(term, {}) or {})
                    .get(student.student_id, "") or "").strip()
            if not text:
                continue
            document.add_heading(term, level=2)
            document.add_paragraph(text)
            wrote = True
        remarks = (st.session_state["teacher_remarks"]
                   .get(student.student_id, "") or "").strip()
        if remarks:
            document.add_heading("Teacher remarks", level=2)
            document.add_paragraph(remarks)
            wrote = True
        if not wrote:
            document.add_paragraph("(no comments recorded)")
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


# --------------------------------------------------------------------------
# Prompt + clipboard
# --------------------------------------------------------------------------

def _current_term_evidence(student) -> list:
    """Build :class:`Evidence` items from the student's valid, included scores in
    the CURRENT term's active assignments — the raw material for slicing.
    Excused assignments are left out of the prompt entirely."""
    names = current_term_assignment_names() - excused_assignments_for(student.student_id)
    items = []
    for bucket in student.scores.values():
        for sc in bucket:
            if (sc.assignment in names and sc.is_valid and sc.include_in_report):
                items.append(Evidence(
                    assignment=sc.assignment,
                    criterion=sc.criterion.value,
                    score=int(sc.value),
                    timestamp=sc.timestamp,
                    keywords=list(sc.keywords),
                    comment=clean_comment(sc.comment),
                ))
    return items


def _matched_task_lines(student) -> list:
    """Current-term assignments the student has marks in, with their criteria.
    Excused assignments are omitted (they are out of the assessment)."""
    names = current_term_assignment_names() - excused_assignments_for(student.student_id)
    table = {r["name"]: r for r in assignment_table()}
    have = sorted({sc.assignment
                   for bucket in student.scores.values()
                   for sc in bucket if sc.assignment in names})
    out = []
    for nm in have:
        crit = table.get(nm, {}).get("criteria", "—")
        out.append(f"{nm} (criteria {crit})")
    return out


def _trend_lines(student, detail: str = "compact") -> list:
    """One trajectory sentence per criterion with ≥2 current-term data points.
    Follows the Missing=0 / Excused policy so the narrated trend matches the
    plotted one. ``detail`` ("compact"|"detailed") controls how much of the
    path each sentence narrates (see ``format_trend_sentence``)."""
    names = current_term_assignment_names() - excused_assignments_for(student.student_id)
    lines = []
    for c in CRIT_ORDER:
        series = [(sc.timestamp, sc.value)
                  for sc in student.criterion_scores(Criterion(c), valid_only=True)
                  if sc.include_in_report and sc.assignment in names]
        series += missing_zero_points(student, c, names)
        info = trend_for_series(series)
        if info is None:
            continue
        ordered = sorted(series, key=lambda p: p[0])
        lines.append(format_trend_sentence(
            f"Criterion {c} ({Criterion(c).label})", info,
            first_label=f"{ordered[0][0]:%b %d}",
            last_label=f"{ordered[-1][0]:%b %d}",
            detail=detail))
    return lines


def _late_submission_stats(student, names) -> tuple:
    """``(late_assignments, submitted_assignments)`` for the timeliness block.

    *Submitted* = a current-term, non-excused assignment on which the student
    holds at least one valid, report-included score. *Late* = such an assignment
    where **any** criterion score reads late via the two-layer ``is_late`` (a
    manual override wins, else the synced ``late`` field). Counted per distinct
    assignment — how many criteria it touched is irrelevant."""
    excused = excused_assignments_for(student.student_id)
    submitted = 0
    late = 0
    for asg in names:
        if asg in excused:
            continue
        scs = [sc for bucket in student.scores.values() for sc in bucket
               if sc.assignment == asg and sc.is_valid and sc.include_in_report]
        if not scs:
            continue
        submitted += 1
        if any(is_late(student.student_id, asg, sc.criterion.value, sc)
               for sc in scs):
            late += 1
    return late, submitted


def _missing_work_stats(student, names) -> tuple:
    """``(missing_rows, assessed_n)`` for the missing-work block.

    *Missing* reuses ``missing_assignment_rows`` verbatim — the same
    single-predicate gate Window 3 uses, so Excused work, ⏳ Awaiting Grade
    folder assignments, formative events (criteria "—") and unbanded exams are
    already excluded. *Assessed* = missing + submitted, where *submitted* is
    exactly ``_late_submission_stats``'s definition (a non-excused assignment
    carrying ≥ 1 valid, report-included score). Counted per distinct assignment,
    never per criterion — matching the late block. The denominator deliberately
    differs from the late block's *submitted* one: a never-submitted task can be
    missing but can never be late."""
    missing_rows = missing_assignment_rows(student, names)
    submitted = _late_submission_stats(student, names)[1]
    return missing_rows, len(missing_rows) + submitted


def compile_prompt(student, cfg) -> str:
    """Synthesize the segmented system prompt for one student's comment.

    The blueprint cleanly segregates curated context into labeled blocks so the
    model can tell durable curriculum framing, the current term's sliced
    evidence, the math-engine trajectory, and compressed prior-term history
    apart. Token budget is controlled by (a) sending only the best/worst sliced
    pieces rather than every mark, and (b) replacing past terms' raw evidence
    with their one finalized summary each."""
    plan = st.session_state["unit_plan"]
    name_mode = cfg.get("name_mode", "first")
    pron = pronouns_for(student)
    scope = cfg.get("focus_scope", "current")
    cur_term = current_term()
    names = current_term_assignment_names()

    blocks = []

    # ---- 1. Role / instructions ----
    subject = active_subject()
    subj_role = f"IB MYP {subject} teacher" if subject else "IB MYP teacher"
    blocks.append(
        "[SYSTEM ROLE]\n"
        f"You are an experienced {subj_role} writing a single, "
        "concise, cohesive report-card comment for ONE student, covering the "
        f"CURRENT term ({cur_term}). Write in flowing prose, not headings or "
        "bullet lists.")

    # ---- 2. Student & level ----
    student_lines = []
    if name_mode == "full":
        student_lines.append(f"Full name: {student_label(student)} — use the "
                             "FULL name in the first sentence.")
    elif name_mode == "none":
        student_lines.append("Do NOT use the student's name anywhere — refer to "
                             "them by pronoun only.")
    else:
        student_lines.append(f"First name: {first_name_for(student)} — use the "
                             "FIRST name in the first sentence.")
    student_lines.append(f"Pronouns: {pron}.")
    ac = active_class_dict()
    if subject:
        student_lines.append(f"Subject: {subject}.")
    myp = active_myp_year()
    if myp:
        grade = ac.get("grade") or ac.get("name") or ""
        label = f"{grade} (MYP Year {myp})" if grade else f"MYP Year {myp}"
        student_lines.append(
            f"Academic level: {label}. Pitch vocabulary, conceptual complexity "
            f"and expectations for an MYP Year {myp} student.")
    blocks.append("[STUDENT & LEVEL]\n" + "\n".join(student_lines))

    # ---- 3. Curriculum context (always included) ----
    curric = []
    if plan:
        if plan.unit_title:
            curric.append(f"Unit: {plan.unit_title}")
        if plan.statement_of_inquiry:
            curric.append(f"Statement of Inquiry: {plan.statement_of_inquiry}")
        if getattr(plan, "key_concepts", None):
            curric.append(f"Core/Key Concepts: {plan.concepts_text}")
    matched = _matched_task_lines(student)
    if matched:
        curric.append("Graded tasks this term:\n  - " + "\n  - ".join(matched))
    # Unsubmitted-task names deliberately live in the toggleable [MISSING WORK]
    # block (section 4c), not here — so turning that toggle OFF removes every
    # missing-work signal from the prompt.
    if not curric:
        curric.append("(No unit plan loaded — rely on the criterion evidence below.)")
    blocks.append("[CURRICULUM CONTEXT]\n" + "\n".join(curric))

    # ---- 4. Current-term criterion results (term-scoped bands) ----
    # Missing selected work counts as 0 and excused work is dropped, matching
    # the cockpit's grade panel exactly.
    crit_lines = []
    for c in CRIT_ORDER:
        res = aggregate_with_policy(student, c, names=names)
        if res:
            crit_lines.append(f"  - {c} {Criterion(c).label}: grade "
                              f"{res.rounded_band} (from {res.n} scored piece(s))")
    if crit_lines:
        blocks.append("[CURRENT TERM — CRITERION RESULTS (MYP criterion grade 0-8)]\n"
                      + "\n".join(crit_lines))

    # ---- 4b. Submission timeliness (optional, toggleable) ----
    # Omitted entirely at 0 late so the model never over-indexes on a perfect
    # record; only surfaces when there is a genuine late share to acknowledge.
    late_block = False
    if cfg.get("inc_late", True):
        late_n, submitted_n = _late_submission_stats(student, names)
        if late_n > 0 and submitted_n > 0:
            pct = round(late_n / submitted_n * 100)
            blocks.append(
                "[SUBMISSION TIMELINESS]\n"
                f"  - {late_n} of the {submitted_n} graded tasks this term "
                f"({pct}%) were submitted late.")
            late_block = True

    # ---- 4c. Missing work (optional, toggleable) ----
    # Omitted entirely at 0 missing, mirroring the late block: nothing for the
    # model to over-index on. The unsubmitted task names moved here out of
    # [CURRICULUM CONTEXT], so OFF removes all missing-work signal. The count
    # comes solely from Window 3's structural indicators (missing_assignment_rows)
    # — stored band-0 scores are never scanned (a stored 0 means something WAS
    # submitted).
    missing_block = False
    if cfg.get("inc_missing", True):
        missing_rows, assessed_n = _missing_work_stats(student, names)
        if missing_rows and assessed_n > 0:
            missing_n = len(missing_rows)
            pct = round(missing_n / assessed_n * 100)
            names_txt = ", ".join(r["name"] for r in missing_rows)
            blocks.append(
                "[MISSING WORK]\n"
                f"  - {missing_n} of the {assessed_n} assessed tasks this term "
                f"({pct}%) were not submitted; each counts as grade 0 in the "
                "criterion results above.\n"
                f"  - Unsubmitted: {names_txt}")
            missing_block = True

    # ---- 5/6. Best-vs-worst sliced evidence (Strengths / Growths) ----
    evidence = _current_term_evidence(student)
    n_s = int(cfg.get("n_strengths", 2)) if cfg.get("inc_strengths") else 0
    n_g = int(cfg.get("n_growth", 2)) if cfg.get("inc_growth") else 0
    strengths, growths = select_evidence(evidence, n_s, n_g)
    if strengths:
        blocks.append(
            "[CURRENT TERM — STRENGTH EVIDENCE (highest-scoring pieces)]\n"
            "Draw the strengths from this material:\n  - "
            + "\n  - ".join(e.as_text() for e in strengths))
    if growths:
        blocks.append(
            "[CURRENT TERM — GROWTH EVIDENCE (lowest-scoring pieces)]\n"
            "Frame the growth areas from this material:\n  - "
            + "\n  - ".join(e.as_text() for e in growths))

    # ---- 7. Trend trajectory (optional, math-engine derived) ----
    if cfg.get("inc_trend", True):
        # Detail level auto-follows the word budget: a roomier comment (130+
        # words) gets the fuller path narration, a tight one stays compact.
        trend_detail = "detailed" if int(cfg.get("word_limit", 100)) >= 130 \
            else "compact"
        trend = _trend_lines(student, detail=trend_detail)
        if trend:
            blocks.append("[TREND SUMMARY (math engine)]\n  - "
                          + "\n  - ".join(trend))

    # ---- 8. Compressed prior-term history (only when scope includes past) ----
    if scope == "include_past":
        prior = past_term_context_for(student)
        for term, text in prior:
            blocks.append(f"[PREVIOUS TERM FINALIZED SUMMARY — {term}]\n{text}")
        # Soft alert baked into the prompt: name the terms with no context so
        # the model never invents a history for them.
        gaps = missing_past_terms(student)
        if gaps:
            blocks.append(
                "[PREVIOUS TERM CONTEXT — GAPS]\nNo saved comment exists for: "
                + ", ".join(gaps) + ". Do not speculate about those terms; if "
                "the student has no recorded work there, they may have joined "
                "the class later — simply omit that period.")

    # ---- 9. Teacher notes ----
    remarks = st.session_state["teacher_remarks"].get(student.student_id, "")
    if remarks.strip():
        blocks.append("[TEACHER NOTES]\n" + remarks.strip())

    # ---- 10. Output requirements ----
    req = []
    want = []
    if cfg.get("inc_strengths"):
        want.append(f"{cfg['n_strengths']} specific strength(s)")
    if cfg.get("inc_growth"):
        want.append(f"{cfg['n_growth']} growth area(s)")
    if cfg.get("inc_next"):
        want.append("clear next steps")
    if cfg.get("inc_criteria"):
        want.append("a short note per criterion")
    if want:
        req.append("Include: " + ", ".join(want) + ".")
    req.append(f"Keep it to about {cfg['word_limit']} words, warm and "
               "professional, using the student's pronouns.")
    if cfg.get("tone_formal"):
        req.append("Tone: more formal and academic — measured, precise language "
                   "for an official report card.")
    if cfg.get("tone_encouraging"):
        req.append("Tone: more encouraging — affirm effort and frame growth "
                   "areas as next opportunities.")
    if cfg.get("no_numbers"):
        req.append(
            "Do NOT state any numbers in the comment: no criterion grades, "
            "scores, marks, percentages or counts. Describe achievement "
            "qualitatively instead (e.g. 'excellent', 'strong', 'developing', "
            "'has room to improve'). The numeric data above is for your "
            "understanding only.")
    req.append("Base the narrative ONLY on the current term's curriculum, "
               "criterion, strength, growth and trend sections above.")
    if late_block:
        req.append("If a notable share of tasks was late, acknowledge "
                   "submission habits briefly within the next steps.")
    if missing_block:
        req.append("Missing work is already reflected in the grade results "
                   "above — acknowledge incomplete submission habits once, "
                   "briefly, without double-penalizing the tone.")
    if scope == "include_past":
        req.append("Use any [PREVIOUS TERM FINALIZED SUMMARY] section ONLY for "
                   "developmental trend context — to note growth or continuity "
                   "since earlier terms. Do NOT restate or quote it, and do not "
                   "re-grade past work.")
    req.append(f"(Target model: {cfg['provider']})")
    blocks.append("[OUTPUT REQUIREMENTS]\n" + "\n".join(req))

    return "\n\n".join(blocks)


def clipboard_button(text: str) -> None:
    # Rendered inside an st.iframe (raw-HTML mode), so the page-level theme CSS
    # cannot reach it — pick up the active theme's button surface directly.
    safe = json.dumps(text)
    mode = theme_mode()
    s = THEME_SURFACES[mode]
    fg = st.get_option(f"theme.{mode}.textColor")
    html = (
        f"<button id='acm_cpy' style='padding:6px 12px;border-radius:6px;"
        f"border:1px solid {s['btn_border']};background:{s['btn_bg']};"
        f"color:{fg};cursor:pointer'>"
        f"Copy prompt to clipboard</button> "
        f"<span id='acm_ok' style='color:{fg}'></span>"
        "<script>const t=" + safe + ";"
        "document.getElementById('acm_cpy').onclick=function(){"
        "navigator.clipboard.writeText(t);"
        "document.getElementById('acm_ok').innerText=' copied';};</script>"
    )
    # st.iframe treats a non-URL string as raw HTML and embeds it in an iframe
    # (scripts still run), replacing the deprecated st.components.v1.html.
    st.iframe(html, height=46)


def build_single_docx(student) -> bytes:
    document = _new_report_document()
    _student_docx(document, student, calculation_method(student.student_id))
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


# --------------------------------------------------------------------------
# Modal dialogs (flag-driven so widgets inside survive internal reruns)
# --------------------------------------------------------------------------

def _render_db_switch_panel(pending: dict) -> None:
    """Adopt-vs-overwrite decision for a newly-configured path (Phase 2).

    Save pointed ``db_custom_path`` at a location that **already holds** a
    readable ``acm_database.json``. Overwriting it with the current (possibly
    demo) session is wipe mechanism 1, so CAM never does it silently: this panel
    offers **Load** (the default — "point my new PC at my cloud DB") or an
    explicit, backed-up **Replace**. Rendered in place of the settings form
    while ``db_switch_pending`` is set; every button ends with ``st.rerun()``,
    which closes the dialog. The new path pref is committed only on Load /
    Replace — Cancel (and an ESC-dismiss) leave it on the old location, so the
    session is never left pointed at (and about to autosave over) this DB."""
    path = pending["path"]
    new_custom = pending.get("new_custom", "")
    counts = pending.get("counts", {})
    st.markdown("**A database already exists at this location**")
    st.info(
        f"`{path}`\n\n"
        f"already contains **{counts.get('assignments', 0)} assignment(s)**, "
        f"**{counts.get('students', 0)} student(s)** across "
        f"**{counts.get('classes', 0)} class(es)**.")
    st.caption("CAM will not overwrite it automatically. Choose what to do:")

    if st.button("📥 Load this database into CAM", type="primary",
                 width="stretch", key="dbsw_adopt",
                 help="Use the database already at this location — the normal "
                      "way to point a second computer at your cloud data. "
                      "Nothing on disk is changed."):
        # Commit the new path pref, then re-run the boot hydrate against it (the
        # same code a fresh boot uses). db_load_blocked is cleared so a prior
        # quarantine does not linger; the hydrate re-diagnoses the (ok) file.
        prefs = st.session_state["prefs"]
        prefs["db_custom_path"] = new_custom
        save_prefs(prefs)
        st.session_state["db_loaded"] = False
        st.session_state["db_load_blocked"] = None
        st.session_state["db_switch_pending"] = None
        st.session_state["save_status"] = (
            "ok", f"Loaded the existing database at {path}.")
        st.session_state["dlg_settings"] = False
        st.rerun()

    st.markdown("---")
    st.caption("Or replace it with the session currently open in CAM. This "
               "overwrites the database above; a timestamped backup "
               "(`…acm_database.json.bak-replaced-…`) is written first.")
    overwrite_ok = st.checkbox(
        "I understand the existing database will be overwritten",
        key="dbsw_overwrite_cfm")
    if st.button("♻ Replace it with the current session", disabled=not overwrite_ok,
                 width="stretch", key="dbsw_overwrite"):
        backup = _backup_replaced_db(path)
        prefs = st.session_state["prefs"]
        prefs["db_custom_path"] = new_custom  # commit the path, then write to it
        save_prefs(prefs)
        # Explicit, checkbox-confirmed, already-backed-up overwrite: bypass the
        # Phase-3 shrink tripwire (a demo -> rich Replace is exactly the shrink
        # it would otherwise refuse). _backup_replaced_db above is the safety net.
        persist(allow_shrink=True)  # write the current session to the new path
        if st.session_state.get("save_status", ("", ""))[0] != "error":
            note = f"Overwrote {path} with the current session."
            if backup:
                note += f" Previous database backed up to {backup}."
            st.session_state["save_status"] = ("ok", note)
        st.session_state["db_switch_pending"] = None
        st.session_state["dlg_settings"] = False
        st.rerun()

    if st.button("Cancel", key="dbsw_cancel"):
        # Nothing to undo: the path pref was never moved off the old location
        # (it is committed only on Load / Replace), so the session keeps
        # pointing where it was. Just clear the decision and close.
        st.session_state["db_switch_pending"] = None
        st.session_state["dlg_settings"] = False
        st.rerun()


def _render_term_restore_confirm(payload: dict, term: str, sig: str) -> None:
    """Render the dry-run diff + typed-confirmation for a validated backup.

    Nothing is written until the teacher types ``RESTORE {term}`` exactly and
    clicks Restore. Cancel remembers this file's signature so re-running the
    dialog (which keeps the uploaded file in the widget) does not re-offer it."""
    diff = diff_term_backup(payload)
    counts = diff.get("counts", {})
    created = diff.get("created_at", "") or "an unknown time"
    st.warning(
        f"⚠ This replaces **{term}** wholesale. Any change made to {term} "
        f"**after this backup was created ({created})** is NOT in the file and "
        f"**will be lost**. Every other term is left untouched.")
    st.caption(
        f"Backup holds {counts.get('assignments', 0)} assignment(s), "
        f"{counts.get('scores', 0)} score(s), {counts.get('exam_results', 0)} "
        f"exam result(s) and {counts.get('comments', 0)} comment(s) across "
        f"{len(payload.get('classes', []))} class(es).")
    st.markdown("**Dry-run preview — what a restore would change:**")
    per_class = diff.get("per_class", {})
    if not per_class:
        st.caption("Nothing to change for this term.")
    for cls, d in per_class.items():
        st.markdown(f"**{cls or '(unassigned)'}**")
        bits = []
        if d["new_comments"]:
            bits.append(f"{len(d['new_comments'])} comment(s) restored into blanks")
        if d["changed_comments"]:
            bits.append(f"⚠ {len(d['changed_comments'])} comment(s) will OVERWRITE "
                        "a different current comment")
        if d["assignments_added"]:
            bits.append("adds assignment(s): " + ", ".join(d["assignments_added"]))
        if d["assignments_removed"]:
            bits.append("REMOVES live assignment(s) not in backup: "
                        + ", ".join(d["assignments_removed"]))
        bits.append(f"scores {d['live_scores']} → {d['backup_scores']}")
        st.caption(" · ".join(bits))
    if diff.get("remarks_fill") or diff.get("override_fill"):
        st.caption(
            f"Fill-blanks-only (existing entries kept): {diff.get('remarks_fill', 0)} "
            f"teacher remark(s), {diff.get('override_fill', 0)} final-override "
            "criterion/criteria added.")
    st.markdown(f"To proceed, type `RESTORE {term}` below:")
    typed = st.text_input("Confirm restore", key="tb_confirm",
                          label_visibility="collapsed",
                          placeholder=f"RESTORE {term}")
    cc = st.columns(2)
    do = cc[0].button("♻ Restore this term", type="primary", key="tb_do_restore",
                      width="stretch",
                      disabled=(typed.strip() != f"RESTORE {term}"))
    cancel = cc[1].button("Cancel", key="tb_cancel_restore", width="stretch")
    if do:
        backup = restore_term_backup(payload)
        if st.session_state.get("save_status", ("", ""))[0] != "error":
            note = f"Restored {term} from backup."
            if backup:
                note += f" Pre-restore database backed up to {backup}."
            st.session_state["save_status"] = ("ok", note)
        st.session_state["tb_restore_dismissed"] = sig
        st.session_state["dlg_settings"] = False
        st.rerun()
    if cancel:
        st.session_state["tb_restore_dismissed"] = sig
        st.rerun()


def _render_term_backup_section(prefs: dict) -> None:
    """The ⚙ Settings 'Term backup & restore' block (plan §3).

    Backup writes one term's snapshot to a teacher-chosen folder (zero risk — it
    only ever writes outside the database). Restore is a disaster tool behind a
    dry-run diff, a typed confirmation and an automatic pre-restore DB backup."""
    st.markdown("---")
    st.markdown("**🗄 Term backup & restore**")
    st.caption(
        "A deliberate end-of-term snapshot of one whole term, written to a "
        "folder you choose — it only ever writes OUTSIDE the database. Restore "
        "is a disaster tool: it replaces that term's data wholesale, behind a "
        "preview, a typed confirmation and an automatic backup.")

    folder = st.text_input(
        "Backup folder", value=prefs.get("term_backup_folder", ""),
        placeholder=r"e.g. C:\Users\you\CAM-term-backups   or a USB drive",
        key="tb_folder",
        help="Where ⬇ Back up term writes its files. Per-device (like every "
             "path setting); point it anywhere — even a USB stick or a non-cloud "
             "folder — for an off-site copy.")
    bc = st.columns([2, 1], vertical_alignment="bottom")
    cur = current_term()
    bk_term = bc[0].selectbox(
        "Term to back up", TERMS,
        index=TERMS.index(cur) if cur in TERMS else 0, key="tb_term")
    if bc[1].button("⬇ Back up term", key="tb_backup", width="stretch",
                    disabled=not folder.strip()):
        prefs["term_backup_folder"] = folder.strip()
        save_prefs(prefs)
        try:
            snapshot = build_term_backup(bk_term)
            path = write_term_backup(bk_term, folder.strip(), snapshot)
            c = snapshot["counts"]
            st.session_state["save_status"] = (
                "ok",
                f"Backed up {bk_term} to {path} — {c.get('assignments', 0)} "
                f"assignment(s), {c.get('scores', 0)} score(s), "
                f"{c.get('comments', 0)} comment(s).")
        except Exception as exc:
            st.session_state["save_status"] = (
                "error", f"Term backup failed: {exc}")
        st.session_state["dlg_settings"] = False
        st.rerun()

    up = st.file_uploader("⬆ Restore from backup… (.json)", type=["json"],
                          key="tb_restore_up")
    if up is None:
        return
    sig = f"{up.name}:{up.size}"
    if st.session_state.get("tb_restore_dismissed") == sig:
        st.caption("This backup was dismissed. Choose a different file, or "
                   "remove and re-add it to restore it.")
        return
    try:
        payload = json.loads(up.getvalue().decode("utf-8"))
    except Exception as exc:
        st.error(f"Could not read this file as JSON: {exc}")
        return
    term, err = validate_term_backup(payload)
    if err:
        st.error(err)
        return
    _render_term_restore_confirm(payload, term, sig)


@st.dialog("⚙ Settings — device & data", width="large")
def settings_dialog() -> None:
    """Per-device UI prefs + custom (cloud) database path.

    Inputs are wrapped in a form so dragging sliders does not rerun and close
    the modal; everything applies on Save. Prefs persist to the device-local
    ``local_device_prefs.json`` only — never the shared database."""
    prefs = st.session_state["prefs"]
    # Phase-2 adopt-vs-overwrite decision: when Save repointed the DB path at a
    # location that already holds a database, the form is replaced by the switch
    # panel until the teacher chooses (load / replace / cancel).
    pending = st.session_state.get("db_switch_pending")
    if pending:
        _render_db_switch_panel(pending)
        return
    st.caption("Saved to this device only (local_device_prefs.json), so each "
               "machine keeps its own layout while sharing one cloud database.")
    with st.form("settings_form"):
        st.markdown("**Cloud data sync**")
        custom = st.text_input(
            "Custom Database Path (folder or .json file)",
            value=prefs.get("db_custom_path", ""),
            placeholder=r"e.g. C:\Users\you\OneDrive\CAM   or   .../acm_database.json",
            help="Point this at a OneDrive/Google Drive folder for cross-device "
                 "continuity. A folder gets acm_database.json placed inside it, "
                 "plus one subfolder per class holding that class's grading "
                 "exports, exam scans, caches and term summaries. 🔄 Sync "
                 "(Window 1) scans those class subfolders for new files.",
        )
        st.markdown("**Column width ratios**")
        cw = st.columns(3)
        w1 = cw[0].slider("Window 1", 1, 10, int(prefs.get("col_w1", 4)), key="set_w1")
        w2 = cw[1].slider("Window 2", 1, 10, int(prefs.get("col_w2", 3)), key="set_w2")
        w3 = cw[2].slider("Window 3", 1, 10, int(prefs.get("col_w3", 5)), key="set_w3")
        st.markdown("**Scroll heights (px)**")
        ch = st.columns(3)
        h1 = ch[0].slider("Window 1", 200, 1200, int(prefs.get("h1", 520)), 20, key="set_h1")
        h2 = ch[1].slider("Window 2", 200, 1200, int(prefs.get("h2", 520)), 20, key="set_h2")
        h3 = ch[2].slider("Window 3", 200, 1200, int(prefs.get("h3", 640)), 20, key="set_h3")
        st.markdown("**Window 3 text-area heights (px)**")
        cht = st.columns(2)
        h_remarks = cht[0].slider("Remarks Box Height", 30, 500,
                                  int(prefs.get("h_remarks", 100)), 10, key="set_hrem")
        h_comment = cht[1].slider("Overall Comment Height", 30, 500,
                                  int(prefs.get("h_comment", 150)), 10, key="set_hcom")
        saved = st.form_submit_button("Save settings", type="primary",
                                      width="stretch")
    if saved:
        old_custom = prefs.get("db_custom_path", "")
        new_custom = custom.strip()
        old_resolved = os.path.normcase(os.path.abspath(resolve_db_path(old_custom)))
        new_resolved = os.path.normcase(os.path.abspath(resolve_db_path(new_custom)))
        path_changed = new_resolved != old_resolved
        # Layout prefs always apply now. The db_custom_path pref is committed
        # only once the destination is settled — see below.
        prefs.update({"col_w1": w1, "col_w2": w2, "col_w3": w3,
                      "h1": h1, "h2": h2, "h3": h3,
                      "h_remarks": h_remarks, "h_comment": h_comment})
        # Phase 2: repointing at a location that ALREADY holds a readable
        # database must not overwrite it with the current session (wipe
        # mechanism 1). Defer committing the path pref and hand off to the
        # adopt-vs-overwrite panel — the pref stays on the OLD location until
        # the teacher chooses, so an ESC-dismissed panel can never leave the
        # session pointed at (and about to autosave over) the existing DB.
        new_path = resolve_db_path(new_custom)
        if path_changed and db_file_state(new_path) == "ok":
            prefs["db_custom_path"] = old_custom  # keep pointing at the old path
            save_prefs(prefs)
            st.session_state["db_switch_pending"] = {
                "path": new_path,
                "new_custom": new_custom,
                "counts": _db_file_counts(new_path),
            }
            _render_db_switch_panel(st.session_state["db_switch_pending"])
            return
        # Unchanged path (layout-only Save) or a new location with no database
        # there: commit the path pref and persist as today (saves / creates).
        prefs["db_custom_path"] = new_custom
        save_prefs(prefs)
        persist()  # write the database to the (possibly new) path immediately
        st.session_state["dlg_settings"] = False
        st.rerun()
    st.caption(f"Active database file: `{db_path()}`")

    # ---- Report-card grades (shared, saved with the database) ------------
    # School-specific figures on top of the MYP criterion grades. Off by
    # default; a school that uses them enables them here. Saved into the shared
    # database (not device prefs) so the choice follows the teacher's data
    # across machines, and kept in its own form so it never rides along with
    # the DB-path adopt/overwrite flow above.
    st.markdown("---")
    st.markdown("**Report-card grades**")
    st.caption("School-specific figures layered on top of the MYP criterion "
               "grades. Saved with the database (shared across your devices). "
               "All off by default — switch on the ones your report cards use.")
    rc = st.session_state["report_cfg"]
    with st.form("report_cfg_form"):
        rc_myp = st.checkbox(
            "MYP Grade (1–7)", value=bool(rc.get("show_myp_grade", False)),
            help="Banded lookup from the sum of the final criterion grades.")
        rc_eff = st.checkbox(
            "Effort / English-use score",
            value=bool(rc.get("show_effort", False)),
            help="A per-student, per-term teacher score, editable in Window 3. "
                 "Also feeds the School Grade when that is enabled.")
        rc_sch = st.checkbox(
            "School Grade (1–10)", value=bool(rc.get("show_school_grade", False)),
            help="Banded lookup from the criterion-grade sum plus the "
                 "Effort/English-use score.")
        erange = st.columns(2)
        rc_emin = erange[0].number_input(
            "Effort / English-use min", min_value=0, max_value=20, step=1,
            value=int(rc.get("effort_min", 0)), key="rc_emin",
            help="Inclusive lowest selectable Effort score.")
        rc_emax = erange[1].number_input(
            "Effort / English-use max", min_value=0, max_value=20, step=1,
            value=int(rc.get("effort_max", 5)), key="rc_emax")
        rc_saved = st.form_submit_button("Save report-card grades",
                                         width="stretch")
    if rc_saved:
        lo, hi = int(rc_emin), int(rc_emax)
        if hi < lo:
            lo, hi = hi, lo
        rc.update({"show_myp_grade": bool(rc_myp),
                   "show_effort": bool(rc_eff),
                   "show_school_grade": bool(rc_sch),
                   "effort_min": lo, "effort_max": hi})
        persist()
        st.session_state["save_status"] = (
            "ok", "Report-card grade settings saved.")
        st.session_state["dlg_settings"] = False
        st.rerun()

    # ---- Manual re-scan (Sync is otherwise automatic) --------------------
    st.markdown("---")
    st.markdown("**🔄 Force full re-scan**")
    st.caption("Sync now runs automatically — at startup, when you open or "
               "return from grading, and when you add or watch a class. Use "
               "this only to force a full re-scan of every class folder + "
               "master directory (e.g. after hand-editing CSVs on another "
               "machine).")
    has_db = bool((prefs.get("db_custom_path") or "").strip())
    if st.button("Force full re-scan", key="force_rescan", disabled=not has_db,
                 help=None if has_db else
                 "Set a Custom Database Path above first."):
        summary = sync_all()
        _report_sync(summary)
        st.session_state["dlg_settings"] = False
        st.rerun()

    # ---- Term backup & restore -------------------------------------------
    _render_term_backup_section(prefs)

    # ---- Danger zone: wipe data (device/window settings are always kept) ----
    st.markdown("---")
    st.markdown("**⚠ Danger zone**")
    st.caption("Clears teaching data only — your window/device settings are kept.")

    # Targeted class deletion: remove one specific class (its units, grades,
    # roster and unit plan) without touching the other classes.
    names = class_names()
    if names:
        active = st.session_state["active_class"]
        del_target = st.selectbox(
            "Delete a specific class", names,
            index=names.index(active) if active in names else 0,
            key="del_class_sel",
            help="Permanently removes just this class and everything scoped to "
                 "it. All other classes are left untouched.",
        )
        n_units = len({a.name for a in gb().assignments
                       if getattr(a, "class_name", "") == del_target})
        n_roster = len(st.session_state["rosters"].get(del_target, []))
        st.caption(f"'{del_target}' has {n_units} unit(s) and {n_roster} "
                   f"student(s) on its roster.")
        del_ok = st.checkbox(
            f"Confirm: permanently delete class '{del_target}' (its units, "
            f"grades, roster & unit plan)", key="del_class_cfm")
        if st.button(f"Delete class '{del_target}'", disabled=not del_ok,
                     key="del_class_btn"):
            if delete_class(del_target):
                st.session_state["save_status"] = (
                    "ok", f"Deleted class '{del_target}'.")
            st.session_state["dlg_settings"] = False
            st.rerun()

    all_ok = st.checkbox("Confirm: wipe the ENTIRE database (all classes, "
                         "rosters, units & grades)", key="wipe_all_cfm")
    if st.button("Wipe entire database", disabled=not all_ok, key="wipe_all"):
        wipe_database_full()
        st.session_state["save_status"] = ("ok", "Database wiped — fresh start.")
        st.session_state["dlg_settings"] = False
        st.rerun()

    if st.button("Close", key="set_close"):
        st.session_state["dlg_settings"] = False
        st.rerun()


@st.dialog("⬆ Upload & stage files", width="large")
def upload_dialog() -> None:
    """File intake moved into a modal. Unit plan ingests immediately; grading
    CSVs land in the Window 1 staging queue for mapping before commit."""
    with st.expander("ℹ Accepted file formats", expanded=False):
        st.markdown(
            "**Unit plan (.docx)** — the official **IB MYP unit planner** "
            "template (the table-based document your school uses on "
            "ManageBac / Toddle, or from your MYP coordinator / the IB "
            "Programme Resource Centre). The parser reads its tables and "
            "needs at least:\n"
            "- a **Unit title** cell (title next to a *Unit title* label),\n"
            "- a **Statement of Inquiry** section (label row followed by "
            "the statement),\n"
            "- the target **criteria** listed as objective headings such as "
            "*A: Investigating* / *B – Developing*,\n"
            "- optionally *Key/Related concepts* and the *MYP year*.\n"
            "A slightly different template still loads — missing fields are "
            "simply left blank.\n\n"
            "**Grading CSV (.csv)** — one row per student, matched by the "
            "**Student Name** column holding the student's numeric ID (the "
            "part of their school email before the @ — the Window 2 roster "
            "links IDs to real names/emails). Recognised columns:\n"
            "- grades: either criterion-named columns like `Grade (Crit A)` / "
            "`Crit B` (auto-mapped, several per file allowed) or one generic "
            "`Grade` column (you pick the criterion at staging); values 0-8,\n"
            "- optional: `Due Date` (ISO, e.g. 2026-05-10), `Comment`, "
            "`Checked Keywords`,\n"
            "- a CSV with no grade column becomes a formative event "
            "(timeline only).\n"
            "The CAM Grading Workspace's exports already follow this format "
            "exactly.\n\n"
            "**Exam CSV** — an item-level export with one column per "
            "question plus `Total Score` (and optionally `Max Total`); "
            "detected automatically and graded to 0-8 later in Window 1's "
            "📝 Exam grading panel.")
    up_plan = st.file_uploader("Unit plan (.docx)", type=["docx"], key="up_plan")
    if up_plan is not None:
        sig = f"plan:{up_plan.name}:{up_plan.size}"
        if sig not in st.session_state["ingested_sigs"]:
            ingest_unit_plan(up_plan)
            st.session_state["ingested_sigs"].add(sig)
            st.success("Unit plan parsed.")
    ups = st.file_uploader(
        "Grading files (.csv) — multiple allowed",
        type=["csv"], accept_multiple_files=True, key="up_csv",
    )
    n_new = _stage_uploads(ups or [])
    if st.session_state["staging"]:
        st.info(f"{len(st.session_state['staging'])} file(s) waiting in the "
                f"staging queue. Close this to map & commit them.")
    if st.button("Done", key="up_close", type="primary"):
        st.session_state["dlg_upload"] = False
        st.rerun()


@st.dialog("🗃 Archived assignments", width="large")
def archived_dialog() -> None:
    """Restore or permanently delete soft-archived assignments (active class)."""
    archived = archived_for_active()
    st.caption(f"Class: **{st.session_state['active_class']}**")
    if not archived:
        st.caption("No archived assignments for this class. Deleting an "
                   "assignment from the table moves it here, where it can be "
                   "restored or purged.")
    for nm in archived:
        st.markdown(f"**{nm}**")
        cc = st.columns([1, 1, 1.4], vertical_alignment="center")
        if cc[0].button("Restore", key=f"arch_res_{nm}"):
            restore_assignment(nm)
            st.rerun()
        confirm = cc[1].checkbox("Confirm", key=f"arch_cfm_{nm}")
        if cc[2].button("Delete permanently", key=f"arch_del_{nm}",
                        disabled=not confirm):
            delete_assignment_permanent(nm)
            st.rerun()
        st.markdown("---")
    if st.button("Close", key="arch_close"):
        st.session_state["dlg_archived"] = False
        st.rerun()


@st.dialog("➕ Add assignment / exam")
def add_assignment_dialog() -> None:
    """Manually create an assignment or exam row for the active class.

    An **assignment** carries the chosen criterion immediately: each student
    shows an editable "0 (missing)" cell under it in Window 3 where the
    teacher enters individual grades. An **exam** starts with raw marks off
    the 0-8 scale: its analytics dialog offers 🛠 Exam setup (the grading
    workspace's question programmer), and grades arrive via the exported exam
    CSV, then get graded in Window 1's 📝 Exam grading panel."""
    st.caption(f"Class: **{st.session_state['active_class']}** · "
               f"Term: **{current_term()}**")
    with st.form("add_asg_form"):
        name = st.text_input("Name", placeholder="e.g. Perspective Drawing Task")
        kind = st.radio("Type", ["Assignment", "Exam"], horizontal=True,
                        key="add_asg_kind")
        when = st.date_input("Deadline / date", value=date.today(),
                             key="add_asg_date")
        crit = st.selectbox(
            "Criterion", CRIT_ORDER,
            format_func=lambda c: f"{c} — {Criterion(c).label}",
            key="add_asg_crit",
            help="Assignments: grades entered in Window 3 are recorded under "
                 "this criterion. Exams: this is the default criterion offered "
                 "when grading the raw marks.")
        submitted = st.form_submit_button("Create", type="primary",
                                          width="stretch")
    st.caption("ℹ Once created, an assignment counts as **missing = 0** for "
               "every student until you enter their grades (or switch it Off "
               "in the table). An ungraded exam stays out of the math.")
    if submitted:
        name = (name or "").strip()
        if not name:
            st.error("Give the assignment/exam a name.")
            return
        if name in {a.name for a in gb().assignments}:
            st.error(f"'{name}' already exists.")
            return
        is_exam = kind == "Exam"
        gb().register_assignment(Assignment(
            name=name, criteria=[crit], ingested_at=datetime.combine(
                when, datetime.min.time()),
            note=("manually created exam — set it up via 🛠 Exam setup"
                  if is_exam else "manually created assignment"),
            class_name=st.session_state["active_class"],
            term=current_term(), is_exam=is_exam))
        st.session_state["date_override"][name] = when
        st.session_state["active"][name] = True
        st.session_state["archived"].discard(name)
        persist()
        st.session_state["save_status"] = (
            "ok", f"Created {kind.lower()} '{name}' ({current_term()}, "
                  f"Criterion {crit}).")
        st.rerun()


# --------------------------------------------------------------------------
# WINDOW 1 - Timeline & assignment system
# --------------------------------------------------------------------------

def render_window1() -> None:
    st.subheader("1 · Timeline & Assignments")

    # Sync is no longer a button here — it runs automatically (session start,
    # entering/returning from grading, class add/watch). The manual escape hatch
    # ("Force full re-scan") lives in ⚙ Settings for the OneDrive/multi-machine
    # case. vertical_alignment="center" keeps the two remaining buttons centered
    # against the adjacent controls rather than sitting high.
    bar = st.columns([1.4, 1.2], vertical_alignment="center")
    if bar[0].button("⬆ Upload & stage files", key="open_upload",
                     width="stretch"):
        upload_dialog()
    n_arch = len(archived_for_active())
    if bar[1].button(f"🗃 Archived ({n_arch})", key="open_archived",
                     width="stretch"):
        archived_dialog()

    _render_staging_queue()

    plan = st.session_state["unit_plan"]
    ac = active_class_dict()
    phase = current_phase()
    if plan:
        st.info(f"**{plan.unit_title}**  ·  {class_label(ac)}  ·  rubric phase "
                f"{phase}\n\n_{plan.statement_of_inquiry}_")
    elif ac:
        st.caption(f"{class_label(ac)}  ·  rubric phase {phase}. "
                   f"Upload a unit plan via ⬆ Upload & stage files.")

    if st.button("➕ Add assignment / exam", key="add_asg"):
        add_assignment_dialog()

    _render_assignment_table()
    _render_exam_banding()


def _stage_uploads(ups) -> int:
    """Add newly-seen uploads to the staging queue (never auto-ingest).

    Returns the number of newly-staged files."""
    staging = st.session_state["staging"]
    committed = st.session_state["ingested_sigs"]
    added = 0
    for uf in ups:
        sig = f"csv:{uf.name}:{uf.size}"
        if sig in committed or sig in staging:
            continue
        raw = uf.getvalue()
        text = raw.decode("utf-8-sig", errors="replace")
        try:
            header = next(csv.reader(io.StringIO(text)))
        except StopIteration:
            header = []
        header = [h.strip() for h in header]
        detected = map_criterion_columns(header)
        exam = is_exam_csv(header)
        staging[sig] = {
            "name": uf.name,
            "data": raw,
            "header": header,
            "detected": {k: v.value for k, v in detected.items()},
            "has_grade": "Grade" in header,
            "is_exam": exam,
            "questions": exam_question_columns(header) if exam else [],
            "default_name": os.path.splitext(uf.name)[0],
        }
        added += 1
    return added


def _staged_choices(sig: str, item: dict):
    """Read the current staging-row widget choices for one file."""
    name = st.session_state.get(f"stg_name_{sig}", item["default_name"])
    if item.get("is_exam") or item["detected"]:
        crit = None
    elif item["has_grade"]:
        crit = st.session_state.get(f"stg_crit_{sig}", CRIT_ORDER[1])
    else:
        crit = None
    when = st.session_state.get(f"stg_date_{sig}")
    return name, crit, when


def _render_staging_queue() -> None:
    """Render the uncommitted staging table: map criterion + date, then commit.

    While anything sits here the cockpit's math/visualisation is paused (see
    render_window3 / render_tray) so partially-mapped files never leak into the
    aggregation."""
    staging = st.session_state["staging"]
    if not staging:
        return
    st.markdown(f"**Staging queue — {len(staging)} file(s) awaiting commit**")
    st.caption("Map each file's criterion and (optionally) its due date, "
               "then commit. Grades stay out of the math until committed.")
    for sig, item in list(staging.items()):
        with st.container(border=True):
            st.markdown(f"📄 **{item['name']}**")
            cols = st.columns([2.2, 1.5, 1.5])
            cols[0].text_input("Assignment name", value=item["default_name"],
                               key=f"stg_name_{sig}")
            if item.get("is_exam"):
                cols[1].caption(f"📝 Exam export — {len(item['questions'])} "
                                "question(s), raw marks. Grade it in Window 1 "
                                "after committing.")
            elif item["detected"]:
                cols[1].caption("Criteria auto-detected: "
                                + ", ".join(f"{k}→{v}"
                                            for k, v in item["detected"].items()))
            elif item["has_grade"]:
                cols[1].selectbox("Map 'Grade' → criterion", CRIT_ORDER,
                                  index=1, key=f"stg_crit_{sig}")
            else:
                cols[1].caption("No gradeable column → formative event.")
            cols[2].date_input("Due date (optional)", value=None,
                               key=f"stg_date_{sig}")
            b = st.columns([1, 1, 3])
            if b[0].button("Commit", key=f"stg_commit_{sig}", type="primary"):
                name, crit, when = _staged_choices(sig, item)
                n = commit_staged(sig, name, crit, when)
                st.session_state["save_status"] = (
                    "ok", f"Committed {n} score(s) into '{name}'.")
                st.rerun()
            if b[1].button("Remove", key=f"stg_remove_{sig}"):
                del st.session_state["staging"][sig]
                st.rerun()

    if len(staging) > 1 and st.button("Commit all staged files",
                                      key="stg_commit_all", type="primary"):
        for sig in list(staging.keys()):
            item = staging.get(sig)
            if not item:
                continue
            name, crit, when = _staged_choices(sig, item)
            commit_staged(sig, name, crit, when)
        st.rerun()


@st.dialog("⋯ Manage assignment")
def manage_assignment_dialog(name: str, r: dict) -> None:
    """Per-assignment management: rename, due date, term inclusion, archive.

    A centered modal (rather than a row popover) so it is never clipped or
    squeezed by the narrow Manage column / Window 1 width on a small screen."""
    st.markdown(f"**Manage · {name}**")
    new = st.text_input("Rename", value=name, key=f"rn_{name}")
    if st.button("Apply rename", key=f"rnb_{name}"):
        if rename_assignment(name, new):
            st.rerun()
    d = st.date_input("Due date", value=r["date"].date(), key=f"date_{name}")
    if st.button("Apply date", key=f"applyd_{name}"):
        apply_date_override(name, d)
        persist()
        st.rerun()
    st.markdown(f"**Counts in** (ingested into {r.get('term', current_term())})")
    st.caption(f"{current_term()} is set by the On checkbox in the table; tick "
               "other terms here to include this work in their assessments too.")
    for t in TERMS:
        if t == current_term():
            continue
        on_t = st.checkbox(t, value=assignment_on(name, t), key=f"ont_{t}_{name}")
        term_active_map(t)[name] = on_t
    st.markdown("---")
    if st.button("🗑 Delete (archive)", key=f"del_{name}"):
        archive_assignment(name)
        st.rerun()


def _render_assignment_table() -> None:
    """Active-assignment table (the whole window scrolls), with a per-row ⋯
    Manage dialog (rename, due date, term inclusion, archive)."""
    st.markdown("**Active assignments** (April → March)")
    st.caption(f"**On** = counts in **{current_term()}**'s assessment. Switch "
               "the term in the top bar to choose that term's set — each term "
               "keeps its own selection.")
    rows = assignment_table()
    if not rows:
        st.caption("No assignments yet — stage and commit a grading CSV above.")
        return

    head = st.columns([0.5, 2.4, 1.4, 0.6, 1.0])
    for col, cap in zip(head, ("On", "Assignment", "Crit · date", "Subs", "Manage")):
        col.caption(cap)
    for r in rows:
        name = r["name"]
        # vertical_alignment centers the On/Manage controls against the row text.
        c = st.columns([0.5, 2.4, 1.4, 0.6, 1.0], vertical_alignment="center")
        val = c[0].checkbox(" ", value=assignment_on(name),
                            key=f"act_{current_term()}_{name}",
                            label_visibility="collapsed")
        st.session_state["active"][name] = val
        if c[1].button(name, key=f"sel_{name}", width="stretch"):
            st.session_state["sel_assignment"] = name
            show_analytics_dialog(name)
        if r.get("is_exam"):
            # Exams show their RAW class average (e.g. "raw ø 31.5/45"), not a
            # band — banding happens in the panel below the table.
            raw = (f"raw ø {r['raw_avg']}/{r['max_total']}"
                   if r["raw_avg"] is not None else "raw · ungraded")
            crit = f" → Crit {r['criteria']}" if r["criteria"] != "—" else ""
            c[2].caption(f"📝 {raw}{crit}  ·  {r['date']:%b %d}")
        else:
            c[2].caption(f"{r['criteria']}  ·  {r['date']:%b %d}")
        c[3].caption(str(r["submissions"]))
        if c[4].button("⋯", key=f"mng_{name}", width="stretch"):
            manage_assignment_dialog(name, r)


def _class_exams():
    """Non-archived exam assignments belonging to the active class."""
    return [a for a in gb().assignments
            if getattr(a, "is_exam", False)
            and getattr(a, "class_name", "") == st.session_state.get("active_class", "")
            and a.name not in st.session_state["archived"]]


def _exam_results_for(name: str):
    """[(Student, ExamResult)] for one exam, roster-sorted."""
    pairs = [(s, s.exam_results[name]) for s in gb()
             if name in getattr(s, "exam_results", {})]
    return sorted(pairs, key=lambda p: student_label(p[0]).lower())


def _render_exam_banding() -> None:
    """Window 1 exam panel: raw totals per student + a 0-8 band dropdown.

    Exam CSVs arrive with raw item-level marks (e.g. 31/45) that cannot enter
    the 0-8 gradebook directly. This panel reflects every student's raw score
    and lets the teacher assign each a 0-8 band (pre-filled with the
    proportional suggestion, or the band already on record). Applying writes
    ordinary CriterionScores under the chosen criterion, so the exam then
    behaves like any other assessed task downstream.
    """
    exams = _class_exams()
    if not exams:
        return
    with st.expander("📝 Exam grading — raw marks → 0–8 grade", expanded=False):
        names = [a.name for a in exams]
        sel = st.selectbox("Exam", names, key="exam_band_sel")
        asg = next(a for a in exams if a.name == sel)
        results = _exam_results_for(sel)
        if not results:
            st.caption("This exam's CSV carried no student rows.")
            return
        avg = mean(r.total for _, r in results)
        st.caption(f"Raw marks out of **{asg.max_total or '?'}** · class raw "
                   f"average **{avg:.1f}** · {len(results)} student(s) · "
                   f"questions: {', '.join(asg.question_labels) or '—'}")

        crit_default = asg.criteria[0] if asg.criteria else "A"
        crit = st.selectbox(
            "Counts toward criterion", CRIT_ORDER,
            index=CRIT_ORDER.index(crit_default) if crit_default in CRIT_ORDER else 0,
            key=f"exam_crit_{sel}",
            help="The 0–8 grade each student receives is recorded as a normal "
                 "score under this MYP criterion.")

        # Bands already applied (rerunning the panel edits them in place).
        existing = {s.student_id: sc.value
                    for s, sc in scores_for_assignment(sel) if sc.is_valid}

        head = st.columns([2.2, 1.1, 0.9, 1.3], vertical_alignment="center")
        for col, cap in zip(head, ("Student", "Raw score", "%", "Grade (0–8)")):
            col.caption(cap)
        for s, r in results:
            row = st.columns([2.2, 1.1, 0.9, 1.3], vertical_alignment="center")
            row[0].write(student_label(s))
            row[1].write(f"{r.total} / {r.max_total or asg.max_total or '?'}")
            row[2].write(f"{r.percent:.0f}%")
            default = existing.get(s.student_id, r.suggested_band())
            row[3].selectbox(
                " ", list(range(0, 9)),
                index=default if 0 <= default <= 8 else 0,
                key=f"band_{sel}_{s.student_id}",
                label_visibility="collapsed")

        if st.button("Apply grades to gradebook", key=f"apply_bands_{sel}",
                     type="primary"):
            n = _apply_exam_bands(asg, crit, results)
            st.session_state["save_status"] = (
                "ok", f"Graded {n} student(s) for '{sel}' → Criterion {crit}.")
            st.rerun()


def _apply_exam_bands(asg, crit_letter: str, results) -> int:
    """Write the chosen 0-8 bands as CriterionScores (replacing prior bands)."""
    crit = Criterion(crit_letter)
    when = asg.ingested_at or datetime.now()
    if asg.name in st.session_state["date_override"]:
        when = datetime.combine(
            st.session_state["date_override"][asg.name], datetime.min.time())
    applied = 0
    for s, r in results:
        band = st.session_state.get(f"band_{asg.name}_{s.student_id}")
        if band is None:
            continue
        # One band per student per exam: drop any earlier banding first.
        for c, bucket in list(s.scores.items()):
            s.scores[c] = [sc for sc in bucket if sc.assignment != asg.name]
        s.add_score(CriterionScore(
            criterion=crit, value=int(band), timestamp=when,
            source=f"exam:{asg.name}", assignment=asg.name,
            comment=r.comment,
            note=f"banded from raw {r.total}/{r.max_total or asg.max_total}",
        ))
        applied += 1
    asg.criteria = [crit_letter]
    asg.score_count = applied
    st.session_state["active"].setdefault(asg.name, True)
    persist()
    return applied


@st.dialog("Assignment analytics", width="large")
def show_analytics_dialog(name: str) -> None:
    """Pop-up modal: assignment metrics + a Plotly band-distribution chart.

    Replaces the old inline analytics column; opens only when an assignment
    row is clicked in the table above."""
    import plotly.graph_objects as go

    match = next((r for r in assignment_table() if r["name"] == name), None)
    if not match:
        st.write("No analytics available for this assignment.")
        return
    st.markdown(f"### {name}")

    # ---- App bridge: hand this assignment/exam to the grading workspace ----
    # Spawns the Flask sub-app (./cam_grading_workspace) on its own port so
    # Streamlit keeps running, passing the assignment's folder ref/name in the
    # URL so the workspace opens straight into it. Manually-created rows have
    # no submission folder to grade: a manual ASSIGNMENT is marked directly in
    # Window 3, and a manual EXAM offers the workspace's Exam Setup instead.
    cls = st.session_state.get("active_class", "")
    asg_rec = next((a for a in gb().assignments
                    if a.name == name and getattr(a, "class_name", "") == cls),
                   None)
    folder_ref = getattr(asg_rec, "folder_ref", "") if asg_rec else ""
    has_folder = bool(folder_ref)
    is_exam_rec = bool(getattr(asg_rec, "is_exam", False)) if asg_rec else False
    has_exam_data = any(name in getattr(s, "exam_results", {}) for s in gb())

    # A local assignment folder may be deleted after its CSV is exported (a
    # sanctioned workflow — the marks already live in CAM). When that happens
    # the Grade button would open the workspace onto a folder that no longer
    # exists, so surface a re-link affordance instead of a dead launch. Drive
    # refs and existing local folders skip this entirely.
    folder_missing = bool(folder_ref) and _master_is_local(folder_ref) \
        and not os.path.isdir(folder_ref)
    if folder_missing:
        st.warning(
            f"📁 The local folder for **{name}** is missing "
            f"(`{folder_ref}`). Your grades are safe in CAM — this only stops "
            "the workspace re-opening it. Re-link the folder to regrade, or "
            "grade elsewhere and enter marks in Window 3.")
        newp = st.text_input(
            "Re-link folder (existing local path)", key=f"relink_{name}",
            placeholder=r"e.g. C:\Users\you\OneDrive\Y7 Art\Artist Looking")
        if st.button("🔗 Re-link folder", key=f"relinkbtn_{name}",
                     width="stretch"):
            newp = (newp or "").strip()
            if not newp or not os.path.isdir(newp):
                st.error("Enter a path to a folder that exists on disk.")
            else:
                asg_rec.folder_ref = os.path.abspath(newp)
                persist()
                st.success("Folder re-linked — click 🖌 Grade this "
                           "Assignment/Exam to open it.")
                st.rerun()

    if is_exam_rec:
        if st.button("🛠 Exam setup", key=f"examsetup_{name}",
                     type="primary", width="stretch",
                     help="Open the grading workspace's exam question "
                          "programmer for this exam (scan pages, define "
                          "questions, then grade item by item)."):
            url = launch_exam_setup(name)
            if url:
                st.success(f"Exam setup opened in your browser — or click: "
                           f"[{url}]({url})")
            else:
                st.error(st.session_state.get("save_status", ("", ""))[1]
                         or "Could not launch the exam setup.")
    if (has_folder and not folder_missing) or has_exam_data:
        if st.button("🖌 Grade this Assignment/Exam", key=f"bridge_{name}",
                     type="primary" if not is_exam_rec else "secondary",
                     width="stretch",
                     help="Open the CAM Grading Workspace targeting this "
                          f"assignment (port {GRADING_PORT} — Streamlit is "
                          "never blocked)."):
            url = launch_grading_workspace(name)
            if url:
                st.success(f"Grading workspace opened in your browser — or "
                           f"click: [{url}]({url})")
            else:
                st.error(st.session_state.get("save_status", ("", ""))[1]
                         or "Could not launch the grading workspace.")
    elif not is_exam_rec and not has_folder:
        st.info("✏ This assignment is marked directly in CAM: focus each "
                "student in Window 2, then in Window 3 click the "
                f"**0 (missing)** cell under **{name}** to enter their grade "
                "(with late/excused flags and a comment).")

    if match.get("is_exam"):
        # Item-level exam analytics: raw totals + per-question averages.
        results = _exam_results_for(name)
        st.caption(f"📝 Exam  ·  {match['date']:%b %d, %Y}  ·  "
                   f"marked out of {match['max_total'] or '?'}")
        a, b, cc, d = st.columns(4)
        a.metric("Submissions", match["submissions"])
        a2 = (f"{match['raw_avg']}/{match['max_total']}"
              if match["raw_avg"] is not None else "—")
        b.metric("Avg raw score", a2)
        totals = [r.total for _, r in results]
        cc.metric("Spread (sd)", round(pstdev(totals), 2) if len(totals) > 1 else 0.0)
        d.metric("Graded", "Crit " + match["criteria"]
                 if match["criteria"] != "—" else "not yet")

        labels = next((a_.question_labels for a_ in gb().assignments
                       if a_.name == name and getattr(a_, "is_exam", False)), [])
        if labels and results:
            q_avg = [round(mean([r.questions.get(lbl, 0) for _, r in results]), 2)
                     for lbl in labels]
            fig = go.Figure(go.Bar(x=labels, y=q_avg, marker_color="#ff7f0e"))
            fig.update_layout(
                height=320, margin=dict(l=10, r=10, t=30, b=10),
                xaxis=dict(title="Question"), yaxis=dict(title="Average mark"),
                title="Item-level analysis — average mark per question",
            )
            style_chart(fig)
            st.plotly_chart(fig, width="stretch")
        if match["criteria"] == "—":
            st.caption("Assign 0–8 grades in Window 1's 📝 Exam grading panel "
                       "to bring this exam into the criterion math.")
            return
        # Bands were applied — fall through to the normal band distribution.
    else:
        st.caption(f"Criteria {match['criteria']}  ·  {match['date']:%b %d, %Y}")
    a, b, cc, d = st.columns(4)
    a.metric("Submissions", match["submissions"])
    b.metric("Avg grade", match["avg"] if match["avg"] is not None else "—")
    cc.metric("Spread (sd)", match["spread"])
    d.metric("Lates", match["lates"])

    # Unmatched-works pool (Phase 3): graded rows that matched no roster student
    # sit here until visually assigned in Window 2. Surface the count so they are
    # never invisible; the ⚠ popover in Window 2 is where they get resolved.
    pool = st.session_state["unmatched_works"].get(cls, {}).get(name, [])
    if has_folder and pool:
        st.warning(
            f"🧩 **{len(pool)} unmatched work(s)** for this assignment did not "
            "match a roster student. Resolve them in Window 2: open a student "
            "who is missing this work and pick their submission from the grid.")

    dist = {}
    for _, sc in scores_for_assignment(name):
        if sc.is_valid:
            dist[sc.value] = dist.get(sc.value, 0) + 1
    if dist:
        xs = list(range(0, 9))
        ys = [dist.get(band, 0) for band in xs]
        fig = go.Figure(go.Bar(x=xs, y=ys, marker_color="#1f77b4"))
        fig.update_layout(
            height=320, margin=dict(l=10, r=10, t=30, b=10),
            xaxis=dict(title="Grade (0–8)", dtick=1),
            yaxis=dict(title="Count"), title="Grade distribution",
        )
        style_chart(fig)
        st.plotly_chart(fig, width="stretch")
    else:
        st.caption("No valid scored data to plot yet.")


# --------------------------------------------------------------------------
# WINDOW 2 - Roster & submission tracker
# --------------------------------------------------------------------------

def parse_roster(uploaded):
    """Parse an uploaded Google Classroom roster into matchable entries.

    Delegates to the engine's :func:`parse_classroom_roster`, which keys each
    student by the numeric ID embedded in their email (the email's local part),
    so roster rows line up with the numeric IDs the grading CSV uses."""
    text = uploaded.getvalue().decode("utf-8-sig", errors="replace")
    return parse_classroom_roster(text)


def sort_roster_gojuon(entries: list) -> list:
    """Return roster entries ordered by gojūon (hiragana あいうえお) reading.

    A Google Classroom list arrives alphabetised by the Latin spelling, which
    is not the order a Japanese register uses. We sort by the reading of the
    SURNAME first, then the given name. Names are romaji, so each part is
    converted to its kana mora sequence (see :func:`engine.gojuon_sort_key`).

    Re-ordering the roster is display-only: every mark lives in the gradebook
    keyed by student ID, independent of roster position (the same reason the
    ↑/↓ buttons are safe), so a re-sort never disturbs anyone's grades."""
    def key(entry):
        name = (entry.get("name") or entry.get("key") or "").strip()
        first = (entry.get("first") or "").strip()
        # ``name`` is stored "Surname First"; peel the given name off the end to
        # isolate the surname (which may itself be several tokens).
        if first and name.lower().endswith(first.lower()) \
                and len(name) > len(first):
            last = name[:len(name) - len(first)].strip()
        else:
            parts = name.split()
            last = " ".join(parts[:-1]) if len(parts) > 1 else name
        given = first or (name.split()[-1] if name.split() else "")
        return (gojuon_sort_key(last), gojuon_sort_key(given))
    return sorted(entries, key=key)


@st.dialog("➕ Add student")
def add_student_dialog() -> None:
    """Manually add ONE student to the active class's roster.

    First name and email are compulsory: the email's local part is the
    numeric ID the grading CSVs key on, so a student added without it could
    never be matched to their grades."""
    st.caption(f"Class: **{st.session_state['active_class']}** — one student "
               "at a time.")
    with st.form("add_student_form"):
        first = st.text_input("First name (required)")
        last = st.text_input("Surname (optional)")
        email = st.text_input("Email address (required)",
                              placeholder="e.g. 100001@school.ed.jp",
                              help="The part before @ must hold the student's "
                                   "numeric ID — it is the key grading CSVs "
                                   "are matched on.")
        applied = st.form_submit_button("Apply", type="primary",
                                        width="stretch")
    if applied:
        first = (first or "").strip()
        email = (email or "").strip()
        last = (last or "").strip()
        if not first or not email:
            st.error("First name AND email address are both required.")
            return
        if "@" not in email:
            st.error("That doesn't look like an email address.")
            return
        key = student_id_from_email(email)
        roster = st.session_state["roster"]
        if any(e.get("key") == key for e in roster):
            st.error(f"A student with ID {key} is already on this roster.")
            return
        name = f"{last} {first}".strip()
        roster.append({"key": key, "name": name, "email": email,
                       "first": first, "gender": ""})
        # Make sure a gradebook record exists so Focus/Window 3 work at once.
        gb().get_or_create(key, name)
        persist()
        st.session_state["save_status"] = ("ok", f"Added {name} ({key}).")
        st.rerun()


@st.dialog("➖ Remove student")
def remove_student_dialog() -> None:
    """Remove a student from the roster — ARCHIVED, not deleted.

    The entry moves to this class's archive and the student's gradebook
    record (scores, comments, exam results) is kept intact, so no assignment
    or comment is ever orphaned. Archived students can be restored below."""
    cls = st.session_state["active_class"]
    roster = st.session_state["roster"]
    if roster:
        labels = {f"{e.get('name') or e['key']}  ({e['key']})": e["key"]
                  for e in roster}
        pick = st.selectbox("Student to remove", list(labels.keys()),
                            key="rm_student_sel")
        st.caption("Removing archives the student in the background — their "
                   "grades and comments stay in the database, so nothing is "
                   "orphaned or lost.")
        sure = st.checkbox("Confirm: remove this student from the roster",
                           key="rm_student_cfm")
        if st.button("Remove (archive)", type="primary", disabled=not sure,
                     key="rm_student_btn"):
            key = labels[pick]
            entry = next(e for e in roster if e["key"] == key)
            st.session_state["archived_students"].setdefault(cls, []).append(entry)
            set_active_roster([e for e in roster if e["key"] != key])
            if st.session_state["focus_sid"] == key:
                st.session_state["focus_sid"] = None
            persist()
            st.session_state["save_status"] = (
                "ok", f"Archived {entry.get('name') or key} — restore any "
                      "time from ➖ Remove student.")
            st.rerun()
    else:
        st.caption("No students on this roster.")

    archived = st.session_state["archived_students"].get(cls, [])
    if archived:
        st.markdown("---")
        st.markdown("**Archived students**")
        for i, e in enumerate(archived):
            cc = st.columns([3, 1], vertical_alignment="center")
            cc[0].write(f"{e.get('name') or e['key']}  ({e['key']})")
            if cc[1].button("Restore", key=f"rest_stu_{e['key']}_{i}"):
                archived.remove(e)
                set_active_roster(st.session_state["roster"] + [e])
                persist()
                st.rerun()


def _thumb_cache_path(path: str, width: int, mtime: int) -> str:
    """Cache filename for one rendered work thumbnail, keyed by absolute path +
    mtime + width so a re-saved source (new mtime) or a different requested size
    never reuses a stale render. Mirrors the grading workspace's own
    ``_thumb_cache_path``."""
    key = hashlib.sha1(
        os.path.normcase(os.path.abspath(path)).encode("utf-8")).hexdigest()[:16]
    return os.path.join(THUMB_CACHE_DIR, f"{key}__{mtime}__{width}.png")


def _work_page_png_bytes(path: str, max_width: int) -> bytes:
    """PNG bytes of a file's first page (PDF) / image, LANCZOS-downscaled to
    ``max_width``. The same PyMuPDF/Pillow pipeline the grading workspace uses
    for its previews (``exam_engine.page_png_bytes``), replicated here so CAM
    needn't import the sub-app package. Both imports are lazy so a missing
    PyMuPDF only degrades this one feature (filename-only tiles), never boot."""
    from PIL import Image
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        import fitz  # PyMuPDF
        with fitz.open(path) as doc:
            if doc.page_count < 1:
                raise ValueError("empty PDF")
            pix = doc[0].get_pixmap(dpi=150)
            img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
        if img.width > max_width:
            ratio = max_width / img.width
            img = img.resize((max_width, int(img.height * ratio)), Image.LANCZOS)
        buf = io.BytesIO()
        img.save(buf, "PNG")
        return buf.getvalue()
    with Image.open(path) as im:
        img = im.convert("RGB")
        if img.width > max_width:
            ratio = max_width / img.width
            img = img.resize((max_width, int(img.height * ratio)), Image.LANCZOS)
        buf = io.BytesIO()
        img.save(buf, "PNG")
        return buf.getvalue()


def _render_work_png(path: str, width: int) -> Optional[bytes]:
    """Rendered PNG bytes for ``path`` at ``width``, disk-cached under
    ``thumb_cache/``. Returns ``None`` on any failure (missing PyMuPDF, an
    unreadable/unsupported file) so the caller falls back to a filename-only
    tile — a preview is a convenience, never a precondition for assigning."""
    try:
        mtime = int(os.path.getmtime(path))
    except OSError:
        return None
    cache_path = _thumb_cache_path(path, width, mtime)
    if os.path.exists(cache_path):
        try:
            with open(cache_path, "rb") as fh:
                return fh.read()
        except OSError:
            pass
    try:
        png = _work_page_png_bytes(path, width)
    except Exception:
        return None
    try:
        os.makedirs(THUMB_CACHE_DIR, exist_ok=True)
        tmp = cache_path + f".{os.getpid()}.tmp"
        with open(tmp, "wb") as fh:
            fh.write(png)
        os.replace(tmp, cache_path)          # atomic publish
    except OSError:
        pass                                 # caching is best-effort
    return png


def _find_work_file(folder_ref: str, files_cell: str) -> Optional[str]:
    """Absolute path of a pooled work's first on-disk file under ``folder_ref``.

    The pool row's ``files`` cell is the export's ``"; "``-joined basenames,
    newest first. A local assignment folder may be laid out flat (stem =
    student) or subfolder-per-student, so walk the tree once and return the
    first listed basename that survives on disk, honouring the newest-first
    order. ``None`` when nothing matches (the folder was pruned after export —
    a sanctioned workflow)."""
    wanted = [t.strip() for t in (files_cell or "").split(";") if t.strip()]
    if not wanted or not os.path.isdir(folder_ref):
        return None
    found: dict = {}
    for root, _dirs, names in os.walk(folder_ref):
        for n in names:
            low = n.lower()
            if low not in found:
                found[low] = os.path.join(root, n)
    for w in wanted:
        p = found.get(w.lower())
        if p:
            return p
    return None


def _work_thumbnail(row: dict, folder_ref: str, width: int):
    """``(png_bytes | None, note)`` for one pooled work at ``width``.

    ``png_bytes`` is ``None`` whenever no on-disk preview can be shown; ``note``
    is a short teacher-readable reason for the tile caption. Drive-backed pools
    (rare — a transfer student off the roster) get a filename-only tile: no
    Drive fetch in this phase."""
    if not folder_ref or not _master_is_local(folder_ref):
        return None, "Preview not available (Drive work)."
    if not os.path.isdir(folder_ref):
        return None, "Source folder no longer on disk."
    path = _find_work_file(folder_ref, row.get("files", ""))
    if not path:
        return None, "Source file no longer on disk."
    png = _render_work_png(path, width)
    if png is None:
        return None, "No preview for this file type."
    return png, ""


def _work_caption(row: dict) -> str:
    """The tile's secondary filename caption — the newest listed file's basename.

    Filenames here are meaningless camera-roll noise (that's the whole reason
    the work is unmatched), so this is a hint under the image, never the
    identification channel."""
    first = (row.get("files", "") or "").split(";")[0].strip()
    return first or "(no filename)"


@st.dialog("🧩 Match unmatched works", width="large")
def match_works_dialog(student_key: str, assignment: str) -> None:
    """Resolve one student's unmatched work visually (Phase 4).

    The dialog is scoped to ONE student who is missing ONE assignment, so
    matching is a single click — no dropdown. It shows a scrollable grid of the
    assignment's pooled works (thumbnail + filename caption); the teacher clicks
    the one that is this student's, which calls :func:`assign_work` (write the
    durable alias, materialise the score, drop from the pool) and reruns.
    Clicking a thumbnail's ⤢ Enlarge renders it full-width (~1600px) to read a
    handwritten name; ⤡ Shrink returns to the grid. Both toggles rerun only the
    dialog fragment so the grid stays open; a successful assign reruns the whole
    app so Window 2's ⚠ count and Window 3's score refresh."""
    cls = st.session_state.get("active_class", "")
    pool = st.session_state["unmatched_works"].get(cls, {}).get(assignment, [])
    entry = next((e for e in st.session_state["roster"]
                  if e.get("key") == student_key), None)
    if entry:
        first = (entry.get("first") or "").strip() or (
            entry.get("name", "").split()[-1] if entry.get("name") else "")
    else:
        first = ""
    first = first or student_key
    asg_rec = next((a for a in gb().assignments
                    if a.name == assignment
                    and getattr(a, "class_name", "") == cls), None)
    folder_ref = getattr(asg_rec, "folder_ref", "") if asg_rec else ""

    st.markdown(
        f"**{assignment}** — pick **{first}**'s submission from the "
        f"{len(pool)} unmatched work(s) below.")
    st.caption("Filenames here are camera-roll noise — identify by the image. "
               "Click ⤢ to enlarge a work (to read a handwritten name), then "
               f"“This is {first}’s work” to assign it.")
    if not pool:
        st.success("No unmatched works remain for this assignment.")
        return

    ekey = f"_match_enlarge_{cls}_{assignment}"
    enlarged_ck = st.session_state.get(ekey)
    enlarged_row = next(
        (r for r in pool if r.get("csv_key") == enlarged_ck), None)

    # --- enlarged single-work view ------------------------------------------
    if enlarged_row is not None:
        if st.button("⤡ Shrink back to the grid", key="match_shrink",
                     width="stretch"):
            st.session_state.pop(ekey, None)
            st.rerun(scope="fragment")
        png, note = _work_thumbnail(enlarged_row, folder_ref,
                                    THUMB_ENLARGE_WIDTH)
        if png is not None:
            st.image(png, width="stretch")
        else:
            st.info(note)
        st.caption(_work_caption(enlarged_row))
        if st.button(f"✓ This is {first}’s work", key="match_assign_enlarged",
                     type="primary", width="stretch"):
            if assign_work(cls, assignment, enlarged_row["csv_key"],
                           student_key):
                st.session_state.pop(ekey, None)
                st.session_state["save_status"] = (
                    "ok", f"Assigned a work to {first} in “{assignment}”.")
            st.rerun()
        return

    # --- thumbnail grid ------------------------------------------------------
    # Widget keys use the row's ENUMERATE index (not pool.index(row), which
    # collides for two identical garbage rows → StreamlitDuplicateElementKey).
    cols_per_row = 3
    for i in range(0, len(pool), cols_per_row):
        cols = st.columns(cols_per_row)
        for offset, (col, row) in enumerate(zip(cols, pool[i:i + cols_per_row])):
            idx = i + offset
            ck = row.get("csv_key", "")
            with col:
                with st.container(border=True):
                    png, note = _work_thumbnail(row, folder_ref,
                                                THUMB_GRID_WIDTH)
                    if png is not None:
                        st.image(png, width="stretch")
                    else:
                        st.caption(f"🗎 {note}")
                    st.caption(_work_caption(row))
                    if st.button("⤢ Enlarge", key=f"match_enl_{idx}_{ck}",
                                 width="stretch"):
                        st.session_state[ekey] = ck
                        st.rerun(scope="fragment")
                    if st.button(f"✓ {first}’s work",
                                 key=f"match_assign_{idx}_{ck}",
                                 type="primary", width="stretch"):
                        if assign_work(cls, assignment, ck, student_key):
                            st.session_state["save_status"] = (
                                "ok", f"Assigned a work to {first} in "
                                      f"“{assignment}”.")
                        st.rerun()


def render_window2() -> None:
    st.subheader("2 · Roster & Submissions")
    if st.session_state["staging"]:
        st.info("Files are staged in Window 1 — commit them to update submission tracking.")
    st.caption(f"Roster for **{st.session_state['active_class']}** — switch class in the top bar.")

    # Roster intake: compact uploader + Apply on one line, in a visual box.
    with st.container(border=True, key="roster_intake"):
        rc = st.columns([3, 1], vertical_alignment="center")
        up = rc[0].file_uploader(
            "Google Classroom roster (.csv)", type=["csv"], key="up_roster",
            accept_multiple_files=False, label_visibility="collapsed")
        if rc[1].button("Apply", key="rost_file", type="primary",
                        width="stretch", disabled=up is None,
                        help="Load the uploaded Google Classroom roster as "
                             "this class's student list (matched by the "
                             "numeric ID in each email address)."):
            set_active_roster(sort_roster_gojuon(parse_roster(up)))
            persist()
            st.rerun()
        st.caption("⚠ Must be a Google Classroom CSV export (roster or "
                   "assignment grades) — students are matched by the numeric "
                   "ID in each email address. The list is sorted automatically "
                   "into hiragana (gojūon / あいうえお) order — by surname, then "
                   "given name. Use ↑/↓ to fine-tune.")
        mc = st.columns(2)
        if mc[0].button("➕ Add student", key="open_add_student",
                        width="stretch"):
            add_student_dialog()
        if mc[1].button("➖ Remove student", key="open_rm_student",
                        width="stretch",
                        disabled=not (st.session_state["roster"] or
                                      st.session_state["archived_students"]
                                      .get(st.session_state["active_class"]))):
            remove_student_dialog()

    roster = st.session_state["roster"]
    if not roster:
        st.caption("Upload a Classroom roster (matched by email-ID) or add "
                   "students manually.")
        return

    # Active assignment names are already class-scoped via assignment_table(),
    # and On is per-term, so missing-work tracking follows the selected term.
    table_rows = assignment_table()
    table_by_name = {r["name"]: r for r in table_rows}
    active_names = [r["name"] for r in table_rows
                    if assignment_on(r["name"])]
    subs = {n: submitter_keys(n) for n in active_names}
    st.caption(f"Edit names inline · set gender · ⚠ lists {current_term()} "
               "missing work · ↑/↓ reorder · Focus.")

    gender_choices = ["—"] + GENDER_OPTIONS
    for i, entry in enumerate(roster):
        key = entry["key"]
        missing = [n for n in active_names if key not in subs[n]]
        # columns: up, down, name (editable), gender, missing, focus
        # vertical_alignment centers the ↑/↓/Focus buttons against the row.
        row = st.columns([0.35, 0.35, 2.3, 1.2, 0.7, 0.8],
                         vertical_alignment="center")
        if row[0].button("↑", key=f"up_{i}") and i > 0:
            roster[i - 1], roster[i] = roster[i], roster[i - 1]
            persist()
            st.rerun()
        if row[1].button("↓", key=f"dn_{i}") and i < len(roster) - 1:
            roster[i + 1], roster[i] = roster[i], roster[i + 1]
            persist()
            st.rerun()

        # Editable full name (Surname/first). Keyed by stable student key so a
        # reorder never carries one student's edit onto another.
        new_name = row[2].text_input(
            "name", value=entry.get("name", key), key=f"rname_{key}",
            label_visibility="collapsed",
        )
        if new_name != entry.get("name", ""):
            entry["name"] = new_name
            s = gb().students.get(key)
            if s:
                s.name = new_name
            persist()

        # Gender selection -> drives report-comment pronouns in Window 3.
        cur_g = entry.get("gender", "") or "—"
        idx = gender_choices.index(cur_g) if cur_g in gender_choices else 0
        pick = row[3].selectbox("gender", gender_choices, index=idx,
                                key=f"gender_{key}", label_visibility="collapsed")
        new_g = "" if pick == "—" else pick
        if new_g != entry.get("gender", ""):
            entry["gender"] = new_g
            s = gb().students.get(key)
            if s:
                s.gender = new_g
            persist()

        # Missing-work indicator: a popover button that lists the actual names.
        with row[4].popover(f"⚠{len(missing)}" if missing else "✓",
                            width="stretch"):
            if missing:
                st.markdown(f"**Missing assignments ({current_term()})**")
                cls = st.session_state["active_class"]
                pool_by_asg = st.session_state["unmatched_works"].get(cls, {})
                for nm in missing:
                    # If graded works for this assignment pooled unmatched
                    # (Phase 3), the student's submission may be sitting there —
                    # offer visual matching instead of a dead "missing" line.
                    n_pool = len(pool_by_asg.get(nm, []))
                    if n_pool:
                        if st.button(f"🧩 {nm} — {n_pool} unmatched work"
                                     f"{'s' if n_pool != 1 else ''}",
                                     key=f"matchbtn_{key}_{nm}",
                                     width="stretch",
                                     help="Open the thumbnail grid and pick this "
                                          "student's submission from the "
                                          "unmatched works."):
                            match_works_dialog(key, nm)
                        continue
                    if is_excused(key, nm):
                        tag = " _(excused)_"
                    elif awaiting_grade(table_by_name.get(nm, {})):
                        # Folder-backed and still being graded — awaiting a grade
                        # from its folder, never counted as a 0. Once the folder's
                        # grading completes it drops the tag and shows as plain
                        # missing (a real 0).
                        tag = " _(awaiting grade)_"
                    else:
                        tag = ""
                    st.markdown(f"- {nm}{tag}")
                if active_names and len(missing) == len(active_names):
                    st.caption(f"No work recorded in {current_term()} at all. "
                               "Missing tasks count as **0** in the assessment "
                               "math — for a mid-year transfer, mark their "
                               "pre-arrival tasks Excused (Window 3 edit "
                               "panel) so nothing counts against them.")
                else:
                    st.caption("Missing tasks inject a **0** into the trend "
                               "graph and the grade calculation. Mark one "
                               "Excused in Window 3's edit panel to remove it "
                               "from the assessment instead; tasks awaiting "
                               "grade are never counted as 0.")
            else:
                st.caption(f"All {current_term()} assignments submitted.")

        if row[5].button("Focus", key=f"foc_{i}"):
            sid = key if key in gb().students else None
            if sid is None:
                for s in gb():
                    if s.name == key or student_label(s) == key:
                        sid = s.student_id
                        break
            st.session_state["focus_sid"] = sid or key
            st.session_state["edit_cell"] = None
            st.rerun()


# --------------------------------------------------------------------------
# WINDOW 3 - Evaluation cockpit
# --------------------------------------------------------------------------

def find_score(student, assignment, crit):
    for sc in student.criterion_scores(Criterion(crit), valid_only=False):
        if sc.assignment == assignment:
            return sc
    return None


def render_window3() -> None:
    st.subheader("3 · Evaluation Cockpit")
    if st.session_state["staging"]:
        st.warning("Math & charts are paused while files sit in staging. Commit or "
                   "remove the staged files in Window 1 to resume evaluation.")
        return
    sid = st.session_state["focus_sid"]
    student = find_student(sid) if sid else None
    if student is None:
        st.caption("Select a student in Window 2 to load their full details.")
        return

    # Name heading + roster email on one line: name left, a small click-to-copy
    # code chip pushed right (st.code keeps its native copy icon on hover). The
    # chip lives in a keyed container so theme_css can shrink it. Blank email
    # (not on the roster) → render the heading full-width, no empty chip.
    _email = student_email_for(student)
    if _email:
        _name_col, _email_col = st.columns([3, 2], vertical_alignment="bottom")
        _name_col.markdown(f"### {student_label(student)}")
        with _email_col.container(key="w3_email_chip"):
            st.code(_email, language=None)
    else:
        st.markdown(f"### {student_label(student)}")

    # --- Scrollable analytical TOP -------------------------------------------
    # Fixed-height scroll region keeps the cockpit compact. The AI comment deck
    # is rendered OUTSIDE this container (below) so its text areas escape the
    # BaseWeb flexbox collapse that affects text areas inside height containers.
    with st.container(height=int(st.session_state["prefs"].get("h3", 640)), border=False):
        st.markdown(f"**Selected tasks & marks — {current_term()}** — click a "
                    "mark to edit")
        st.caption("**Missing** counts as a mathematical 0 — click to enter a "
                   "mark or excuse it.  **⏳ Awaiting Grade** = folder still "
                   "being graded, excluded from the trend and grade math.")
        active_names = [r["name"] for r in assignment_table()
                        if assignment_on(r["name"])]
        table_by_name = {r["name"]: r for r in assignment_table()}
        student_asgs = {sc.assignment for b in student.scores.values() for sc in b}
        any_cell = False
        for asg in active_names:
            excused_now = is_excused(student.student_id, asg)
            scs = [sc for b in student.scores.values()
                   for sc in b if sc.assignment == asg]
            for sc in scs:
                any_cell = True
                crit = sc.criterion.value
                cc = st.columns([3, 1.2], vertical_alignment="center")
                cc[0].caption(f"{asg}  ·  {sc.timestamp:%b %d}")
                mark = "·" if (not sc.is_valid and sc.value == 0) else str(sc.value)
                late_mark = (" (late)"
                             if is_late(student.student_id, asg, crit, sc)
                             else "")
                exc_mark = " (excused)" if excused_now else ""
                if cc[1].button(f"{crit}: {mark}{late_mark}{exc_mark}",
                                key=f"cell_{asg}_{crit}", width="stretch"):
                    edit_grade_dialog(student.student_id, asg, crit)
            if scs or asg in student_asgs:
                continue
            r = table_by_name.get(asg)
            if not r:
                continue
            # Awaiting Grade: folder-backed work whose folder is still being
            # graded is simply not graded yet, not missing — one read-only row,
            # no invented 0, no effect on the trend or grade math (see the policy
            # note above missing_assignment_rows). Grades arrive via the
            # workspace. Once the folder's grading completes this student falls
            # through to the editable Missing = 0 rows below.
            if awaiting_grade(r):
                any_cell = True
                cc = st.columns([3, 1.2], vertical_alignment="center")
                cc[0].caption(f"{asg}  ·  {r['date']:%b %d}")
                cc[1].button("⏳ Awaiting Grade", key=f"cell_await_{asg}",
                             width="stretch", disabled=True)
                continue
            # Missing = 0: a selected, criteria-bearing task the student never
            # submitted shows one editable zero row per criterion (or an
            # "excused" chip once waived), so the injected 0s are visible and
            # one click away from being marked, zeroed-out, or excused.
            if r["criteria"] == "—":
                continue
            for crit in [c.strip() for c in r["criteria"].split(",")
                         if c.strip() in CRIT_ORDER]:
                any_cell = True
                cc = st.columns([3, 1.2], vertical_alignment="center")
                cc[0].caption(f"{asg}  ·  {r['date']:%b %d}")
                label = (f"{crit}: excused" if excused_now
                         else f"{crit}: 0 (missing)")
                if cc[1].button(label, key=f"cell_missing_{asg}_{crit}",
                                width="stretch"):
                    edit_grade_dialog(student.student_id, asg, crit)
        if not any_cell:
            st.caption("No marks yet for the active tasks.")

        _render_trend(student)
        _render_grade_panel(student)   # method dropdown + 4-criteria grades + overrides
        _render_comments(student)

    # --- Fixed BOTTOM (pinned, immune to flexbox collapse) -------------------
    _render_ai_deck(student)


@st.dialog("✏ Edit grade")
def edit_grade_dialog(sid: str, asg: str, crit: str) -> None:
    """Pop-up editor for one student's mark on one assignment/criterion.

    Replaces the old inline panel (which expanded below all the grades and
    vanished off-screen on short window heights). Shows the assignment name,
    a 0-8 dropdown, late/excused checklists, the task's comment, and
    Save / Cancel."""
    student = find_student(sid)
    if student is None:
        st.caption("Student not found.")
        return
    sc = find_score(student, asg, crit)
    is_missing = sc is None
    st.markdown(f"**{asg}**  ·  Criterion {crit} ({Criterion(crit).label})  ·  "
                f"{student_label(student)}"
                + ("  ·  _missing (counts as 0)_" if is_missing else ""))
    k = f"{sid}||{asg}||{crit}"   # unique widget keys per cell
    cur = int(sc.value) if sc else 0
    new_val = st.selectbox("Grade (0-8)", list(range(9)), index=cur,
                           key=f"ed_val_{k}")
    flags = st.columns(2)
    late = flags[0].checkbox("Late submission",
                             value=is_late(sid, asg, crit, sc), key=f"ed_late_{k}")
    excused = flags[1].checkbox(
        "Excused", value=is_excused(sid, asg), key=f"ed_exc_{k}",
        help="Removes this assignment from the trend graph, the final grade "
             "calculation and the AI prompt for this student (instead of "
             "counting it as a 0).")
    feedback = st.text_area(
        "Comment (this task's descriptive feedback)",
        value=(sc.comment if sc else ""), key=f"ed_fb_{k}",
        help="Edit or clear this to change what feeds the comments log and "
             "the generated overall comment.")
    b = st.columns(2)
    if b[0].button("Save", key=f"ed_save_{k}", type="primary",
                   width="stretch"):
        set_excused(sid, asg, excused)
        # Sticky-manual override (Phase B, D5): only materialise a late_flags
        # key when the teacher actually moved the checkbox off the effective
        # value, or the key is already manual. Nothing is written yet at this
        # point, so is_late() still reports the pre-save effective value — an
        # untouched checkbox is a no-op, so routine grade/comment Saves no
        # longer leave a redundant override. Once a key exists it stays manual
        # forever (updated here, never auto-deleted).
        lf = st.session_state["late_flags"]
        lk = late_key(sid, asg, crit)
        if lk in lf or late != is_late(sid, asg, crit, sc):
            lf[lk] = late
        if sc is not None:
            sc.value = int(new_val)
            sc.is_valid = True
            sc.comment = feedback
        elif not excused and (int(new_val) > 0 or feedback.strip()):
            # A real mark (or feedback) entered for previously-missing work:
            # materialise it as an ordinary score. A bare 0 needs no record —
            # the Missing = 0 pipeline already injects it.
            row = next((r for r in assignment_table() if r["name"] == asg), None)
            when = row["date"] if row else datetime.now()
            student.add_score(CriterionScore(
                criterion=Criterion(crit), value=int(new_val), timestamp=when,
                source="manual:missing-entry", assignment=asg, comment=feedback))
        persist()
        st.rerun()
    if b[1].button("Cancel", key=f"ed_cancel_{k}", width="stretch"):
        st.rerun()


def _trend_figure(student):
    """Build the progression-trend Plotly figure under the Missing=0 / Excused
    policy. Returns (figure, has_data). Shared by the Window 3 cockpit and the
    report-card docx export so both always show the same picture.

    Missing selected work is plotted as explicit band-0 points; excused
    assignments are removed from every trace."""
    import plotly.graph_objects as go

    excused = excused_assignments_for(student.student_id)
    zero_by_crit = {c: missing_zero_points(student, c) for c in CRIT_ORDER}

    fig = go.Figure()
    has = False
    # Build a shared, chronologically-ordered list of discrete x labels so the
    # axis lists assignments sequentially (evenly spaced) instead of plotting a
    # continuous datetime scale that leaks raw timestamps (e.g. 23:59:59:999).
    label_order = sorted(
        {sc.timestamp
         for c in CRIT_ORDER
         for sc in student.criterion_scores(Criterion(c), valid_only=True)
         if sc.include_in_report and sc.assignment not in excused}
        | {ts for pts in zero_by_crit.values() for ts, _ in pts},
        key=academic_sort_key,
    )
    categories = [f"{ts:%b %d}" for ts in label_order]
    for c in CRIT_ORDER:
        pts = [(sc.timestamp, sc.value)
               for sc in student.criterion_scores(Criterion(c), valid_only=True)
               if sc.include_in_report and sc.assignment not in excused]
        pts += zero_by_crit[c]
        pts.sort(key=lambda p: academic_sort_key(p[0]))
        if pts:
            has = True
            fig.add_trace(go.Scatter(
                x=[f"{p[0]:%b %d}" for p in pts], y=[p[1] for p in pts],
                mode="lines+markers", name=f"{c} {Criterion(c).label}",
                line=dict(color=CRIT_COLORS[c]),
            ))
    fig.update_layout(height=260, margin=dict(l=10, r=10, t=10, b=10),
                      yaxis=dict(range=[0, 8.4], title="Grade (0–8)"))
    # Discrete, evenly-spaced category axis (no continuous datetimes/timestamps).
    fig.update_xaxes(type="category", categoryorder="array",
                     categoryarray=categories)
    style_chart(fig)
    return fig, has


def _render_trend(student) -> None:
    fig, has = _trend_figure(student)
    st.markdown("**Progression trend**")
    if has:
        st.plotly_chart(fig, width="stretch")
    else:
        st.caption("No scored data to plot yet.")


def _grade_chip(value) -> str:
    """A read-only rounded grade box (see the ``.cam-grade-box`` CSS): matches
    the criterion selectboxes visually but is fill-less to signal it can't be
    edited. Used for the derived MYP Grade / School Grade."""
    txt = value if value is not None else "N/A"
    return f'<div class="cam-grade-box">{txt}</div>'


def _render_grade_panel(student) -> None:
    """Calculated grades for all four criteria + professional-judgment
    overrides, sitting directly under the progression graph.

    A dropdown selects which of the five grading algorithms computes the
    grades; every criterion recalculates as soon as the method changes. Each
    criterion shows its calculated band with an asterisk (e.g. ``6*``) as
    the default selectbox option; picking a number 0-8 locks an override, and
    picking the auto option releases it. ``N/A`` is shown when a criterion has
    no selected (valid, included) tasks."""
    st.markdown("**Grades & professional judgment**")
    # Per-student, per-term calculation method with a live "Auto" default. The
    # Auto option's label carries the current qualifying-assignment count and the
    # method it resolves to, so the teacher sees exactly what an untouched student
    # computes under. Picking a method PINS it in this term's map; picking Auto
    # RELEASES the pin (mirrors the grade selectboxes' pick-to-release pattern).
    sid = student.student_id
    n_qual = qualifying_assignment_count()
    auto_method = METHOD_60_40 if n_qual <= 15 else METHOD_WEIGHTED_MEDIAN
    cmap = calc_method_map()
    pinned = cmap.get(sid)
    AUTO_OPT = "<auto>"   # sentinel option; never a real method name
    options = [AUTO_OPT] + CALCULATION_METHODS
    calc_key = f"calc_{sid}"
    # Re-sync the keyed widget to the map each run (like the old dropdown re-synced
    # to calculation_method()), so switching terms — which empties this term's map —
    # reopens on Auto for everyone, and switching students shows that student's pin.
    st.session_state[calc_key] = pinned if pinned in CALCULATION_METHODS else AUTO_OPT

    def _calc_changed(sid=sid, calc_key=calc_key):
        choice = st.session_state.get(calc_key)
        m = calc_method_map()
        if choice in CALCULATION_METHODS:
            m[sid] = choice        # pin this student to an explicit method
        else:
            m.pop(sid, None)       # release back to the auto default
        persist()

    def _calc_label(m, auto_method=auto_method, n_qual=n_qual):
        if m == AUTO_OPT:
            unit = "assignment" if n_qual == 1 else "assignments"
            return f"Auto — {auto_method} ({n_qual} {unit} this term)"
        return METHOD_LABELS.get(m, m)

    st.selectbox("Calculation method", options, key=calc_key,
                 format_func=_calc_label, on_change=_calc_changed)
    overrides = st.session_state["final_override"].setdefault(student.student_id, {})
    cols = st.columns(4)
    for i, c in enumerate(CRIT_ORDER):
        res = aggregate_with_policy(student, c)
        calc = res.rounded_band if res else None
        auto_label = f"{calc}*  (auto)" if calc is not None else "N/A  (auto)"
        options = [auto_label] + [str(b) for b in range(9)]
        locked = overrides.get(c)
        index = (1 + locked) if isinstance(locked, int) and 0 <= locked <= 8 else 0
        with cols[i]:
            st.caption(f"Crit {c} · {Criterion(c).label}")
            pick = st.selectbox(
                f"crit_{c}", options, index=index,
                key=f"grade_{student.student_id}_{c}",
                label_visibility="collapsed",
            )
        if pick == auto_label:
            if c in overrides:
                del overrides[c]
                persist()
        else:
            v = int(pick)
            if overrides.get(c) != v:
                overrides[c] = v
                persist()
    if overrides:
        st.caption("Locked overrides: "
                   + ", ".join(f"{k}={v}" for k, v in sorted(overrides.items())))

    # ---- Effort + school report-card grades (one row) ----
    # Each figure is School-specific and shown only when enabled in ⚙ Settings →
    # Report-card grades (all off by default). Effort is the only editable
    # control; MYP/School recompute live from the shared student_term_grades
    # helper (also used by every export).
    rc = report_cfg()
    show_eff = bool(rc.get("show_effort", False))
    show_myp = bool(rc.get("show_myp_grade", False))
    show_sch = bool(rc.get("show_school_grade", False))
    if show_eff or show_myp or show_sch:
        cols = st.columns(sum((show_eff, show_myp, show_sch)))
        idx = 0
        if show_eff:
            emap = effort_map()
            stored = student_effort(student.student_id)
            elo, ehi = effort_bounds()
            opts = list(range(elo, ehi + 1))
            with cols[idx]:
                st.caption("Effort / English Use")
                pick = st.selectbox(
                    "effort", opts,
                    index=opts.index(stored) if stored in opts else 0,
                    key=f"effort_{student.student_id}",
                    label_visibility="collapsed",
                )
            if pick != emap.get(student.student_id, EFFORT_DEFAULT):
                emap[student.student_id] = int(pick)
                persist()
            idx += 1
        # Compute after the Effort edit so the chips reflect the new value this run.
        _n, _total, _effort, myp, gyo = student_term_grades(student)
        if show_myp:
            with cols[idx]:
                st.caption("MYP Grade")
                st.markdown(_grade_chip(myp), unsafe_allow_html=True)
            idx += 1
        if show_sch:
            with cols[idx]:
                st.caption("School Grade")
                st.markdown(_grade_chip(gyo), unsafe_allow_html=True)
            idx += 1


def clean_comment(text: str) -> str:
    """Strip rubric scaffolding labels ("Strengths:", "Areas to develop:" and
    similar) and tidy stray full-stops so the comments log reads as a clean,
    comma-separated list of observations."""
    if not text:
        return ""
    import re
    s = str(text)
    # Drop the leading labels wherever they appear (case-insensitive).
    s = re.sub(r"(?i)\b(strengths|areas to develop|growth areas?|"
               r"next steps?|per[- ]criterion)\s*:\s*", " ", s)
    # Split into fragments on full-stops / semicolons, trim, drop empties.
    parts = [p.strip(" .;,") for p in re.split(r"[.;]+", s)]
    parts = [p for p in parts if p]
    return ", ".join(parts)


def _render_comments(student) -> None:
    items = []
    for bucket in student.scores.values():
        for sc in bucket:
            if not sc.include_in_report:
                continue  # task not selected into the current term's assessment
            cleaned = clean_comment(sc.comment)
            if cleaned:
                items.append((sc.timestamp, cleaned))
    items.sort(key=lambda x: x[0], reverse=True)
    joined = ", ".join(c for _, c in items)
    st.markdown(f"**Comments log — {current_term()}** (newest first)")
    if not joined:
        st.caption("No comments yet.")
        return
    if len(joined) > 240:
        # No widget key on purpose: a keyed text_area caches its value in
        # session_state and ignores `value=` after the first run, which would
        # freeze this read-only log at the first focused student's comments.
        st.text_area("comments", value=joined, height=110,
                     label_visibility="collapsed")
        with st.expander("Enlarge comments"):
            st.write(joined)
    else:
        st.write(joined)


def _resolve_key(explicit: str, *env_names: str) -> str:
    """Prefer a key typed into the UI; otherwise read the environment."""
    if explicit and explicit.strip():
        return explicit.strip()
    for name in env_names:
        val = os.environ.get(name)
        if val:
            return val.strip()
    return ""


def _max_tokens(lc) -> int:
    # Allow generous headroom over the requested word limit.
    return max(256, int(lc.get("word_limit", 120)) * 4)


def _call_claude(prompt: str, lc, key: str):
    """Call the Anthropic Messages API. Returns (ok, text_or_error)."""
    key = _resolve_key(key, "ANTHROPIC_API_KEY")
    if not key:
        return False, "No Claude key (type one in the field or set ANTHROPIC_API_KEY)."
    try:
        import anthropic
    except ImportError:
        return False, "anthropic package not installed - run: pip install anthropic"
    model = (lc.get("model") or "").strip() or "claude-sonnet-4-6"
    try:
        client = anthropic.Anthropic(api_key=key)
        msg = client.messages.create(
            model=model,
            max_tokens=_max_tokens(lc),
            messages=[{"role": "user", "content": prompt}],
        )
        text = "".join(getattr(b, "text", "") for b in msg.content).strip()
        return (True, text) if text else (False, "Claude returned an empty response.")
    except Exception as exc:  # network / auth / model errors
        return False, f"Claude API error: {exc}"


def _call_gemini(prompt: str, lc, key: str):
    """Call the Google Gemini API. Returns (ok, text_or_error).

    Tries the current google-genai client first, then the legacy
    google-generativeai package, so either install works.
    """
    key = _resolve_key(key, "GOOGLE_API_KEY", "GEMINI_API_KEY")
    if not key:
        return False, "No Gemini key (type one in the field or set GOOGLE_API_KEY)."
    model = (lc.get("model") or "").strip() or "gemini-2.0-flash"

    # Preferred: google-genai
    try:
        from google import genai
        try:
            client = genai.Client(api_key=key)
            resp = client.models.generate_content(model=model, contents=prompt)
            text = (getattr(resp, "text", "") or "").strip()
            return (True, text) if text else (False, "Gemini returned an empty response.")
        except Exception as exc:
            return False, f"Gemini API error: {exc}"
    except ImportError:
        pass

    # Legacy: google-generativeai
    try:
        import google.generativeai as genai
        genai.configure(api_key=key)
        resp = genai.GenerativeModel(model).generate_content(prompt)
        text = (getattr(resp, "text", "") or "").strip()
        return (True, text) if text else (False, "Gemini returned an empty response.")
    except ImportError:
        return False, "Gemini SDK not installed - run: pip install google-genai"
    except Exception as exc:
        return False, f"Gemini API error: {exc}"


def call_llm_api(prompt: str, lc, key: str):
    """Dispatch to the configured provider. Returns (ok, text_or_error)."""
    if lc.get("provider") == "Gemini":
        return _call_gemini(prompt, lc, key)
    return _call_claude(prompt, lc, key)


@st.dialog("⚙ LLM parameters", width="large")
def llm_params_dialog() -> None:
    """LLM comment configuration, moved out of the cockpit into a modal.

    The pronoun selector is gone: pronouns are derived from the student's
    gender in Window 2. Inputs are batched in a form so they apply on submit."""
    lc = st.session_state["llm_cfg"]
    sid = st.session_state.get("focus_sid")
    student = find_student(sid) if sid else None
    if student is not None:
        st.caption(f"Pronouns are read from {student_label(student)}'s gender "
                   f"(Window 2) → **{pronouns_for(student)}**.")
    with st.form("llm_form"):
        mode = st.radio("Output", ["Clipboard prompt", "API call"],
                        index=0 if lc["mode"] == "Clipboard prompt" else 1,
                        key="lc_mode", horizontal=True)
        provider = st.radio("Model", ["Claude", "Gemini"],
                            index=0 if lc["provider"] == "Claude" else 1,
                            key="lc_prov", horizontal=True)
        name_labels = {
            "full": "Use full name in first sentence",
            "first": "Use first name in first sentence",
            "none": "Never use the student's name",
        }
        name_keys = ["full", "first", "none"]
        cur_nm = lc.get("name_mode", "first")
        name_mode = st.radio(
            "Name usage in comment", name_keys,
            index=name_keys.index(cur_nm) if cur_nm in name_keys else 1,
            format_func=lambda k: name_labels[k], key="lc_name_mode")
        s1 = st.columns(2)
        inc_s = s1[0].checkbox("Strengths", value=lc["inc_strengths"], key="lc_s")
        inc_g = s1[1].checkbox("Growth areas", value=lc["inc_growth"], key="lc_g")
        inc_n = s1[0].checkbox("Next steps", value=lc["inc_next"], key="lc_n")
        inc_c = s1[1].checkbox("Per-criterion", value=lc["inc_criteria"], key="lc_c")
        wl = st.number_input("Word limit", 40, 400, lc["word_limit"], 10, key="lc_wl")
        s2 = st.columns(2)
        n_s = s2[0].number_input("# strengths", 0, 5, lc["n_strengths"], key="lc_ns")
        n_g = s2[1].number_input("# growth", 0, 5, lc["n_growth"], key="lc_ng")
        # Tone is a single mutually-exclusive choice: formal and encouraging pull
        # the model in opposite directions, so selecting one clears the other
        # rather than sending it two contradictory instructions.
        tone_keys = ["neutral", "formal", "encouraging"]
        tone_labels = {
            "neutral": "Neutral (default)",
            "formal": "More formal language",
            "encouraging": "More encouraging language",
        }
        cur_tone = ("formal" if lc.get("tone_formal")
                    else "encouraging" if lc.get("tone_encouraging")
                    else "neutral")
        tone_choice = st.radio(
            "Tone", tone_keys, index=tone_keys.index(cur_tone),
            format_func=lambda k: tone_labels[k], key="lc_tone", horizontal=True)
        tone_formal = tone_choice == "formal"
        tone_enc = tone_choice == "encouraging"
        st.markdown("**Comment focus scope**")
        scope_keys = ["current", "include_past"]
        scope_labels = {
            "current": "Current term only",
            "include_past": "Current + past terms (compressed summaries)",
        }
        cur_scope = lc.get("focus_scope", "current")
        focus_scope = st.radio(
            "Context window", scope_keys,
            index=scope_keys.index(cur_scope) if cur_scope in scope_keys else 0,
            format_func=lambda k: scope_labels[k], key="lc_scope",
            help="'Include past terms' replaces raw earlier-term marks with each "
                 "term's comment, used purely for developmental trend context. "
                 "It uses the finalized cloud summary when one exists, else the "
                 "overall comment saved for that term in the app. Students "
                 "missing a prior term's comment get a soft alert — generation "
                 "still runs with whatever terms are available.")
        if term_index(current_term()) == 0:
            st.caption("ℹ You're in Term 1 — 'include past terms' has no "
                       "effect until Term 2.")
        inc_trend = st.checkbox(
            "Inject progression-trend summary",
            value=lc.get("inc_trend", True), key="lc_trend",
            help="Adds a one-line per-criterion trajectory (e.g. 'Criterion B "
                 "showed positive progression…') computed by the math engine. "
                 "At a word limit of 130+ the trajectory switches to a fuller "
                 "narration (spread, mid-term dips) automatically.")
        inc_late = st.checkbox(
            "Inject late-submission rate",
            value=lc.get("inc_late", True), key="lc_late",
            help="Adds a [SUBMISSION TIMELINESS] block giving the share of this "
                 "term's graded tasks submitted late (synced from the grading "
                 "workspace's Late column, with any CAM override applied). "
                 "Omitted entirely when nothing was late.")
        inc_missing = st.checkbox(
            "Inject missing-work rate",
            value=lc.get("inc_missing", True), key="lc_missing",
            help="Adds a [MISSING WORK] block giving the share of this term's "
                 "assessed tasks not submitted — counted from Window 3's "
                 "'0 (missing)' indicators (excused and still-being-graded work "
                 "excluded). Each missing task already counts as a 0 in the "
                 "grade math, so this only flags the habit; omitted entirely "
                 "when nothing is missing. Turn off when a zero was given for a "
                 "reason other than non-submission (e.g. academic dishonesty).")
        no_numbers = st.checkbox(
            "Never mention numeric grades in the comment",
            value=lc.get("no_numbers", False), key="lc_no_numbers",
            help="The comment describes achievement qualitatively ('excellent "
                 "grades', 'her grades can improve') and never states numbers "
                 "— no criterion grades, scores, marks or percentages. The "
                 "prompt still contains the numeric evidence; this only stops "
                 "the model from quoting it.")
        skip_existing = st.checkbox(
            "Skip students who already have a comment (whole-class only)",
            value=lc.get("skip_existing", True), key="lc_skip_existing",
            help="When 'Generate for whole class' runs, skip any student who "
                 "already has a non-empty overall comment for this term — "
                 "whether it was generated earlier or typed by hand. Leave on "
                 "to fill only the gaps (e.g. after a partial run that hit an "
                 "API quota); turn off to regenerate everyone, overwriting "
                 "existing comments. Does not affect the single-student "
                 "generate button.")
        default_model = ("claude-sonnet-4-6" if provider == "Claude"
                         else "gemini-2.0-flash")
        model = st.text_input("Model (API mode)",
                              value=(lc.get("model") or default_model), key="lc_model")
        api_key = st.text_input(
            "API key (API mode)", type="password",
            value=st.session_state.get("llm_api_key", ""), key="lc_key",
            help="Leave blank to use ANTHROPIC_API_KEY / GOOGLE_API_KEY "
                 "from your environment.")
        st.caption("Keys live in memory for this session only — never written "
                   "to disk or the database.")
        saved = st.form_submit_button("Apply", type="primary",
                                      width="stretch")
    if saved:
        lc.update({
            "mode": mode, "provider": provider, "name_mode": name_mode,
            "inc_strengths": inc_s, "inc_growth": inc_g,
            "inc_next": inc_n, "inc_criteria": inc_c,
            "word_limit": int(wl), "n_strengths": int(n_s),
            "n_growth": int(n_g), "tone_formal": tone_formal,
            "tone_encouraging": tone_enc, "model": model,
            "focus_scope": focus_scope, "inc_trend": inc_trend,
            "inc_late": inc_late, "inc_missing": inc_missing,
            "no_numbers": no_numbers,
            "skip_existing": skip_existing,
        })
        # Stash the key in a session var so it survives the dialog closing
        # (the keyed widget itself is purged once the modal is not rendered).
        st.session_state["llm_api_key"] = api_key
        st.session_state["dlg_llm"] = False
        # Regenerate the compiled/clipboard prompt in place so it reflects the
        # just-applied parameters. Without this the displayed prompt (and the
        # text the copy button reads) stays frozen at whatever it was when
        # "Compile prompt to clipboard" was last clicked. Only refresh when a
        # prompt already exists, so the expander doesn't pop open prematurely.
        if student is not None and st.session_state.get("llm_prompt"):
            lc["pronouns"] = pronouns_for(student)  # gender-derived, dynamic
            st.session_state["llm_prompt"] = compile_prompt(student, lc)
        st.rerun()
    if st.button("Close", key="llm_close"):
        st.session_state["dlg_llm"] = False
        st.rerun()


def _generate_class_comments(lc) -> tuple:
    """Generate an overall comment for every student in the active class via the
    API, honouring the parameter deck. When ``skip_existing`` is set, students
    who already have a non-empty comment for the current term are left untouched
    (so a re-run after a partial/quota-limited run fills only the gaps).
    Returns (n_ok, n_fail, n_skipped, first_error)."""
    n_ok = n_fail = n_skipped = 0
    first_error = ""
    skip_existing = lc.get("skip_existing", True)
    students = students_for_active_class()
    progress = st.progress(0.0, text="Generating class comments…")
    for i, stu in enumerate(students):
        if skip_existing and st.session_state["llm_response"].get(
                stu.student_id, "").strip():
            n_skipped += 1
            progress.progress((i + 1) / max(1, len(students)),
                              text=f"Generating class comments… ({i + 1}/{len(students)})")
            continue
        cfg = dict(lc)
        cfg["pronouns"] = pronouns_for(stu)
        prompt = compile_prompt(stu, cfg)
        ok, text = call_llm_api(prompt, cfg, st.session_state.get("llm_api_key", ""))
        if ok:
            st.session_state["llm_response"][stu.student_id] = text
            n_ok += 1
        else:
            n_fail += 1
            first_error = first_error or text
        progress.progress((i + 1) / max(1, len(students)),
                          text=f"Generating class comments… ({i + 1}/{len(students)})")
    progress.empty()
    persist()
    return n_ok, n_fail, n_skipped, first_error


def _render_ai_deck(student) -> None:
    st.markdown(f"**AI comment deck — {current_term()}**")
    lc = st.session_state["llm_cfg"]
    api_mode = lc["mode"] == "API call"

    # Multi-term context status: show what earlier-term material will feed the
    # prompt, and soft-alert (never block) when some of it is missing.
    if lc.get("focus_scope") == "include_past":
        cur_idx = term_index(current_term())
        if cur_idx == 0:
            st.caption("Scope includes past terms, but this is Term 1 — no "
                       "earlier terms exist yet.")
        else:
            have = past_term_context_for(student)
            if have:
                st.caption("Past-term context available: "
                           + ", ".join(t for t, _ in have) + " ✓")
            gaps = missing_past_terms(student)
            if gaps:
                notes = [t + ("" if student_worked_in_term(student, t)
                              else " (no work recorded — joined later?)")
                         for t in gaps]
                st.warning("No saved comment for " + ", ".join(notes)
                           + f" for {first_name_for(student)}. Generation "
                             "still works — the missing term(s) are simply "
                             "left out of the assessment context.")

    info = st.columns([1.4, 2], vertical_alignment="center")
    if info[0].button("⚙ LLM parameters", key="open_llm", width="stretch"):
        llm_params_dialog()
    tones = []
    if lc.get("tone_formal"):
        tones.append("formal")
    if lc.get("tone_encouraging"):
        tones.append("encouraging")
    tone_txt = (" · tone: " + "+".join(tones)) if tones else ""
    scope_txt = ("current+past" if lc.get("focus_scope") == "include_past"
                 else "current only")
    trend_txt = " · trend" if lc.get("inc_trend", True) else ""
    numbers_txt = " · no numbers" if lc.get("no_numbers") else ""
    info[1].caption(f"{lc['mode']} · {lc['provider']} · name: {lc.get('name_mode','first')} · "
                    f"pronouns {pronouns_for(student)} · scope: {scope_txt}"
                    f"{trend_txt}{tone_txt}{numbers_txt}")

    # Two generators: one student, or the whole class (API only). Both follow the
    # parameters set in the deck above.
    one_label = ("Generate for " + first_name_for(student)
                 if api_mode else "Compile prompt to clipboard")
    g = st.columns(2, vertical_alignment="center")
    if g[0].button(one_label, key="llm_go", type="primary",
                   width="stretch"):
        lc["pronouns"] = pronouns_for(student)  # gender-derived, dynamic
        prompt = compile_prompt(student, lc)
        st.session_state["llm_prompt"] = prompt
        st.session_state["llm_status"] = ("", "")
        if api_mode:
            ok, text = call_llm_api(prompt, lc, st.session_state.get("llm_api_key", ""))
            if ok:
                st.session_state["llm_response"][student.student_id] = text
                st.session_state["llm_status"] = ("ok", f"{lc['provider']} responded.")
            else:
                st.session_state["llm_status"] = ("error", text)
        st.rerun()
    if g[1].button("Generate for whole class", key="llm_go_all",
                   disabled=not api_mode, width="stretch",
                   help=("Generates a comment for every student in the active "
                         "class via the API. Switch Output to 'API call' in the "
                         "LLM parameters to enable.")):
        n_ok, n_fail, n_skipped, err = _generate_class_comments(lc)
        note = ""
        if lc.get("focus_scope") == "include_past" and term_index(current_term()):
            n_gap = sum(1 for s in students_for_active_class()
                        if missing_past_terms(s))
            if n_gap:
                note = (f" {n_gap} student(s) had no prior-term comment — "
                        "theirs used the available terms only.")
        skip_txt = f" Skipped {n_skipped} (already had one)." if n_skipped else ""
        if n_fail == 0:
            st.session_state["llm_status"] = (
                "ok", f"Generated {n_ok} comments for the class.{skip_txt}{note}")
        else:
            st.session_state["llm_status"] = (
                "error",
                f"{n_ok} done, {n_fail} failed.{skip_txt} First error: {err}{note}")
        st.rerun()

    kind, message = st.session_state.get("llm_status", ("", ""))
    if kind == "ok":
        st.success(message)
    elif kind == "error":
        st.warning(message)

    # Keyed-class CSS sets the exact box heights, bypassing Streamlit's 68px
    # text_area floor so the teacher can make them as small (or large) as they
    # like via the Settings sliders.
    hr = int(st.session_state["prefs"].get("h_remarks", 80))
    hc = int(st.session_state["prefs"].get("h_comment", 90))
    st.markdown(
        f"<style>[class*=\"st-key-rem_box\"] textarea{{height:{hr}px!important;min-height:{hr}px!important;}}"
        f"[class*=\"st-key-resp_box\"] textarea{{height:{hc}px!important;min-height:{hc}px!important;}}</style>",
        unsafe_allow_html=True)

    # Overall comment is the always-visible primary box. Teacher remarks tuck into
    # a popover opened by the small button to the right of the label.
    hdr = st.columns([5, 1], vertical_alignment="bottom")
    hdr[0].markdown(f"**Overall comment — {first_name_for(student)} · "
                    f"{current_term()} (editable)**")
    with hdr[1].popover("Remarks", width="stretch"):
        prev_rem = st.session_state["teacher_remarks"].get(student.student_id, "")
        rem = st.text_area(
            "Teacher remarks (notes that steer the generated comment)",
            value=prev_rem,
            key=f"rem_box_{student.student_id}")
        # Clearing a remark is a deliberate deletion — flag it for the mirror
        # shrink tripwire the same way the overall-comment box does.
        if prev_rem.strip() and not rem.strip():
            _mark_teacher_input_deleted()
        st.session_state["teacher_remarks"][student.student_id] = rem
    overall = st.text_area(
        "Overall comment", label_visibility="collapsed",
        value=st.session_state["llm_response"].get(student.student_id, ""),
        key=f"resp_box_{student.student_id}_{current_term().replace(' ', '_')}",
        help="Write, edit, or paste the final report-card comment here. It is "
             f"saved under {current_term()} — switch the term in the top bar "
             "to view or edit another term's comment for this student. The "
             "report-card pack collates one of these per student.")
    prev_overall = st.session_state["llm_response"].get(student.student_id, "")
    if overall != prev_overall:
        # Clearing a comment is a deliberate deletion: tell the cloud mirror so
        # its shrink tripwire lets the reduced slice reach disk (invariant 2).
        if prev_overall.strip() and not overall.strip():
            _mark_teacher_input_deleted()
        st.session_state["llm_response"][student.student_id] = overall
        persist()

    # In clipboard mode, expose the compiled prompt so it can be copied out.
    if not api_mode and st.session_state["llm_prompt"]:
        with st.expander("Compiled prompt (for manual/paste workflow)"):
            # No widget key here on purpose: a keyed text_area caches its value in
            # session_state and ignores `value=` on every rerun after the first,
            # which would freeze this display at the first compiled prompt. This
            # box is read-only (the copy button below reads llm_prompt directly),
            # so letting `value=` drive it each run keeps it in sync with the
            # latest compile.
            st.text_area("Compiled prompt", value=st.session_state["llm_prompt"],
                         height=140, label_visibility="collapsed")
            clipboard_button(st.session_state["llm_prompt"])


# --------------------------------------------------------------------------
# SYSTEM DELIVERABLES TRAY
# --------------------------------------------------------------------------

def render_tray() -> None:
    st.markdown("### System deliverables")

    # Save now + the width-capped status banner (see .st-key-save_status_bar
    # CSS). No Finalize button any more: every assignment, exam and comment is
    # explicitly tagged with its academic term, so there is nothing to lock in.
    sb = st.columns([1.2, 4.6], vertical_alignment="center")
    if sb[0].button("💾 Save now", key="save_now", width="stretch"):
        persist(show=True)
        st.rerun()
    kind, msg = st.session_state.get("save_status", ("", ""))
    if kind in ("ok", "error"):
        status_box = sb[1].container(key="save_status_bar")
        (status_box.success if kind == "ok" else status_box.error)(msg)
    st.caption("💾 saves the database now · comments and assignments are "
               "stored per term in the database, so any term stays editable "
               "at any time.")

    if st.session_state["staging"]:
        st.caption("Exports are paused while files sit in staging — commit them first.")
        return

    cols = st.columns(5, vertical_alignment="top")
    if not list(gb()):
        st.caption("Ingest a grading CSV to enable exports.")
        return

    # Exports are scoped to the active Class/Level, not the whole gradebook.
    active = st.session_state["active_class"]
    class_students = students_for_active_class()
    st.caption(f"Exports cover the **{active}** class only "
               f"({len(class_students)} student(s)). Click a button to build "
               "the file, then download it.")
    date_tag = date.today().isoformat()
    safe_cls = _safe_dirname(active)
    mime_xlsx = ("application/vnd.openxmlformats-officedocument."
                 "spreadsheetml.sheet")
    mime_docx = ("application/vnd.openxmlformats-officedocument."
                 "wordprocessingml.document")
    if not class_students:
        for i in range(4):
            cols[i].caption("No students in this class yet.")
    else:
        _export_slot(
            cols[0], "xlsx", f"Build Excel master — {active}",
            build_excel_bytes,
            f"CAM_master_{safe_cls}_{date_tag}.xlsx", mime_xlsx,
            ctx=(active, current_term()),
            help_text="Multi-tab workbook: final suggestions, raw scores, "
                      "assignment analytics (with class/subject/term).")
        _export_slot(
            cols[1], "pack", f"Build report-card pack — {active}",
            lambda: build_reportcards_docx(class_students),
            f"report_{safe_cls}_{date_tag}.docx", mime_docx,
            ctx=(active, current_term()),
            help_text="One document, one page per student: individual marks, "
                      "the progression graph, final criterion grades and "
                      "comments.")
        # ---- Mail-merge pack (col 2): one .docx per student, named by email --
        # Sits between the combined "pack" and "class comments" because it is
        # the OTHER report-card deliverable — the same per-student reports, but
        # split into individual files whose names are the students' emails,
        # ready for a batch-send script (each file mails itself to the address
        # it is named after). The skipped-student note is rendered full-width
        # below the row (it lists names and would not fit this narrow column).
        slots = st.session_state.setdefault("export_ready", {})
        zctx = (active, current_term(), len(class_students))
        if cols[2].button(f"Build mail-merge pack — {active}",
                          key="build_mailmerge", width="stretch",
                          help="A ZIP of individual report cards, each named "
                               "<student-email>.docx, for batch emailing. PDF "
                               "conversion happens when your send-script mails "
                               "them."):
            data, skipped = build_reportcards_zip(class_students)
            slots["mailmerge"] = {
                "ctx": zctx, "fname": f"mailmerge_{safe_cls}_{date_tag}.zip",
                "data": data, "skipped": skipped}
        mm_slot = slots.get("mailmerge")
        if mm_slot and mm_slot.get("ctx") == zctx:
            cols[2].download_button(
                f"⬇ {mm_slot['fname']}", data=mm_slot["data"],
                file_name=mm_slot["fname"], mime="application/zip",
                key="dl_mailmerge", width="stretch", type="primary")
        elif mm_slot:
            slots.pop("mailmerge", None)   # stale (class/term changed) — drop
        _export_slot(
            cols[3], "comments", f"Build class comments — {active}",
            lambda: build_class_comments_docx(class_students),
            f"comments_{safe_cls}_{date_tag}.docx", mime_docx,
            ctx=(active, current_term()),
            help_text="Compiles every student's saved comments (all terms) "
                      "into one document.")
    sid = st.session_state["focus_sid"]
    stu = find_student(sid) if sid else None
    if stu is not None:
        safe_stu = _safe_dirname(student_label(stu)).replace(" ", "_")
        _export_slot(
            cols[4], "single", f"Build report — {student_label(stu)}",
            lambda: build_single_docx(stu),
            f"report_{safe_stu}_{date_tag}.docx", mime_docx,
            ctx=(active, current_term(), sid))
    else:
        cols[4].caption("Select a student to export their single card.")

    # Mail-merge skipped-student note — full width, below the button row, since
    # it lists every left-out student and would overflow the narrow column.
    mm_slot = st.session_state.get("export_ready", {}).get("mailmerge")
    if class_students and mm_slot and mm_slot.get("ctx") == (
            active, current_term(), len(class_students)):
        skipped = mm_slot.get("skipped") or []
        if skipped:
            lines = "\n".join(f"- {lbl}: {why}" for lbl, why in skipped)
            st.warning(f"Mail-merge pack — {len(skipped)} student(s) left out; "
                       f"fix their roster email and rebuild:\n{lines}")
        else:
            st.caption(f"Mail-merge pack — all {len(class_students)} student(s) "
                       "included, one `<email>.docx` each.")


def _export_slot(col, kind: str, label: str, builder, fname: str, mime: str,
                 ctx: tuple, help_text: str = "") -> None:
    """Two-step deliverable export: build on click, then offer the download.

    Deliverables used to be rebuilt eagerly on EVERY rerun (each ``st.
    download_button`` computes its payload up front) — with per-student chart
    rendering that meant multi-minute freezes at boot and on every student
    click. Building only when asked makes reruns instant. The built file is
    kept until the class/term (or focused student) changes."""
    slots = st.session_state.setdefault("export_ready", {})
    if col.button(label, key=f"build_{kind}", width="stretch",
                  help=help_text or None):
        slots[kind] = {"ctx": ctx, "fname": fname, "data": builder()}
    slot = slots.get(kind)
    if slot and slot.get("ctx") == ctx:
        col.download_button(
            f"⬇ {slot['fname']}", data=slot["data"],
            file_name=slot["fname"], mime=mime,
            key=f"dl_{kind}", width="stretch", type="primary")
    elif slot:
        slots.pop(kind, None)   # stale (class/term/student changed) — drop it


# --------------------------------------------------------------------------
# Main layout
# --------------------------------------------------------------------------

def _class_dialog_body(edit: bool) -> None:
    """Body of the combined Add / Edit class dialog: one form for every class
    field (name, grade, MYP year, subject, master directory). Edit mode
    pre-fills from the ACTIVE class and saves through update_class(), so a
    rename moves every name-keyed store together; it also hosts the 👁 Watch
    scanner and the one-time 🔗 Connect Google Drive sign-in."""
    ac = active_class_dict() if edit else {}
    if edit and not ac:
        st.info("No class selected — switch to ➕ Add a class mode above to "
                "create one first.")
        return
    if edit:
        st.caption(f"Editing **{ac.get('name', '')}** — switch class in the "
                   "top bar to edit a different one. Renaming moves its "
                   "roster, units, grades, unit plan and data folder with it.")
    else:
        st.caption("Each class is its own tab — unit plan, grading CSVs and "
                   "roster are added per class. The MYP year sets the rubric "
                   "phase and the level the AI pitches comments at.")
    myp_opts = ["1", "2", "3", "4", "5"]
    cur_myp = str(ac.get("myp_year") or "") if edit else "1"
    if edit and cur_myp not in myp_opts:
        # Legacy class with no MYP year: keep it blank ("(not set)", rubric
        # phase falls back to the unit plan) unless the teacher picks one —
        # a Save must not silently stamp MYP 1 onto it.
        myp_opts = [cur_myp] + myp_opts
    # Per-mode widget keys: flipping the Edit/Add toggle must always start the
    # other mode's form from its own defaults, never from half-typed values.
    k = "edit" if edit else "add"
    with st.form(f"{k}class_form"):
        name = st.text_input("Class name", value=ac.get("name", ""),
                             key=f"clsdlg_name_{k}",
                             placeholder="e.g. 1-4, 2-Z, Year 7 Art")
        grade = st.text_input("Grade level (optional)",
                              value=ac.get("grade", ""),
                              key=f"clsdlg_grade_{k}",
                              placeholder="e.g. Year 7")
        myp = st.selectbox("MYP Year", myp_opts,
                           index=myp_opts.index(cur_myp) if cur_myp in myp_opts else 0,
                           key=f"clsdlg_myp_{k}",
                           format_func=lambda v: v or "(not set)",
                           help="Year 6→MYP1, 7→2, 8→3, 9→4, 10→5 (adjust to your school).")
        subject = st.text_input(
            "Subject", value=ac.get("subject", ""),
            key=f"clsdlg_subject_{k}",
            placeholder="e.g. Visual Arts, Design, Mathematics",
            help="The subject this class is assessed in. It frames the AI "
                 "comment prompts and appears on every exported report.")
        master = st.text_input(
            "Master directory (optional — local path or Drive Folder ID)",
            value=ac.get("master_dir", ""),
            key=f"clsdlg_master_{k}",
            placeholder=r"e.g. C:\Users\you\OneDrive\Y7 Art   or   1AbCdEfGhIjK...",
            help="The class's assignment home — its subfolders are scanned "
                 "into Window 1's assignment list automatically when you "
                 "save a new or changed directory (and again on every "
                 "🔄 Sync, or via 👁 Watch in Edit mode).")
        with st.expander("ℹ Master directory — local folder vs Google Drive",
                         expanded=False):
            st.markdown(
                "**👁 Watch** scans the master directory — every subfolder "
                "becomes an Assignment/Exam in Window 1's list. Rename or "
                "set deadlines via each row's ⋯ Manage menu; the real "
                "folders on disk/Drive are never touched.\n\n"
                "💡 **Recommended: use a local folder** (e.g. inside "
                "OneDrive) with the student files — no Google setup is "
                "needed. Using a **Google Drive folder ID** instead requires "
                "a one-time setup: in the "
                "[Google Cloud Console](https://console.cloud.google.com/) "
                "create a project, enable the "
                "[Google Drive API](https://console.cloud.google.com/apis/library/drive.googleapis.com), "
                "then create an OAuth **Desktop app** client under "
                "[Credentials](https://console.cloud.google.com/apis/credentials) "
                "and save the downloaded JSON as `credentials.json` inside "
                "`cam_grading_workspace/`; then click **🔗 Connect Google "
                "Drive** once and sign in with the account that can see the "
                "class folders (a consent screen still in **Testing** status "
                "needs that account added as a Test user). The sign-in is "
                "read-only (`drive.readonly`) and stores a reusable token — "
                "one time per device.\n\n"
                "**What each master can grade:** a **local folder** grades "
                "PDFs, images and video (export Office docs — Word / "
                "PowerPoint / Excel — to PDF first). A **Google Drive folder** "
                "grades those *plus* Google-native Docs / Sheets / Slides. "
                "Either way, work graded elsewhere can always be entered by "
                "hand via Window 1's ➕ Add assignment / exam.")
        saved = st.form_submit_button("Save changes" if edit else "Create class",
                                      type="primary", width="stretch")
    if saved:
        # Saving a new/changed master directory triggers an automatic Watch in
        # BOTH modes — pasting a folder into a brand-new class and onto an
        # existing one behave identically (Watch ADOPTS same-name manual rows
        # rather than duplicating them, so a class graded via ➕ Add
        # assignment/exam first is safe). Change detection runs before
        # update_class mutates the stored dict. On failure (e.g. a Drive ID
        # before the one-time sign-in) the save still stands and the status
        # banner names the fix.
        new_master = (master or "").strip()
        master_changed = (new_master != (ac.get("master_dir") or "").strip()
                          if edit else bool(new_master))
        ok = (update_class(ac["name"], name, grade, myp, subject, master)
              if edit else create_class(name, grade, myp, subject, master))
        if ok:
            if new_master and master_changed:
                watch_master_directory()
            st.rerun()
    drive_help = ("Drive watching needs a one-time Google sign-in (and a "
                  "credentials.json in cam_grading_workspace/). Opens the "
                  "sign-in page in your browser.")
    if edit:
        # Watch acts on the SAVED master directory (the form above only
        # applies on Save changes), so it sits outside the form.
        is_drive = bool(class_master_dir()) and not _master_is_local(
            class_master_dir())
        wc = st.columns(2 if is_drive else 1)
        if wc[0].button("👁 Watch", key="watch_master", width="stretch",
                        disabled=not class_master_dir(),
                        help="Rescan the saved master directory now — new "
                             "subfolders become assignments/exams in Window "
                             "1. (Runs automatically when you save a new "
                             "master directory, and on every 🔄 Sync.)"):
            watch_master_directory()
            st.rerun()
        if is_drive and wc[1].button(
                "🔗 Connect Google Drive", key="drive_signin", width="stretch",
                help=drive_help):
            launch_drive_signin()
    else:
        # The sign-in is device-wide (writes token.json), not class-specific —
        # offer it in Add mode too, so a first-time user can connect Drive
        # BEFORE creating their first Drive-backed class instead of having to
        # create it, then reopen this dialog in Edit mode just to sign in.
        if st.button("🔗 Connect Google Drive", key="drive_signin_add",
                     width="stretch", help=drive_help):
            launch_drive_signin()
        st.caption("Only needed when the master directory is a **Google Drive "
                   "folder ID** (one sign-in per device) — skip it for local "
                   "folders.")
    if st.button("Cancel", key="editclass_cancel" if edit else "addclass_cancel"):
        st.rerun()


@st.dialog("✎ Add / Edit class", width="large")
def class_dialog() -> None:
    """One modal for the whole class lifecycle. A two-button toggle at the top
    switches between **Edit current class** (the default — pre-filled, incl.
    rename, 👁 Watch and 🔗 Connect Google Drive) and **Add a class**; the
    selected mode's button renders red (type="primary") so the active mode is
    unmistakable. The opener resets the mode to Edit each time."""
    mode = st.session_state.get("class_dlg_mode", "edit")
    mc = st.columns(2)
    if mc[0].button("✎ Edit current class", key="clsdlg_mode_edit",
                    width="stretch",
                    type="primary" if mode == "edit" else "secondary"):
        if mode != "edit":
            st.session_state["class_dlg_mode"] = "edit"
            st.rerun(scope="fragment")
    if mc[1].button("➕ Add a class", key="clsdlg_mode_add", width="stretch",
                    type="primary" if mode == "add" else "secondary"):
        if mode != "add":
            st.session_state["class_dlg_mode"] = "add"
            st.rerun(scope="fragment")
    _class_dialog_body(edit=(mode == "edit"))


def _needs_first_boot_setup() -> bool:
    """Whether CAM must show the one-time first-boot setup panel (Phase 4).

    True only when this machine has **not chosen a data home yet**: no
    ``CAM_DB_PATH`` override, a **blank** ``db_custom_path`` pref, and the
    one-time ``setup_done`` marker unset. So a machine already pointed at a cloud
    folder boots exactly as before (non-blank path), a machine that explicitly
    chose *Start fresh* is not nagged again (``setup_done`` set), and a harness/
    env-configured path skips the panel — while a genuine fresh machine (no prefs
    file → blank path, no marker) always lands on the panel instead of silently
    booting the sample DB. See docs/CROSS_DEVICE_AND_DB_SAFETY_PLAN.md Phase 4."""
    if os.environ.get("CAM_DB_PATH", "").strip():
        return False
    prefs = st.session_state.get("prefs", DEFAULT_PREFS)
    if (prefs.get("db_custom_path") or "").strip():
        return False
    return not prefs.get("setup_done")


def _adopt_db_path(new_custom: str, note: str) -> None:
    """Commit a chosen data home and (re)hydrate against it (Phase 4).

    Shared by all three first-boot choices (adopt a discovered candidate, use
    another folder, start fresh). Writes the ``db_custom_path`` pref + the
    one-time ``setup_done`` marker, then clears ``db_loaded`` so the boot hydrate
    re-runs against the new path — reusing Phase 2's adopt behaviour: an existing
    database there is **loaded** (never overwritten), an absent one is created on
    first save. There is nothing worth pushing during bootstrap, so this never
    offers an overwrite. ``db_load_blocked`` is cleared so a stale quarantine
    from the pre-choice state does not linger (the hydrate re-diagnoses the new
    path)."""
    prefs = st.session_state["prefs"]
    prefs["db_custom_path"] = new_custom
    prefs["setup_done"] = True
    save_prefs(prefs)
    st.session_state["db_loaded"] = False
    st.session_state["db_load_blocked"] = None
    st.session_state["db_switch_pending"] = None
    st.session_state.pop("_boot_candidates", None)
    st.session_state["save_status"] = ("ok", note)
    st.rerun()


def _render_first_boot_setup() -> None:
    """One-time cross-device bootstrap panel shown before CAM loads anything.

    Rendered in place of the whole cockpit (main() returns right after) whenever
    ``_needs_first_boot_setup()`` — so this machine never silently boots the
    sample DB, and never persists, before the teacher points CAM at a data home.
    Offers: discovered cloud databases (convenience), a manual folder/path, and
    an explicit *Start fresh*. Every choice routes through ``_adopt_db_path``.
    See docs/CROSS_DEVICE_AND_DB_SAFETY_PLAN.md Phase 4."""
    st.title("Welcome to Criterion Assessment Metrics")
    st.caption("First, tell CAM where your gradebook lives. This is asked once "
               "per computer and saved to this device only "
               "(local_device_prefs.json).")

    # ---- Discovered candidates (probed once, cached for this panel) ----------
    if "_boot_candidates" not in st.session_state:
        with st.spinner("Looking for an existing CAM database in your OneDrive, "
                        "Google Drive and Dropbox folders…"):
            st.session_state["_boot_candidates"] = discover_db_candidates()
    candidates = st.session_state["_boot_candidates"]

    st.markdown("### 1 · Use my existing database")
    if candidates:
        st.caption("CAM found these databases in your cloud folders. Pick the "
                   "one with your real classes — the counts help you recognise "
                   "it. Nothing on disk is changed; CAM just points here.")
        for i, cand in enumerate(candidates):
            c = cand["counts"]
            row = st.columns([5, 1], vertical_alignment="center")
            row[0].markdown(
                f"**`{cand['folder']}`**  \n"
                f"{c.get('assignments', 0)} assignment(s) · "
                f"{c.get('students', 0)} student(s) · "
                f"{c.get('classes', 0)} class(es)")
            if row[1].button("Use this", key=f"boot_adopt_{i}", type="primary",
                             width="stretch"):
                _adopt_db_path(
                    cand["folder"],
                    f"Connected to the database at {cand['folder']}.")
    else:
        st.caption("No existing CAM database was found automatically in your "
                   "OneDrive / Google Drive / Dropbox folders. If you have one "
                   "there (or on a USB / network drive), point CAM at it below.")
    if st.button("🔄 Scan again", key="boot_rescan"):
        st.session_state.pop("_boot_candidates", None)
        st.rerun()

    # ---- Manual folder / file path (USB, network share, anything unlisted) ---
    st.markdown("---")
    st.markdown("### 2 · Use another folder")
    st.caption("Paste the folder that holds (or should hold) acm_database.json — "
               "a OneDrive/Drive folder, a USB drive, or a network share. An "
               "existing database there is loaded; an empty folder gets a new one "
               "on first save.")
    manual = st.text_input(
        "Database folder or .json path",
        key="boot_manual_path",
        placeholder=r"e.g. D:\CAM   or   \\server\share\CAM\acm_database.json")
    if st.button("Use this folder", key="boot_manual_use",
                 disabled=not manual.strip()):
        path = manual.strip()
        resolved = resolve_db_path(path)
        if db_file_state(resolved) == "ok":
            note = f"Connected to the existing database at {resolved}."
        else:
            note = f"CAM will use {resolved} (created on first save)."
        _adopt_db_path(path, note)

    # ---- Start fresh (explicit sample/demo boot) -----------------------------
    st.markdown("---")
    st.markdown("### 3 · Start fresh")
    st.caption("Begin with CAM's built-in sample gradebook on this computer "
               "only. You can point at a cloud folder later in ⚙ Settings.")
    if st.button("Start fresh with sample data", key="boot_fresh"):
        # Blank db_custom_path -> the sample acm_database.json beside app.py;
        # setup_done keeps this panel from reappearing next boot.
        _adopt_db_path("", "Started fresh with the built-in sample gradebook.")


def _render_db_quarantine_banner() -> None:
    """Full-width read-only-quarantine banner when the DB could not be loaded.

    Phase 1 (docs/CROSS_DEVICE_AND_DB_SAFETY_PLAN.md): when the boot guard set
    ``db_load_blocked``, the app is running demo/quarantine state with saving
    disabled. Make that loud and name the configured path so the teacher knows
    exactly which file/location to fix before restarting."""
    blocked = st.session_state.get("db_load_blocked")
    if not blocked:
        return
    path = blocked.get("path", "")
    reason = blocked.get("reason", "")
    if reason == "shrink-blocked":
        parked = blocked.get("parked", "")
        parked_note = (f" Your unsaved session was set aside at **`{parked}`** "
                       "for inspection." if parked else "")
        detail = (f"A save to **`{path}`** was refused because it would have "
                  "erased most of the database already on disk — the outgoing "
                  "session held far less data (far fewer assignments, roster "
                  "entries and graded students) than the file it would have "
                  "replaced, so CAM stopped the write." + parked_note)
    elif reason == "storage-missing":
        detail = (f"Your database location **`{path}`** is unavailable — the "
                  "drive or folder is not currently accessible (an unplugged "
                  "USB drive, an unmounted cloud folder, or a disconnected "
                  "network share).")
    elif reason == "empty-load":
        detail = (f"Your database file **`{path}`** contains data but loaded "
                  "with no students or assignments — it may be damaged.")
    else:  # "unreadable"
        detail = (f"Your database file **`{path}`** exists but could not be "
                  "read — it may be malformed, temporarily locked, or a cloud "
                  "file that has not finished downloading.")
    st.error(
        f"🛑 **Database not loaded — CAM is in read-only quarantine.**\n\n"
        f"{detail}\n\n"
        "**Nothing will be saved** while this banner is showing, so your real "
        "data on disk has **not** been changed. Fix the file or the path, then "
        "restart CAM. (Do not use Settings to point somewhere else and save — "
        "that would write this demo session out.)")


def _render_top_bar() -> None:
    """Title + class/level selector + Add/Edit class + Settings."""
    st.title("Criterion Assessment Metrics")
    st.caption("MYP Assessment · multi-term grade tracking, aggregation, "
               "grading workspace & reporting")
    # Controls on the LEFT — the top-right corner is occupied by Streamlit's own
    # Deploy / hamburger chrome, which would hide them.
    ctrl = st.columns([2.4, 1.2, 1.4, 1.2, 2.2], vertical_alignment="center")
    names = class_names()
    active = st.session_state["active_class"]
    sel = ctrl[0].selectbox(
        "Class / level", names,
        index=names.index(active) if active in names else 0,
        format_func=lambda n: class_label(next((c for c in st.session_state["classes"]
                                                 if c["name"] == n), {"name": n})),
        key="class_sel", label_visibility="collapsed",
    )
    if sel != active:
        st.session_state["active_class"] = sel
        st.session_state["focus_sid"] = None
        persist()
        st.rerun()
    # Active term: newly-ingested work is tagged to it; earlier terms become
    # "past" and feed the prompt engine as compressed finalized summaries.
    cur_term = current_term()
    term_sel = ctrl[1].selectbox(
        "Term", TERMS, index=term_index(cur_term),
        key="term_sel", label_visibility="collapsed",
        help="The term everything below works in: the On checkboxes pick this "
             "term's assessment set, and overall comments are saved under it. "
             "Switch terms any time — each keeps its own selection & comments.",
    )
    if term_sel != cur_term:
        st.session_state["active_term"] = term_sel
        persist()
        st.rerun()
    # Dialogs are opened by calling their @st.dialog functions directly (the
    # canonical Streamlit pattern). They inherit fragment behaviour, so widgets
    # inside only rerun the dialog, and dismissing (click-outside / ESC) does not
    # rerun the app — so a dialog never re-pops on a later unrelated rerun.
    if ctrl[2].button("✎ Add / Edit class", key="open_classdlg", width="stretch",
                      help="Edit the active class (rename, grade level, MYP "
                           "year, subject, master directory, 👁 Watch, Google "
                           "Drive sign-in) or add a new one — a red toggle "
                           "inside picks the mode (Edit is the default)."):
        st.session_state["class_dlg_mode"] = "edit"   # always open in Edit mode
        class_dialog()
    if ctrl[3].button("⚙ Settings", key="open_settings", width="stretch"):
        settings_dialog()


def main() -> None:
    st.set_page_config(page_title="Criterion Assessment Metrics", layout="wide")
    st.markdown(DENSE_CSS, unsafe_allow_html=True)  # dense layout for 1080p/4K
    st.markdown(theme_css(), unsafe_allow_html=True)  # grey/red theme surfaces
    init_state()
    # Phase-4 first-boot bootstrap: before any class/term context, sync, dedupe
    # or autosave can run, a machine that has not chosen a data home gets the
    # one-time setup panel — and nothing else. Returning here guarantees no
    # persist() fires against the (unchosen) path before the teacher picks.
    if _needs_first_boot_setup():
        _render_first_boot_setup()
        return
    ensure_class_context()
    ensure_term_context()
    # Heal any duplicate assignment records (same name+class) before rendering —
    # the timeline keys its per-row widgets on the assignment name, so a
    # duplicate would crash with StreamlitDuplicateElementKey. Runs every rerun
    # (cheap, idempotent) so a session that already loaded a duplicate recovers
    # without a full restart; persists only when it actually removes one.
    if _dedupe_assignments():
        persist()
    sync_active_into_scores()
    sync_roster_into_students()

    # Phase 1 sync lifecycle: one automatic global catch-up per session (OneDrive
    # / multi-machine changes), then — on every later rerun while a grading
    # session is open — the throttled scoped probe that ingests a fresh export
    # the moment the teacher returns from CGW.
    _run_session_start_sync()
    _run_active_launch_probe()

    prefs = st.session_state["prefs"]

    _render_top_bar()
    _render_db_quarantine_banner()

    # Three columns at the user's preferred width ratios. Each window body lives
    # in its own fixed-height scroll region (overflow-y:auto) so the whole
    # 3-column band scrolls internally and the deliverables tray below stays
    # permanently visible.
    w1, w2, w3 = st.columns([max(1, int(prefs["col_w1"])),
                             max(1, int(prefs["col_w2"])),
                             max(1, int(prefs["col_w3"]))])
    with w1:
        with st.container(height=int(prefs["h1"])):
            render_window1()
    with w2:
        with st.container(height=int(prefs["h2"])):
            render_window2()
    with w3:
        # Window 3 holds editable text areas (teacher remarks, overall comment).
        # A fixed-height st.container(height=...) makes BaseWeb collapse those
        # areas to a sliver, so this window uses a natural-height bordered
        # container; windows 1-2 keep their internal scroll for the long roster.
        with st.container(border=True):
            render_window3()

    render_tray()

    # Auto-save: mirror any state change from this run back to disk. Guarded so
    # a fresh, empty launch doesn't create a stub database file.
    if gb().students or gb().assignments or st.session_state["roster"]:
        persist()


if __name__ == "__main__":
    main()
